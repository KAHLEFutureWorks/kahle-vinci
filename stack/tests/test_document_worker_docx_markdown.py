#!/usr/bin/env python3
"""Regression checks for semantic DOCX-to-Markdown conversion."""
from __future__ import annotations

import io
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
WORKER_MAIN = ROOT / "document-worker" / "app" / "main.py"


def load_helpers() -> dict[str, Any]:
    src = WORKER_MAIN.read_text(encoding="utf-8")
    start = src.index("_MOJIBAKE_MARKERS")
    end = src.index("def _extract_text_pdf")
    ns: dict[str, Any] = {
        "io": io,
        "re": re,
        "Counter": Counter,
        "Optional": Optional,
        "List": List,
        "Tuple": Tuple,
        "Dict": Dict,
        "Any": Any,
        "HTTPException": RuntimeError,
        "USE_MARKITDOWN": False,
    }
    exec(src[start:end], ns)
    return ns


def test_docx_preserves_body_order_and_semantics() -> None:
    doc = Document()
    doc.add_heading("Interne Richtlinie", level=1)
    doc.add_paragraph("Ein kompakter Einleitungstext.")
    doc.add_paragraph("Erster Schritt", style="List Bullet")
    callout = doc.add_table(rows=1, cols=1)
    callout.cell(0, 0).text = "Wichtig: Freigabe erforderlich.\nVor Nutzung prüfen."
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feld"
    table.cell(0, 1).text = "Wert"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Aktiv | geprüft"
    doc.add_heading("Abschluss", level=2)
    buffer = io.BytesIO()
    doc.save(buffer)

    markdown = load_helpers()["_extract_text_docx"](buffer.getvalue())

    assert markdown.index("# Interne Richtlinie") < markdown.index("Ein kompakter")
    assert markdown.index("Ein kompakter") < markdown.index("- Erster Schritt")
    assert markdown.index("- Erster Schritt") < markdown.index("> Wichtig: Freigabe erforderlich.")
    assert "> Vor Nutzung prüfen." in markdown
    assert "<br>" not in markdown
    assert markdown.index("> Vor Nutzung prüfen.") < markdown.index("| Feld | Wert |")
    assert "| Status | Aktiv \\| geprüft |" in markdown
    assert markdown.index("| Status") < markdown.index("## Abschluss")


def test_docx_collapses_merged_layout_cells_without_duplicate_content() -> None:
    doc = Document()
    table = doc.add_table(rows=5, cols=4)
    table.cell(0, 0).merge(table.cell(0, 3)).text = "Rechtsgrund:\nEinmaliger Hinweistext."
    table.cell(1, 0).text = "1."
    table.cell(1, 1).merge(table.cell(1, 3)).text = "Governance"
    table.cell(2, 0).text = "1."
    table.cell(2, 1).merge(table.cell(2, 3)).text = "Einmaliger Abschnittstext."
    table.cell(3, 0).text = "8."
    table.cell(3, 1).text = "Funktion"
    table.cell(3, 2).text = "Name"
    table.cell(3, 3).text = "Kontakt"
    table.cell(4, 0).text = "8."
    table.cell(4, 1).text = "KI-Beauftragter"
    table.cell(4, 2).text = "Jan"
    table.cell(4, 3).text = "jan@example.test"
    buffer = io.BytesIO()
    doc.save(buffer)

    markdown = load_helpers()["_extract_text_docx"](buffer.getvalue())

    assert markdown.count("Einmaliger Hinweistext.") == 1
    assert markdown.count("Einmaliger Abschnittstext.") == 1
    assert "<br>" not in markdown
    assert "## 1. Governance" in markdown
    assert "| Funktion | Name | Kontakt |" in markdown
    assert "| 8. | Funktion" not in markdown

if __name__ == "__main__":
    test_docx_preserves_body_order_and_semantics()
    test_docx_collapses_merged_layout_cells_without_duplicate_content()
    print("document worker docx markdown tests passed")