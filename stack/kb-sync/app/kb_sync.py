from __future__ import annotations

import hashlib
import hmac
import json
import os
import signal
import sys
import threading
import time
import uuid
import re
from dataclasses import dataclass
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

import requests
from docx import Document
from pypdf import PdfReader
from watchdog.events import FileSystemEvent, FileSystemEventHandler
from watchdog.observers import Observer

try:
    from .bm25_snapshot import BM25Snapshot
except ImportError:  # pragma: no cover
    from bm25_snapshot import BM25Snapshot

try:
    from .canonical_inventory import CanonicalInventory, load_canonical_inventory, load_portal_inventory, write_inventory_report
    from .hybrid_sync import HYBRID_BUILD_ID, HybridIndexBuilder, HybridSyncError, QdrantHybridClient
except ImportError:  # pragma: no cover
    from canonical_inventory import CanonicalInventory, load_canonical_inventory, load_portal_inventory, write_inventory_report
    from hybrid_sync import HYBRID_BUILD_ID, HybridIndexBuilder, HybridSyncError, QdrantHybridClient


EMBEDDING_DIMENSION = 1024
DEFAULT_COLLECTIONS = ("kahleallgemein", "kahlekontext", "kahlerichtlinien")
DEFAULT_EXTENSIONS = (".md", ".txt", ".pdf", ".docx", ".csv")


@dataclass(frozen=True)
class Config:
    kb_root: Path
    state_path: Path
    qdrant_url: str
    ionos_base_url: str
    ionos_api_key: str
    embedding_model: str
    collections: tuple[str, ...]
    debounce_seconds: float
    reconcile_interval_seconds: int
    supported_extensions: tuple[str, ...]
    control_port: int
    hybrid_snapshot_path: Path
    internal_api_key: str
    portal_db_path: Path
    portal_files_root: Path


def env(name: str, default: str = "") -> str:
    return os.environ.get(name, default).strip()


def load_config() -> Config:
    # Der Token heisst lokal IONOS_API_TOKEN und auf dem Produktionsserver
    # IONOS_API_KEY. Beide Namen werden akzeptiert, damit derselbe Code in
    # beiden Umgebungen laeuft; IONOS_API_TOKEN hat Vorrang.
    api_key = env("IONOS_API_TOKEN") or env("IONOS_API_KEY")
    if not api_key:
        raise RuntimeError("IONOS_API_TOKEN or IONOS_API_KEY is required")

    collections = tuple(
        item.strip()
        for item in env("KB_SYNC_COLLECTIONS", ",".join(DEFAULT_COLLECTIONS)).split(",")
        if item.strip()
    )
    if not collections:
        raise RuntimeError("KB_SYNC_COLLECTIONS must contain at least one collection")

    extensions = tuple(
        item.strip().lower()
        for item in env("KB_SYNC_EXTENSIONS", ",".join(DEFAULT_EXTENSIONS)).split(",")
        if item.strip()
    )

    return Config(
        kb_root=Path(env("KB_ROOT", "/knowledgebases")),
        state_path=Path(env("KB_STATE_PATH", "/state/kb-sync-state.json")),
        qdrant_url=env("QDRANT_URL", "http://qdrant:6333").rstrip("/"),
        ionos_base_url=env("IONOS_OPENAI_BASE_URL", "https://openai.inference.de-txl.ionos.com/v1").rstrip("/"),
        ionos_api_key=api_key,
        embedding_model=env("IONOS_EMBEDDING_MODEL", "BAAI/bge-m3"),
        collections=collections,
        debounce_seconds=float(env("KB_SYNC_DEBOUNCE_SECONDS", "2")),
        reconcile_interval_seconds=int(env("KB_SYNC_RECONCILE_INTERVAL_SECONDS", "300")),
        supported_extensions=extensions,
        control_port=int(env("KB_SYNC_CONTROL_PORT", "8093")),
        hybrid_snapshot_path=Path(env("KB_HYBRID_SNAPSHOT_PATH", "/state/hybrid-bm25.json")),
        internal_api_key=env("KB_SYNC_INTERNAL_API_KEY"),
        portal_db_path=Path(env("PORTAL_DB_PATH", "/portal-data/wissensportal.sqlite3")),
        portal_files_root=Path(env("PORTAL_FILES_ROOT", "/portal-data/files")),
    )


class StateStore:
    def __init__(self, path: Path) -> None:
        self.path = path
        self.lock = threading.Lock()
        self.data: dict[str, dict[str, Any]] = {"collections": {}}
        self.load()

    def load(self) -> None:
        if not self.path.exists():
            return
        try:
            loaded = json.loads(self.path.read_text(encoding="utf-8"))
        except Exception as exc:
            print(f"state_load_failed path={self.path} error={exc}", flush=True)
            return
        if isinstance(loaded, dict):
            self.data = loaded
            self.data.setdefault("collections", {})

    def save(self) -> None:
        with self.lock:
            self.path.parent.mkdir(parents=True, exist_ok=True)
            tmp = self.path.with_suffix(self.path.suffix + ".tmp")
            tmp.write_text(json.dumps(self.data, ensure_ascii=False, indent=2), encoding="utf-8")
            tmp.replace(self.path)

    def collection(self, name: str) -> dict[str, Any]:
        collections = self.data.setdefault("collections", {})
        collection = collections.setdefault(name, {})
        collection.setdefault("files", {})
        return collection


class QdrantClient:
    def __init__(self, base_url: str) -> None:
        self.base_url = base_url

    def request(self, method: str, path: str, **kwargs: Any) -> dict[str, Any]:
        url = f"{self.base_url}{path}"
        response = requests.request(method, url, timeout=60, **kwargs)
        if response.status_code >= 400:
            raise RuntimeError(f"{method} {url} -> {response.status_code}: {response.text[:300]}")
        if not response.text:
            return {}
        return response.json()

    def ensure_collection(self, collection: str) -> None:
        response = requests.put(
            f"{self.base_url}/collections/{collection}",
            json={"vectors": {"size": EMBEDDING_DIMENSION, "distance": "Cosine"}},
            timeout=60,
        )
        if response.status_code not in {200, 409}:
            raise RuntimeError(f"PUT {self.base_url}/collections/{collection} -> {response.status_code}: {response.text[:300]}")

    def count(self, collection: str) -> int:
        result = self.request("POST", f"/collections/{collection}/points/count", json={"exact": True})
        return int(((result.get("result") or {}).get("count")) or 0)

    def delete_document(self, collection: str, doc_id: str) -> None:
        self.request(
            "POST",
            f"/collections/{collection}/points/delete?wait=true",
            json={"filter": {"must": [{"key": "doc_id", "match": {"value": doc_id}}]}},
        )

    def scroll_doc_ids(self, collection: str) -> set[str]:
        doc_ids: set[str] = set()
        offset: Any = None
        while True:
            body: dict[str, Any] = {"limit": 256, "with_payload": ["doc_id"], "with_vector": False}
            if offset is not None:
                body["offset"] = offset
            result = self.request("POST", f"/collections/{collection}/points/scroll", json=body).get("result") or {}
            for point in result.get("points") or []:
                payload = point.get("payload") or {}
                doc_id = payload.get("doc_id")
                if isinstance(doc_id, str):
                    doc_ids.add(doc_id)
            offset = result.get("next_page_offset")
            if offset is None:
                break
        return doc_ids

    def upsert_points(self, collection: str, points: list[dict[str, Any]]) -> None:
        if not points:
            return
        self.request("PUT", f"/collections/{collection}/points?wait=true", json={"points": points})


class IonosEmbeddings:
    def __init__(self, base_url: str, api_key: str, model: str) -> None:
        self.base_url = base_url
        self.api_key = api_key
        self.model = model

    def embed(self, texts: list[str]) -> list[list[float]]:
        response = requests.post(
            f"{self.base_url}/embeddings",
            headers={"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"},
            json={"model": self.model, "input": texts},
            timeout=120,
        )
        if response.status_code >= 400:
            raise RuntimeError(f"IONOS embeddings -> {response.status_code}: {response.text[:300]}")
        body = response.json()
        vectors = [item.get("embedding") for item in sorted(body.get("data") or [], key=lambda item: item.get("index", 0))]
        if len(vectors) != len(texts):
            raise RuntimeError(f"expected {len(texts)} embeddings, got {len(vectors)}")
        for vector in vectors:
            if not isinstance(vector, list) or len(vector) != EMBEDDING_DIMENSION:
                raise RuntimeError(f"unexpected embedding dimension: {len(vector) if isinstance(vector, list) else 0}")
        return vectors


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def stable_uuid(value: str) -> str:
    return str(uuid.uuid5(uuid.NAMESPACE_URL, value))


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Seite {index}]\n{text}")
    return "\n\n".join(pages)


def read_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt", ".csv"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    raise ValueError(f"unsupported file type: {suffix}")


def chunk_text(text: str, max_chars: int = 1800, overlap: int = 220) -> list[str]:
    clean = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    while "\n\n\n" in clean:
        clean = clean.replace("\n\n\n", "\n\n")
    chunks: list[str] = []
    index = 0
    while index < len(clean):
        end = min(len(clean), index + max_chars)
        if end < len(clean):
            cut = clean.rfind("\n\n", index, end)
            if cut > index + 500:
                end = cut
        part = clean[index:end].strip()
        if part:
            chunks.append(part)
        if end >= len(clean):
            break
        index = max(end - overlap, index + 1)
    return chunks


def is_rag_index_enabled(path: Path) -> bool:
    """Return whether a Markdown document is explicitly enabled for RAG.

    A source document can opt out with ``rag_index: false`` in its YAML
    frontmatter. This keeps navigation or draft files available in the
    filesystem without allowing them to outrank curated knowledge articles.
    """
    if path.suffix.lower() != ".md":
        return True

    try:
        header = path.read_bytes()[:8192].decode("utf-8-sig", errors="replace")
    except OSError:
        return False

    if not header.startswith("---"):
        return True

    frontmatter_end = re.search(r"^---\s*$", header[3:], flags=re.MULTILINE)
    if not frontmatter_end:
        return True

    frontmatter = header[3 : 3 + frontmatter_end.start()]
    return re.search(r"^\s*rag_index\s*:\s*(?:false|no|0)\s*(?:#.*)?$", frontmatter, flags=re.IGNORECASE | re.MULTILINE) is None

def is_supported_file(path: Path, root: Path, extensions: tuple[str, ...]) -> bool:
    if not path.is_file():
        return False
    if path.name.startswith(".") or path.name.startswith("~$"):
        return False
    if path.suffix.lower() not in extensions:
        return False
    if not is_rag_index_enabled(path):
        return False
    try:
        path.relative_to(root)
    except ValueError:
        return False
    return True


def iter_collection_files(root: Path, extensions: tuple[str, ...]) -> list[Path]:
    if not root.exists():
        return []
    return sorted(path for path in root.rglob("*") if is_supported_file(path, root, extensions))


def discover_collections(config: Config) -> tuple[str, ...]:
    discovered = {
        path.name
        for path in config.kb_root.iterdir()
        if path.is_dir() and not path.name.startswith(".")
    } if config.kb_root.exists() else set()
    return tuple(sorted(set(config.collections) | discovered))


class KnowledgebaseSync:
    def __init__(self, config: Config) -> None:
        self.config = config
        self.state = StateStore(config.state_path)
        self.qdrant = QdrantClient(config.qdrant_url)
        self.embeddings = IonosEmbeddings(config.ionos_base_url, config.ionos_api_key, config.embedding_model)
        self.reconcile_lock = threading.Lock()
        self.hybrid_lock = threading.Lock()
        self.hybrid_builder = HybridIndexBuilder(
            QdrantHybridClient(config.qdrant_url, EMBEDDING_DIMENSION), self.embeddings,
            snapshot_path=config.hybrid_snapshot_path,
        )

    def reconcile_all(self) -> None:
        for collection in discover_collections(self.config):
            self.reconcile_collection(collection)
        self.reconcile_hybrid()

    def _rebuild_hybrid_locked(
        self, inventory: CanonicalInventory, *, migration_candidates: int,
    ) -> dict[str, Any]:
        """Build the complete hybrid baseline while the caller holds hybrid_lock."""
        report = self.hybrid_builder.rebuild(list(inventory.documents))
        self.state.data.setdefault("hybrid", {}).update({
            "digest": inventory.digest,
            "status": "active",
            "collection": report["collection"],
            "documents": report["documents"],
            "chunks": report["chunks"],
            "migration_candidates": migration_candidates,
            "updated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        })
        self.state.save()
        print(
            f"hybrid_index_activated collection={report['collection']} "
            f"documents={report['documents']} chunks={report['chunks']}", flush=True,
        )
        return report

    def _bootstrap_hybrid_locked(self, inventory: CanonicalInventory) -> dict[str, Any]:
        legacy = load_canonical_inventory(self.config.kb_root)
        report_path = self.config.state_path.parent / "hybrid-migration-inventory.json"
        write_inventory_report(CanonicalInventory(
            inventory.documents, legacy.migration_candidates, inventory.digest,
        ), report_path)
        return self._rebuild_hybrid_locked(
            inventory, migration_candidates=len(legacy.migration_candidates),
        )

    def reconcile_hybrid(self, *, force: bool = False) -> None:
        with self.hybrid_lock:
            inventory = load_portal_inventory(self.config.portal_db_path, self.config.portal_files_root)
            legacy = load_canonical_inventory(self.config.kb_root)
            report_path = self.config.state_path.parent / "hybrid-migration-inventory.json"
            write_inventory_report(CanonicalInventory(
                inventory.documents, legacy.migration_candidates, inventory.digest,
            ), report_path)
            hybrid_state = self.state.data.setdefault("hybrid", {})
            if not inventory.documents:
                hybrid_state["status"] = "awaiting_canonical_documents"
                hybrid_state["migration_candidates"] = len(legacy.migration_candidates)
                self.state.save()
                return
            if not force and hybrid_state.get("digest") == inventory.digest:
                return
            self._rebuild_hybrid_locked(
                inventory, migration_candidates=len(legacy.migration_candidates),
            )

    def reconcile_hybrid_version(self, version_id: str) -> dict[str, Any]:
        """Update exactly one active portal version in the live hybrid index."""
        with self.hybrid_lock:
            inventory = load_portal_inventory(self.config.portal_db_path, self.config.portal_files_root)
            try:
                snapshot = BM25Snapshot.load(self.config.hybrid_snapshot_path)
            except FileNotFoundError:
                return self._bootstrap_hybrid_locked(inventory)
            if snapshot.build_id != HYBRID_BUILD_ID:
                raise HybridSyncError("hybrid_schema_migration_required")
            document = next((item for item in inventory.documents if item.version_id == version_id), None)
            if document is None:
                raise HybridSyncError("active_version_not_indexable")
            report = self.hybrid_builder.sync_document(document)
            hybrid_state = self.state.data.setdefault("hybrid", {})
            hybrid_state.update({
                "digest": inventory.digest, "status": "active",
                "collection": report["collection"],
                "updated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "last_incremental_version_id": version_id,
            })
            self.state.save()
            return report

    def reconcile_hybrid_document(self, document_id: str) -> dict[str, Any]:
        """Synchronize current state of one document, including deactivation."""
        with self.hybrid_lock:
            inventory = load_portal_inventory(self.config.portal_db_path, self.config.portal_files_root)
            try:
                snapshot = BM25Snapshot.load(self.config.hybrid_snapshot_path)
            except FileNotFoundError:
                if not inventory.documents:
                    raise HybridSyncError("hybrid_snapshot_unavailable")
                return self._bootstrap_hybrid_locked(inventory)
            if snapshot.build_id != HYBRID_BUILD_ID:
                raise HybridSyncError("hybrid_schema_migration_required")
            document = next((item for item in inventory.documents if item.document_id == document_id), None)
            if document:
                report = self.hybrid_builder.sync_document(document)
            else:
                collection = self.hybrid_builder.qdrant.active_collection(self.hybrid_builder.alias)
                self.hybrid_builder.qdrant.delete_document_versions(collection, document_id)
                report = {"collection": collection, "documents": 0, "chunks": 0}
            self.state.data.setdefault("hybrid", {}).update({
                "digest": inventory.digest, "status": "active", "collection": report["collection"],
                "updated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "last_incremental_document_id": document_id,
            })
            self.state.save()
            return report

    def reconcile_collection(self, collection: str) -> None:
        with self.reconcile_lock:
            collection_root = self.config.kb_root / collection
            self.qdrant.ensure_collection(collection)
            state_collection = self.state.collection(collection)
            state_files: dict[str, Any] = state_collection.setdefault("files", {})
            files = iter_collection_files(collection_root, self.config.supported_extensions)
            seen = {path.relative_to(collection_root).as_posix() for path in files}
            seen_doc_ids = {f"{collection}/{rel_path}" for rel_path in seen}

            force = bool(state_files) and self.qdrant.count(collection) == 0
            if force:
                print(f"reindex_forced collection={collection} reason=qdrant_empty_state_present", flush=True)

            for rel_path in sorted(set(state_files) - seen):
                doc_id = f"{collection}/{rel_path}"
                self.qdrant.delete_document(collection, doc_id)
                del state_files[rel_path]
                print(f"deleted collection={collection} file={rel_path}", flush=True)

            for doc_id in sorted(self.qdrant.scroll_doc_ids(collection) - seen_doc_ids):
                if not doc_id.startswith(f"{collection}/"):
                    continue
                self.qdrant.delete_document(collection, doc_id)
                print(f"deleted_orphan collection={collection} doc_id={doc_id}", flush=True)

            for path in files:
                rel_path = path.relative_to(collection_root).as_posix()
                digest = sha256_file(path)
                previous = state_files.get(rel_path) or {}
                if not force and previous.get("sha256") == digest:
                    continue
                self.index_file(collection, collection_root, path, digest)

            state_collection["last_reconcile_at"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            self.state.save()

    def index_file(self, collection: str, collection_root: Path, path: Path, digest: str) -> None:
        rel_path = path.relative_to(collection_root).as_posix()
        doc_id = f"{collection}/{rel_path}"
        try:
            text = extract_text(path)
            chunks = chunk_text(text)
        except Exception as exc:
            print(f"extract_failed collection={collection} file={rel_path} error={exc}", flush=True)
            return

        self.qdrant.delete_document(collection, doc_id)
        points: list[dict[str, Any]] = []
        for offset in range(0, len(chunks), 8):
            batch = chunks[offset : offset + 8]
            vectors = self.embeddings.embed(batch)
            for batch_index, vector in enumerate(vectors):
                chunk_index = offset + batch_index
                content = batch[batch_index]
                payload = {
                    "content": content,
                    "text": content,
                    "kb": collection,
                    "doc_id": doc_id,
                    "source_path": rel_path,
                    "chunk_index": chunk_index,
                    "metadata": {
                        "doc_id": doc_id,
                        "kb": collection,
                        "source_path": rel_path,
                        "filename": path.name,
                        "chunk_index": chunk_index,
                        "source": "kb-sync",
                    },
                }
                points.append({"id": stable_uuid(f"{doc_id}#{chunk_index}"), "vector": vector, "payload": payload})

        self.qdrant.upsert_points(collection, points)
        files = self.state.collection(collection).setdefault("files", {})
        files[rel_path] = {
            "sha256": digest,
            "chunks": len(points),
            "updated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        }
        print(f"indexed collection={collection} file={rel_path} chunks={len(points)}", flush=True)


class ReindexRequestHandler(BaseHTTPRequestHandler):
    sync_service: KnowledgebaseSync

    def _json_response(self, status: int, payload: dict[str, Any]) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        if self.path != "/health":
            self._json_response(404, {"ok": False, "error": "not_found"})
            return
        self._json_response(200, {"ok": True, "collections": list(discover_collections(self.sync_service.config)), "hybrid": self.sync_service.state.data.get("hybrid", {})})

    def _handle_sparse_query(self) -> None:
        expected = self.sync_service.config.internal_api_key
        supplied = self.headers.get("X-API-Key", "")
        if not expected or not hmac.compare_digest(supplied, expected):
            self._json_response(401, {"error": "internal_api_key_required"})
            return
        try:
            length = min(int(self.headers.get("Content-Length", "0")), 8192)
            payload = json.loads(self.rfile.read(length) or b"{}")
            query = str(payload.get("query") or "").strip()
            if not query or len(query) > 2000:
                self._json_response(422, {"error": "invalid_query"})
                return
            snapshot = BM25Snapshot.load(self.sync_service.config.hybrid_snapshot_path)
            self._json_response(200, snapshot.encode_query(query))
        except FileNotFoundError:
            self._json_response(503, {"error": "hybrid_snapshot_unavailable"})
        except Exception as exc:
            print(f"sparse_query_failed error={exc}", flush=True)
            self._json_response(500, {"error": "sparse_query_failed"})

    def _is_authorized(self) -> bool:
        expected = self.sync_service.config.internal_api_key
        supplied = self.headers.get("X-API-Key", "")
        return bool(expected and hmac.compare_digest(supplied, expected))

    def do_POST(self) -> None:
        if self.path == "/hybrid/sparse-query":
            self._handle_sparse_query()
            return
        if not self._is_authorized():
            self._json_response(401, {"ok": False, "error": "internal_api_key_required"})
            return
        if self.path == "/hybrid/versions/sync":
            try:
                length = min(int(self.headers.get("Content-Length", "0")), 8192)
                payload = json.loads(self.rfile.read(length) or b"{}")
                version_id = str(payload.get("version_id") or "").strip()
                if not version_id:
                    self._json_response(422, {"ok": False, "error": "version_id_required"})
                    return
                started = time.monotonic()
                report = self.sync_service.reconcile_hybrid_version(version_id)
                self._json_response(200, {
                    "ok": True, "scope": "version", "version_id": version_id,
                    "chunks": report["chunks"],
                    "duration_ms": round((time.monotonic() - started) * 1000),
                })
            except Exception as exc:
                print(f"hybrid_incremental_sync_failed error={exc}", flush=True)
                self._json_response(500, {"ok": False, "error": str(exc)})
            return
        if self.path == "/hybrid/documents/sync":
            try:
                length = min(int(self.headers.get("Content-Length", "0")), 8192)
                payload = json.loads(self.rfile.read(length) or b"{}")
                document_id = str(payload.get("document_id") or "").strip()
                if not document_id:
                    self._json_response(422, {"ok": False, "error": "document_id_required"})
                    return
                started = time.monotonic()
                report = self.sync_service.reconcile_hybrid_document(document_id)
                self._json_response(200, {
                    "ok": True, "scope": "document", "document_id": document_id,
                    "documents": report["documents"], "chunks": report["chunks"],
                    "duration_ms": round((time.monotonic() - started) * 1000),
                })
            except Exception as exc:
                print(f"hybrid_incremental_document_sync_failed error={exc}", flush=True)
                self._json_response(500, {"ok": False, "error": str(exc)})
            return
        if self.path == "/reindex-all":
            try:
                started = time.monotonic()
                self.sync_service.reconcile_hybrid(force=True)
                self._json_response(200, {
                    "ok": True, "scope": "hybrid",
                    "duration_ms": round((time.monotonic() - started) * 1000),
                })
            except Exception as exc:
                print(f"hybrid_reindex_request_failed error={exc}", flush=True)
                self._json_response(500, {"ok": False, "error": "hybrid_reindex_failed"})
            return
        if self.path != "/reindex":
            self._json_response(404, {"ok": False, "error": "not_found"})
            return
        try:
            length = min(int(self.headers.get("Content-Length", "0")), 8192)
            payload = json.loads(self.rfile.read(length) or b"{}")
            collection = str(payload.get("collection") or "").strip()
            if collection not in discover_collections(self.sync_service.config):
                self._json_response(404, {"ok": False, "error": "unknown_collection"})
                return
            started = time.monotonic()
            self.sync_service.reconcile_collection(collection)
            self._json_response(
                200,
                {
                    "ok": True,
                    "collection": collection,
                    "path": str(payload.get("path") or ""),
                    "duration_ms": round((time.monotonic() - started) * 1000),
                },
            )
        except Exception as exc:
            print(f"reindex_request_failed error={exc}", flush=True)
            self._json_response(500, {"ok": False, "error": str(exc)})

    def log_message(self, _format: str, *_args: Any) -> None:
        return


class DebouncedHandler(FileSystemEventHandler):
    def __init__(self, sync: KnowledgebaseSync) -> None:
        self.sync = sync
        self.timer: threading.Timer | None = None
        self.lock = threading.Lock()

    def on_any_event(self, event: FileSystemEvent) -> None:
        if event.is_directory:
            return
        with self.lock:
            if self.timer:
                self.timer.cancel()
            self.timer = threading.Timer(self.sync.config.debounce_seconds, self.sync.reconcile_all)
            self.timer.daemon = True
            self.timer.start()


def main() -> int:
    config = load_config()
    missing_roots = [name for name in config.collections if not (config.kb_root / name).exists()]
    if missing_roots:
        raise RuntimeError(f"knowledgebase directories missing: {', '.join(missing_roots)}")

    sync = KnowledgebaseSync(config)
    sync.reconcile_all()

    stop = threading.Event()
    observer = Observer()
    handler = DebouncedHandler(sync)
    observer.schedule(handler, str(config.kb_root), recursive=True)
    observer.start()

    ReindexRequestHandler.sync_service = sync
    control_server = ThreadingHTTPServer(("0.0.0.0", config.control_port), ReindexRequestHandler)
    control_thread = threading.Thread(target=control_server.serve_forever, daemon=True)
    control_thread.start()

    def shutdown(_signum: int, _frame: Any) -> None:
        stop.set()
        control_server.shutdown()

    signal.signal(signal.SIGTERM, shutdown)
    signal.signal(signal.SIGINT, shutdown)
    print(
        "kb_sync_started "
        f"root={config.kb_root} collections={','.join(discover_collections(config))} model={config.embedding_model}",
        flush=True,
    )

    try:
        while not stop.wait(config.reconcile_interval_seconds):
            sync.reconcile_all()
    finally:
        observer.stop()
        observer.join(timeout=10)
        control_server.server_close()
        sync.state.save()
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"kb_sync_fatal error={exc}", file=sys.stderr, flush=True)
        raise SystemExit(1)
