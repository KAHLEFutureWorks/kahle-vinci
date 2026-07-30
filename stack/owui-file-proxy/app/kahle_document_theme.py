from __future__ import annotations

import html
import io
import json
import re
from pathlib import Path
from typing import Any


BLUE = "0069B3"
BLUE_DARK = "00457A"
INK = "111111"
MUTED = "8A8A8A"
LINE = "E1E1E1"
LIGHT = "F6F7F8"
RED = "CC0000"
WHITE = "FFFFFF"


def _config(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _company(config: dict[str, Any]) -> dict[str, str]:
    raw = config.get("company") if isinstance(config.get("company"), dict) else {}
    defaults = {
        "name": "Autohaus KAHLE GmbH & Co. KG",
        "address": "Am Leineufer 49 | 30419 Hannover",
        "register": "Amtsgericht Hannover HRA 18616",
        "registered_office": "Sitz der Gesellschaft in Hannover",
        "general_partner": "KAHLE Verwaltungs GmbH",
        "general_partner_register": "Amtsgericht Hannover HRB 219226",
        "managing_directors": "Karl-Heinz Kahle | Lukas Kahle | Thomas Keller",
        "vat_id": "DE 115 699 464",
        "tax_id": "25 / 213 / 21103",
        "phone": "0511 / 27 999-0",
        "fax": "0511 / 27 999-99",
        "accounting_fax": "05032 / 804-290",
        "website": "www.kahle.de",
        "email": "info@kahle.de",
        "bank": "Hannoversche Volksbank eG",
        "iban": "DE79 2519 0001 0028 9647 00",
        "bic": "VOHADE2HXXX",
    }
    return {key: str(raw.get(key) or value) for key, value in defaults.items()}


def _inline_segments(text: str) -> list[tuple[str, bool, bool]]:
    result: list[tuple[str, bool, bool]] = []
    cursor = 0
    pattern = re.compile(r"(\*\*|__)(.+?)\1|(?<!\*)\*([^*\n]+)\*(?!\*)")
    for match in pattern.finditer(text or ""):
        if match.start() > cursor:
            result.append((text[cursor:match.start()], False, False))
        if match.group(2) is not None:
            result.append((match.group(2), True, False))
        else:
            result.append((match.group(3), False, True))
        cursor = match.end()
    if cursor < len(text or "") or not result:
        result.append(((text or "")[cursor:], False, False))
    return result


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return len(cells) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown(content: str, title: str) -> list[dict[str, Any]]:
    raw = (content or "").splitlines()
    blocks: list[dict[str, Any]] = []
    index = 0
    duplicate_title_skipped = False
    while index < len(raw):
        stripped = raw[index].strip()
        if not stripped:
            blocks.append({"type": "space"})
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code: list[str] = []
            index += 1
            while index < len(raw) and not raw[index].strip().startswith("```"):
                code.append(raw[index].rstrip())
                index += 1
            index += 1
            blocks.append({"type": "code", "language": language, "text": "\n".join(code)})
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            text = heading.group(2).strip()
            if not duplicate_title_skipped and text.casefold() == (title or "").strip().casefold():
                duplicate_title_skipped = True
            else:
                blocks.append({"type": "heading", "level": min(len(heading.group(1)), 3), "text": text})
            index += 1
            continue
        if "|" in stripped and index + 1 < len(raw) and _is_table_separator(raw[index + 1].strip()):
            rows = [_table_row(stripped)]
            index += 2
            while index < len(raw) and "|" in raw[index] and raw[index].strip():
                rows.append(_table_row(raw[index]))
                index += 1
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            blocks.append({"type": "table", "rows": rows})
            continue
        if stripped.startswith(">"):
            quote: list[str] = []
            while index < len(raw) and raw[index].strip().startswith(">"):
                quote.append(raw[index].strip()[1:].strip())
                index += 1
            blocks.append({"type": "callout", "text": " ".join(quote)})
            continue
        if re.match(r"^[-*+]\s+", stripped):
            items: list[str] = []
            while index < len(raw) and re.match(r"^[-*+]\s+", raw[index].strip()):
                items.append(re.sub(r"^[-*+]\s+", "", raw[index].strip()))
                index += 1
            blocks.append({"type": "bullets", "items": items})
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            items: list[str] = []
            while index < len(raw) and re.match(r"^\d+[.)]\s+", raw[index].strip()):
                items.append(re.sub(r"^\d+[.)]\s+", "", raw[index].strip()))
                index += 1
            blocks.append({"type": "numbers", "items": items})
            continue
        paragraph = [stripped]
        index += 1
        while index < len(raw):
            nxt = raw[index].strip()
            if not nxt or nxt.startswith(("#", ">", "```")) or re.match(r"^[-*+]\s+|^\d+[.)]\s+", nxt):
                break
            if "|" in nxt and index + 1 < len(raw) and _is_table_separator(raw[index + 1].strip()):
                break
            paragraph.append(nxt)
            index += 1
        blocks.append({"type": "paragraph", "text": " ".join(paragraph)})
    while blocks and blocks[-1]["type"] == "space":
        blocks.pop()
    return blocks


def _clear_container(container: Any) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _paragraph_rule(paragraph: Any, color: str, size: int = 8, top: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement("w:top" if top else "w:bottom")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "1")
    edge.set(qn("w:color"), color)
    borders.append(edge)


def _set_run(run: Any, *, name: str = "VW Head", size: float | None = None, bold: bool | None = None, color: str | None = None, all_caps: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if all_caps:
        caps = OxmlElement("w:caps")
        run._element.get_or_add_rPr().append(caps)


def _add_field(paragraph: Any, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field = OxmlElement("w:instrText")
    field.set(qn("xml:space"), "preserve")
    field.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, field, separate, end])
    _set_run(run, name="Arial", size=6.5, color=MUTED)


def _style_document(document: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    specs = {
        "Normal": ("VW Head", 11.25, False, INK, 1.45, 0, 7),
        "Title": ("VW Head", 30, True, INK, 1.0, 0, 10),
        "Subtitle": ("VW Head", 12, False, INK, 1.35, 0, 4),
        "Heading 1": ("VW Head", 16.5, True, BLUE, 1.1, 18, 7),
        "Heading 2": ("VW Head", 11.5, True, INK, 1.15, 14, 5),
        "Heading 3": ("VW Head", 10.5, True, INK, 1.15, 11, 4),
        "List Bullet": ("VW Head", 11.25, False, INK, 1.4, 0, 3),
        "List Number": ("VW Head", 11.25, False, INK, 1.4, 0, 3),
    }
    for name, spec in specs.items():
        try:
            style = document.styles[name]
        except Exception:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font, size, bold, color, spacing, before, after = spec
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.line_spacing = spacing
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
    for name, size, color in (("Kahle Kicker", 8.5, BLUE), ("Kahle Lead", 12.5, INK), ("Kahle Caption", 7.5, MUTED), ("Kahle Code", 8.5, INK)):
        try:
            style = document.styles[name]
        except Exception:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial" if name != "Kahle Lead" else "VW Head"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name == "Kahle Kicker"
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.35


def _configure_section(document: Any, company: dict[str, str], logo_path: Path) -> None:
    from copy import deepcopy
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    document.settings.odd_and_even_pages_header_footer = False
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.45)
        section.bottom_margin = Cm(2.75)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.55)
        section.footer_distance = Cm(0.55)

        header = section.header
        _clear_container(header)
        table = header.add_table(rows=1, cols=2, width=Cm(17.4))
        table.autofit = False
        table.columns[0].width = Cm(7.2)
        table.columns[1].width = Cm(10.2)
        left, right = table.rows[0].cells
        left.width, right.width = Cm(7.2), Cm(10.2)
        left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        if logo_path.exists():
            run = left.paragraphs[0].add_run()
            run.add_picture(str(logo_path), width=Cm(2.35))
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Cm(0)
        run = right_p.add_run(company["name"] + "\n" + company["website"])
        _set_run(run, name="Arial", size=7, color=MUTED)
        rule = header.add_paragraph()
        rule.paragraph_format.space_before = Cm(0.05)
        rule.paragraph_format.space_after = Cm(0)
        _paragraph_rule(rule, BLUE, size=18)

        footer = section.footer
        _clear_container(footer)
        rule = footer.add_paragraph()
        rule.paragraph_format.space_after = Cm(0.08)
        _paragraph_rule(rule, LINE, size=6, top=True)
        foot = footer.add_table(rows=1, cols=3, width=Cm(17.4))
        foot.autofit = False
        widths = (Cm(5.6), Cm(6.2), Cm(5.6))
        for index, cell in enumerate(foot.rows[0].cells):
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell, 0, 0, 0, 90)
        columns = [
            [company["name"], company["address"], f"Telefon {company['phone']}", f"Fax {company['fax']}", f"{company['website']} | {company['email']}"],
            [company["register"], company["registered_office"], f"Pers. haftend: {company['general_partner']}", company["general_partner_register"], f"Geschäftsführung: {company['managing_directors']}", f"USt-IdNr. {company['vat_id']} | Steuer-Nr. {company['tax_id']}"],
            [company["bank"], f"IBAN {company['iban']}", f"BIC {company['bic']}", f"Fax Rechnungswesen {company['accounting_fax']}"],
        ]
        for cell, lines in zip(foot.rows[0].cells, columns):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Cm(0)
            for line_index, line in enumerate(lines):
                run = p.add_run(("\n" if line_index else "") + line)
                _set_run(run, name="Arial", size=5.9, bold=line_index == 0, color=INK if line_index == 0 else MUTED)
        page_p = foot.rows[0].cells[2].add_paragraph()
        page_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_p.paragraph_format.space_before = Cm(0.05)
        run = page_p.add_run("Seite ")
        _set_run(run, name="Arial", size=6.5, color=MUTED)
        _add_field(page_p, "PAGE")
        run = page_p.add_run(" von ")
        _set_run(run, name="Arial", size=6.5, color=MUTED)
        _add_field(page_p, "NUMPAGES")

        # The legacy DOCX template contains separate first/even-page parts.
        # Mirror the canonical chrome into every variant explicitly.
        for variant in (section.first_page_header, section.even_page_header):
            _clear_container(variant)
            for child in header._element:
                variant._element.append(deepcopy(child))
        for variant in (section.first_page_footer, section.even_page_footer):
            _clear_container(variant)
            for child in footer._element:
                variant._element.append(deepcopy(child))


def _add_inline_docx(paragraph: Any, text: str, *, color: str = INK, size: float | None = None) -> None:
    for value, bold, italic in _inline_segments(text):
        run = paragraph.add_run(value)
        _set_run(run, size=size, bold=bold, color=color)
        run.italic = italic


def _add_docx_table(document: Any, rows: list[list[str]]) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    if not rows:
        return
    width = len(rows[0])
    table = document.add_table(rows=len(rows), cols=width)
    table.autofit = False
    usable = 17.4
    weights = [1.0] * width
    if width >= 2:
        weights[1] = 2.4
    total = sum(weights)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.width = Cm(usable * weights[column_index] / total)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, 105, 120, 105, 120)
            if row_index == 0:
                _set_cell_shading(cell, INK)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if column_index == width - 1 and re.search(r"\d", value) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            _set_run(run, name="Arial", size=8.5, bold=row_index == 0, color=WHITE if row_index == 0 else INK, all_caps=row_index == 0)
    caption = document.add_paragraph("Tabelle · Automatisch aus dem Inhalt übernommen", style="Kahle Caption")
    caption.paragraph_format.space_before = Cm(0.08)


def _add_docx_callout(document: Any, text: str) -> None:
    from docx.shared import Cm

    lower = text.casefold()
    signal = lower.startswith(("warnung", "frist", "achtung"))
    fill = RED if signal else BLUE
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, 220, 260, 220, 260)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Cm(0)
    label = "WICHTIG" if signal else "KERNAUSSAGE"
    run = p.add_run(label + "\n")
    _set_run(run, name="Arial", size=7.5, bold=True, color=WHITE, all_caps=True)
    _add_inline_docx(p, text, color=WHITE, size=12)


def render_docx(content: str, title: str, template_path: Path, logo_path: Path, config_path: Path, generated_at: str) -> bytes | None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm
    except Exception:
        return None
    try:
        document = Document(str(template_path)) if template_path.exists() else Document()
        body = document._element.body
        section_props = None
        for child in list(body):
            if child.tag.endswith("sectPr"):
                section_props = child
            else:
                body.remove(child)
        if section_props is not None and section_props.getparent() is None:
            body.append(section_props)
        config = _config(config_path)
        company = _company(config)
        _style_document(document)
        _configure_section(document, company, logo_path)

        kicker = document.add_paragraph(style="Kahle Kicker")
        _add_inline_docx(kicker, "KAHLE-DOKUMENT \u00b7 ERSTELLT MIT KAHLE-VINCI", color=BLUE, size=8.5)
        title_p = document.add_paragraph(style="Title")
        _add_inline_docx(title_p, (title or "Dokument").strip() or "Dokument", color=INK, size=30)

        meta = document.add_table(rows=1, cols=3)
        meta.autofit = False
        labels = (("DOKUMENT", (title or "Dokument").strip()), ("DATUM", generated_at.split(" ")[0]), ("ERSTELLT VON", "KAHLE-Vinci"))
        for cell, (label, value) in zip(meta.rows[0].cells, labels):
            _set_cell_margins(cell, 100, 0, 100, 140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Cm(0)
            run = p.add_run(label + "\n")
            _set_run(run, name="Arial", size=7, bold=True, color=INK, all_caps=True)
            run = p.add_run(value)
            _set_run(run, name="Arial", size=7.5, color=MUTED)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Cm(0.15)
        _paragraph_rule(spacer, LINE, size=6)

        blocks = parse_markdown(content, title)
        lead_used = False
        for block in blocks:
            kind = block["type"]
            if kind == "space":
                continue
            if kind == "heading":
                document.add_paragraph(block["text"], style=f"Heading {block['level']}")
            elif kind == "paragraph":
                style = "Kahle Lead" if not lead_used else "Normal"
                p = document.add_paragraph(style=style)
                _add_inline_docx(p, block["text"], color=INK, size=12.5 if not lead_used else 11.25)
                lead_used = True
            elif kind == "bullets":
                items = block["items"][:7]
                for item_index, item in enumerate(items):
                    p = document.add_paragraph()
                    p.paragraph_format.keep_with_next = item_index < len(items) - 1
                    p.paragraph_format.left_indent = Cm(0.55)
                    p.paragraph_format.first_line_indent = Cm(-0.35)
                    _add_inline_docx(p, f"\u2022  {item}")
            elif kind == "numbers":
                items = block["items"]
                for item_index, item in enumerate(items):
                    p = document.add_paragraph()
                    p.paragraph_format.keep_with_next = item_index < len(items) - 1
                    p.paragraph_format.left_indent = Cm(0.7)
                    p.paragraph_format.first_line_indent = Cm(-0.45)
                    _add_inline_docx(p, f"{item_index + 1}.  {item}")
            elif kind == "table":
                _add_docx_table(document, block["rows"])
            elif kind == "callout":
                _add_docx_callout(document, block["text"])
            elif kind == "code":
                for line in block["text"].splitlines() or [""]:
                    document.add_paragraph(line, style="Kahle Code")

        core = document.core_properties
        core.title = (title or "Dokument").strip()
        core.author = "KAHLE-Vinci"
        out = io.BytesIO()
        document.save(out)
        return out.getvalue()
    except Exception:
        return None


def _rl_inline(text: str) -> str:
    safe = html.escape(text or "")
    safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
    safe = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", safe)
    return safe


def render_pdf(content: str, title: str, logo_path: Path, config_path: Path, generated_at: str) -> bytes | None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.pdfgen import canvas as rl_canvas
    except Exception:
        return None
    try:
        config = _config(config_path)
        company = _company(config)
        blue = colors.HexColor("#" + str(config.get("colors", {}).get("blue", BLUE)))
        ink = colors.HexColor("#" + str(config.get("colors", {}).get("anthracite", INK)))
        muted = colors.HexColor("#" + str(config.get("colors", {}).get("muted", MUTED)))
        line = colors.HexColor("#" + LINE)
        light = colors.HexColor("#" + LIGHT)
        red = colors.HexColor("#" + RED)
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="Kicker", fontName="Helvetica-Bold", fontSize=8.5, leading=10, textColor=blue, spaceAfter=6, uppercase=True))
        styles.add(ParagraphStyle(name="KTitle", fontName="Helvetica-Bold", fontSize=29, leading=31, textColor=ink, spaceAfter=12))
        styles.add(ParagraphStyle(name="Lead", fontName="Helvetica", fontSize=12, leading=18, textColor=ink, spaceAfter=10))
        styles.add(ParagraphStyle(name="H1", fontName="Helvetica-Bold", fontSize=16.5, leading=19, textColor=blue, spaceBefore=15, spaceAfter=7, keepWithNext=True))
        styles.add(ParagraphStyle(name="H2", fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=ink, spaceBefore=12, spaceAfter=5, keepWithNext=True))
        styles.add(ParagraphStyle(name="H3", fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=ink, spaceBefore=10, spaceAfter=4, keepWithNext=True))
        styles.add(ParagraphStyle(name="BodyK", fontName="Helvetica", fontSize=10.5, leading=15.5, textColor=ink, spaceAfter=7))
        styles.add(ParagraphStyle(name="BulletK", parent=styles["BodyK"], leftIndent=12, firstLineIndent=-8, bulletIndent=0, spaceAfter=3))
        styles.add(ParagraphStyle(name="CaptionK", fontName="Helvetica", fontSize=7.5, leading=9, textColor=muted, spaceBefore=4, spaceAfter=7))
        styles.add(ParagraphStyle(name="MetaLabel", fontName="Helvetica-Bold", fontSize=6.8, leading=8, textColor=ink))
        styles.add(ParagraphStyle(name="TableHeader", fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=colors.white))
        styles.add(ParagraphStyle(name="MetaValue", fontName="Helvetica", fontSize=7.3, leading=9, textColor=muted))
        styles.add(ParagraphStyle(name="Callout", fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=colors.white))
        styles.add(ParagraphStyle(name="CalloutLabel", fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=colors.white))
        styles.add(ParagraphStyle(name="CodeK", fontName="Courier", fontSize=8, leading=10, textColor=ink, backColor=light, leftIndent=6, rightIndent=6, spaceAfter=2))

        out = io.BytesIO()
        doc = SimpleDocTemplate(out, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm, topMargin=28*mm, bottomMargin=35*mm, title=title, author="KAHLE-Vinci")
        width, height = A4

        def draw_page(canvas: Any, _doc: Any) -> None:
            canvas.saveState()
            if logo_path.exists():
                try:
                    canvas.drawImage(str(logo_path), 18*mm, height-18*mm, width=25*mm, height=10*mm, preserveAspectRatio=True, mask="auto", anchor="sw")
                except Exception:
                    pass
            canvas.setFillColor(muted)
            canvas.setFont("Helvetica", 7)
            canvas.drawRightString(width-18*mm, height-11*mm, company["name"])
            canvas.drawRightString(width-18*mm, height-15*mm, company["website"])
            canvas.setStrokeColor(blue)
            canvas.setLineWidth(2.2)
            canvas.line(0, height-21*mm, width, height-21*mm)
            canvas.setStrokeColor(line)
            canvas.setLineWidth(0.6)
            canvas.line(0, 27*mm, width, 27*mm)
            canvas.setFont("Helvetica-Bold", 5.6)
            canvas.setFillColor(ink)
            canvas.drawString(18*mm, 22.5*mm, company["name"])
            canvas.setFont("Helvetica", 5.2)
            canvas.setFillColor(muted)
            left_lines = [company["address"], f"Telefon {company['phone']} | Fax {company['fax']}", f"{company['website']} | {company['email']}"]
            middle_lines = [company["register"], f"Pers. haftend: {company['general_partner']}", company["general_partner_register"], f"Geschäftsführung: {company['managing_directors']}"]
            right_lines = [company["bank"], f"IBAN {company['iban']}", f"BIC {company['bic']}", f"Seite {canvas.getPageNumber()} von {getattr(canvas, '_kahle_page_count', canvas.getPageNumber())}"]
            for idx, value in enumerate(left_lines):
                canvas.drawString(18*mm, (19.5-idx*3.2)*mm, value)
            for idx, value in enumerate(middle_lines):
                canvas.drawString(73*mm, (22.5-idx*3.2)*mm, value)
            for idx, value in enumerate(right_lines):
                canvas.drawString(145*mm, (22.5-idx*3.2)*mm, value)
            canvas.restoreState()

        story: list[Any] = [Paragraph("KAHLE-DOKUMENT · ERSTELLT MIT KAHLE-VINCI", styles["Kicker"]), Paragraph(_rl_inline(title or "Dokument"), styles["KTitle"])]
        meta = [[Paragraph("DOKUMENT<br/><font color='#8A8A8A'>"+html.escape(title or "Dokument")+"</font>", styles["MetaLabel"]), Paragraph("DATUM<br/><font color='#8A8A8A'>"+html.escape(generated_at.split(" ")[0])+"</font>", styles["MetaLabel"]), Paragraph("ERSTELLT VON<br/><font color='#8A8A8A'>KAHLE-Vinci</font>", styles["MetaLabel"])]]
        meta_table = Table(meta, colWidths=[77*mm, 42*mm, 55*mm])
        meta_table.setStyle(TableStyle([("LINEABOVE",(0,0),(-1,-1),0.5,line),("LINEBELOW",(0,0),(-1,-1),0.5,line),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),8),("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7)]))
        story.extend([meta_table, Spacer(1, 7)])

        lead_used = False
        for block in parse_markdown(content, title):
            kind = block["type"]
            if kind == "space":
                continue
            if kind == "heading":
                story.append(Paragraph(_rl_inline(block["text"]), styles[f"H{block['level']}"]))
            elif kind == "paragraph":
                story.append(Paragraph(_rl_inline(block["text"]), styles["Lead"] if not lead_used else styles["BodyK"]))
                lead_used = True
            elif kind in ("bullets", "numbers"):
                for number, item in enumerate(block["items"], start=1):
                    prefix = "&bull;" if kind == "bullets" else f"{number}."
                    story.append(Paragraph(f"{prefix} {_rl_inline(item)}", styles["BulletK"]))
            elif kind == "table":
                data = [[Paragraph(_rl_inline(cell), styles["TableHeader"] if row_index == 0 else styles["BodyK"]) for cell in row] for row_index, row in enumerate(block["rows"])]
                columns = len(data[0])
                weights = [1.0]*columns
                if columns >= 2:
                    weights[1] = 2.4
                total = sum(weights)
                table = Table(data, colWidths=[174*mm*w/total for w in weights], repeatRows=1, hAlign="LEFT")
                table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),ink),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,light]),("LINEBELOW",(0,0),(-1,-1),0.4,line),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6)]))
                story.extend([table, Paragraph("Tabelle · Automatisch aus dem Inhalt übernommen", styles["CaptionK"])])
            elif kind == "callout":
                signal = block["text"].casefold().startswith(("warnung", "frist", "achtung"))
                fill = red if signal else blue
                cell = [[Paragraph(("WICHTIG" if signal else "KERNAUSSAGE"), styles["CalloutLabel"]), Paragraph(_rl_inline(block["text"]), styles["Callout"])]]
                table = Table(cell, colWidths=[30*mm, 144*mm], hAlign="LEFT")
                table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),fill),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),12),("RIGHTPADDING",(0,0),(-1,-1),12),("TOPPADDING",(0,0),(-1,-1),12),("BOTTOMPADDING",(0,0),(-1,-1),12)]))
                story.append(KeepTogether(table))
                story.append(Spacer(1, 8))
            elif kind == "code":
                for line_value in block["text"].splitlines() or [""]:
                    story.append(Paragraph(html.escape(line_value) or "&nbsp;", styles["CodeK"]))

        class BrandedCanvas(rl_canvas.Canvas):
            def __init__(self, *args: Any, **kwargs: Any) -> None:
                super().__init__(*args, **kwargs)
                self._kahle_states: list[dict[str, Any]] = []

            def showPage(self) -> None:
                self._kahle_states.append(dict(self.__dict__))
                self._startPage()

            def save(self) -> None:
                total = len(self._kahle_states)
                for state in self._kahle_states:
                    self.__dict__.update(state)
                    self._kahle_page_count = total
                    draw_page(self, doc)
                    rl_canvas.Canvas.showPage(self)
                rl_canvas.Canvas.save(self)

        doc.build(story, canvasmaker=BrandedCanvas)
        return out.getvalue()
    except Exception:
        return None
