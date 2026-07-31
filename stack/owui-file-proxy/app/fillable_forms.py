"""Deterministic, genuinely fillable Word forms for KAHLE governance."""
from __future__ import annotations
import io, uuid
from pathlib import Path
from typing import Any
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
try:
    from .kahle_document_theme import BLUE, INK, MUTED, WHITE, _add_docx_callout, _add_inline_docx, _company, _config, _configure_section, _set_cell_margins, _set_cell_shading, _set_run, _style_document
except ImportError:
    from kahle_document_theme import BLUE, INK, MUTED, WHITE, _add_docx_callout, _add_inline_docx, _company, _config, _configure_section, _set_cell_margins, _set_cell_shading, _set_run, _style_document

def _sdt(alias: str, tag: str):
    el, pr = OxmlElement("w:sdt"), OxmlElement("w:sdtPr")
    for name, value in (("w:alias", alias), ("w:tag", tag), ("w:id", str(uuid.uuid4().int % 2_000_000_000))):
        node = OxmlElement(name); node.set(qn("w:val"), value); pr.append(node)
    lock = OxmlElement("w:lock"); lock.set(qn("w:val"), "sdtLocked"); pr.append(lock)
    el.append(pr)
    return el, pr

def _control(paragraph: Any, field: dict[str, Any]):
    el, pr = _sdt(field["label"], field["tag"]); kind = field.get("type", "text")
    if kind == "date":
        control = OxmlElement("w:date")
        fmt = OxmlElement("w:dateFormat"); fmt.set(qn("w:val"), "dd.MM.yyyy"); control.append(fmt)
        lid = OxmlElement("w:lid"); lid.set(qn("w:val"), "de-DE"); control.append(lid)
    elif kind == "dropdown":
        control = OxmlElement("w:dropDownList")
        for option in field.get("options", []):
            item = OxmlElement("w:listItem"); item.set(qn("w:displayText"), option); item.set(qn("w:value"), option); control.append(item)
    else:
        control = OxmlElement("w:text")
        if kind == "multiline": control.set(qn("w:multiLine"), "1")
    pr.append(control)
    content, run, text = OxmlElement("w:sdtContent"), OxmlElement("w:r"), OxmlElement("w:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve"); text.text = field.get("placeholder", "Hier eingeben")
    run.append(text); content.append(run); el.append(content); paragraph._p.append(el)

def _checkbox(paragraph: Any, label: str, tag: str):
    el, pr = _sdt(label, tag); box = OxmlElement("w14:checkbox")
    for name, value in (("w14:checked", "0"), ("w14:checkedState", "2612"), ("w14:uncheckedState", "2610")):
        node = OxmlElement(name); node.set(qn("w14:val"), value); box.append(node)
    pr.append(box); content, run, text = OxmlElement("w:sdtContent"), OxmlElement("w:r"), OxmlElement("w:t")
    text.text = "☐"; run.append(text); content.append(run); el.append(content); paragraph._p.append(el)
    label_run = paragraph.add_run(f"  {label}"); _set_run(label_run, name="Arial", size=9.5, color=INK)

def _heading(document: Any, text: str):
    p = document.add_paragraph(style="Heading 2"); p.paragraph_format.keep_with_next = True; _add_inline_docx(p, text, color=INK, size=15)

def _fields(document: Any, fields: list[dict[str, Any]]):
    table = document.add_table(rows=0, cols=2); table.autofit = False
    table.columns[0].width, table.columns[1].width = Cm(5), Cm(11.2)
    for index, field in enumerate(fields):
        cells = table.add_row().cells
        for cell in cells: _set_cell_margins(cell, 110, 110, 90, 110); _set_cell_shading(cell, "F3F7FA" if index % 2 == 0 else WHITE)
        run = cells[0].paragraphs[0].add_run(field["label"]); _set_run(run, name="Arial", size=9, bold=True, color=INK)
        _control(cells[1].paragraphs[0], field)
    document.add_paragraph().paragraph_format.space_after = Pt(2)

def _checks(document: Any, items: list[tuple[str, str]]):
    table = document.add_table(rows=0, cols=1)
    for index, (label, tag) in enumerate(items):
        cell = table.add_row().cells[0]; _set_cell_margins(cell, 90, 90, 90, 110); _set_cell_shading(cell, "F3F7FA" if index % 2 == 0 else WHITE)
        _checkbox(cell.paragraphs[0], label, tag)

def _reset(document: Any):
    body = document._element.body; section = None
    for child in list(body):
        if child.tag.endswith("sectPr"): section = child
        else: body.remove(child)
    if section is not None and section.getparent() is None: body.append(section)

def render_ki_permission_form_docx(template_path: Path, logo_path: Path, config_path: Path, generated_at: str) -> bytes | None:
    """Render an interactive KI usage permission form with real Word controls."""
    try:
        document = Document(str(template_path)) if template_path.exists() else Document()
        _reset(document); _style_document(document); _configure_section(document, _company(_config(config_path)), logo_path)
        kicker = document.add_paragraph(style="Kahle Kicker"); _add_inline_docx(kicker, "KAHLE INTERN · KI-GOVERNANCE", color=BLUE, size=8.5)
        title = "Antrag auf Freigabe einer KI-Nutzung"
        title_p = document.add_paragraph(style="Title"); _add_inline_docx(title_p, title, color=INK, size=27)
        _add_docx_callout(document, "Dieses Formular dokumentiert die Prüfung und gemeinsame Entscheidung über eine genehmigungspflichtige KI-Nutzung. Es ersetzt keine Datenschutz-Folgenabschätzung, IT-Sicherheitsprüfung, Beteiligung des Betriebsrats oder Rechtsberatung.")

        _heading(document, "1. Vorgang und Antragsteller")
        _fields(document, [
            {"label":"Vorgangsnummer","tag":"case_id","placeholder":"wird intern vergeben"},
            {"label":"Antragsdatum","tag":"application_date","type":"date","placeholder":"Datum auswählen"},
            {"label":"Name","tag":"applicant_name"}, {"label":"Abteilung / Funktion","tag":"applicant_department"},
            {"label":"Standort","tag":"applicant_location"}, {"label":"E-Mail / Telefon","tag":"applicant_contact"},
            {"label":"Führungskraft","tag":"manager_name"},
        ])
        _heading(document, "2. Beantragte KI-Nutzung")
        _fields(document, [
            {"label":"KI-System / Produkt","tag":"ai_system_name"}, {"label":"Anbieter / Version / URL","tag":"ai_vendor_version"},
            {"label":"Einsatzzweck und erwarteter Nutzen","tag":"purpose","type":"multiline","placeholder":"Zweck, Prozess und gewünschtes Ergebnis beschreiben"},
            {"label":"Nutzerkreis","tag":"user_group"},
            {"label":"Einsatzhäufigkeit","tag":"frequency","type":"dropdown","placeholder":"Bitte auswählen","options":["Einmalig","Gelegentlich","Regelmäßig","Dauerhaft"]},
            {"label":"Geplanter Start","tag":"planned_start","type":"date","placeholder":"Datum auswählen"}, {"label":"Beantragte Laufzeit","tag":"requested_duration"},
        ])
        _heading(document, "3. Grund der Genehmigungspflicht")
        _checks(document, [
            ("KI-System steht nicht auf der KAHLE-Whitelist","reason_not_whitelisted"), ("Nutzung eines GPAI-/Consumer-Tools","reason_gpai"),
            ("Mögliche Hochrisiko-KI-Anwendung","reason_high_risk"), ("HR-Entscheidung, Leistungs- oder Verhaltenskontrolle","reason_hr"),
            ("Scoring oder Profiling von Kunden oder Beschäftigten","reason_scoring"), ("Verarbeitung vertraulicher oder personenbezogener Daten","reason_sensitive_data"),
            ("Sonstiger genehmigungspflichtiger Einsatz","reason_other"),
        ])
        _fields(document, [{"label":"Erläuterung","tag":"approval_reason_details","type":"multiline","placeholder":"Grund und Abgrenzung erläutern"}])
        _heading(document, "4. Daten, Betroffene und technische Nutzung")
        _fields(document, [
            {"label":"Verarbeitete Datenarten","tag":"data_categories","type":"multiline","placeholder":"Keine / interne / vertrauliche / personenbezogene / besondere Kategorien"},
            {"label":"Betroffene Personengruppen","tag":"affected_persons"}, {"label":"Eingaben in das KI-System","tag":"ai_inputs","type":"multiline"},
            {"label":"Erzeugte Ergebnisse und Verwendung","tag":"ai_outputs","type":"multiline"}, {"label":"Speicherort / Datenübertragung","tag":"storage_transfer"},
            {"label":"Menschliche Prüfung der Ergebnisse","tag":"human_oversight","type":"multiline","placeholder":"Verantwortliche Rolle und Kontrollschritt beschreiben"},
        ])
        _heading(document, "5. Fachliche Prüfungen")
        status = ["Nicht erforderlich","Erforderlich","Abgeschlossen","Offen"]
        _fields(document, [
            {"label":"Risikoeinstufung","tag":"risk_level","type":"dropdown","placeholder":"Bitte auswählen","options":["Niedrig","Begrenzt","Hoch / vertiefte Prüfung","Unklar"]},
            {"label":"Datenschutzprüfung","tag":"privacy_review","type":"dropdown","placeholder":"Bitte auswählen","options":status},
            {"label":"DSFA","tag":"dsfa_status","type":"dropdown","placeholder":"Bitte auswählen","options":status},
            {"label":"IT-Sicherheitsprüfung","tag":"security_review","type":"dropdown","placeholder":"Bitte auswählen","options":status},
            {"label":"Betriebsrat / Personalvertretung","tag":"works_council_review","type":"dropdown","placeholder":"Bitte auswählen","options":["Nicht betroffen","Beteiligung erforderlich","Beteiligt","Offen"]},
            {"label":"Prüfvermerke / Anlagen","tag":"review_notes","type":"multiline","placeholder":"Prüfungen, Dokumente, Links und offene Punkte"},
        ])
        _heading(document, "6. Erklärung des Antragstellers")
        _checks(document, [
            ("Ich verwende die KI ausschließlich für den beschriebenen Zweck.","declaration_purpose"),
            ("Ich lade nur zulässige Daten und Informationen hoch.","declaration_data"),
            ("Ich prüfe KI-Ergebnisse fachlich und übernehme die Verantwortung für ihre Nutzung.","declaration_review"),
            ("Ich beachte Auflagen, Befristungen, Dokumentations- und Meldepflichten.","declaration_conditions"),
            ("Ich melde Vorfälle und wesentliche Änderungen unverzüglich.","declaration_incidents"),
        ])
        _fields(document, [
            {"label":"Name Antragsteller","tag":"applicant_signature_name"}, {"label":"Datum","tag":"applicant_signature_date","type":"date","placeholder":"Datum auswählen"},
            {"label":"Unterschrift / digitale Bestätigung","tag":"applicant_signature","placeholder":"Unterschrift oder Bestätigungsvermerk"},
        ])
        _heading(document, "7. Gemeinsame Entscheidung")
        _checks(document, [("Freigegeben","decision_approved"),("Freigegeben mit Auflagen","decision_conditional"),("Zurückgestellt – weitere Prüfung erforderlich","decision_deferred"),("Abgelehnt","decision_rejected")])
        _fields(document, [
            {"label":"Gültig ab","tag":"valid_from","type":"date","placeholder":"Datum auswählen"}, {"label":"Gültig bis","tag":"valid_until","type":"date","placeholder":"Datum auswählen"},
            {"label":"Auflagen und Nutzungsgrenzen","tag":"conditions","type":"multiline","placeholder":"Erlaubte Daten, Nutzer, Zwecke, Systeme und Kontrollen"},
            {"label":"Überprüfung / Wiedervorlage","tag":"review_date","type":"date","placeholder":"Datum auswählen"}, {"label":"Begründung bei Ablehnung","tag":"rejection_reason","type":"multiline"},
        ])
        _heading(document, "8. Fachliche Mitzeichnungen (nur wenn erforderlich)")
        _fields(document, [
            {"label":"Fachbereich / Führungskraft","tag":"business_owner_name"}, {"label":"Datum / Bestätigung","tag":"business_owner_approval","placeholder":"Datum und Unterschrift oder digitaler Bestätigungsvermerk"},
            {"label":"Datenschutzbeauftragte / DSK","tag":"privacy_approver_name"}, {"label":"Datum / Bestätigung","tag":"privacy_approval","placeholder":"Datum und Unterschrift oder digitaler Bestätigungsvermerk"},
            {"label":"IT-Sicherheit","tag":"security_approver_name"}, {"label":"Datum / Bestätigung","tag":"security_approval","placeholder":"Datum und Unterschrift oder digitaler Bestätigungsvermerk"},
            {"label":"Betriebsrat / Personalvertretung","tag":"works_council_approver_name"}, {"label":"Datum / Bestätigung","tag":"works_council_approval","placeholder":"Datum und Unterschrift oder digitaler Bestätigungsvermerk"},
        ])
        _heading(document, "9. Verbindliche Endfreigabe")
        _add_docx_callout(document, "Die Nutzung ist erst freigegeben, wenn KI-Beauftragter und Geschäftsführung gemeinsam bestätigt haben. Bedingte Mitzeichnungen aus Abschnitt 8 sind vorher einzuholen, sofern die Prüfung sie als erforderlich ausweist.")
        _fields(document, [
            {"label":"KI-Beauftragter – Name","tag":"ai_officer_name"}, {"label":"KI-Beauftragter – Datum","tag":"ai_officer_date","type":"date","placeholder":"Datum auswählen"},
            {"label":"KI-Beauftragter – Unterschrift","tag":"ai_officer_signature","placeholder":"Unterschrift oder digitaler Bestätigungsvermerk"},
            {"label":"Geschäftsführung – Name","tag":"management_name"}, {"label":"Geschäftsführung – Datum","tag":"management_date","type":"date","placeholder":"Datum auswählen"},
            {"label":"Geschäftsführung – Unterschrift","tag":"management_signature","placeholder":"Unterschrift oder digitaler Bestätigungsvermerk"},
        ])
        note=document.add_paragraph(); note.paragraph_format.space_before=Pt(8)
        _add_inline_docx(note, "Schulungsnachweise und KI-Vorfallmeldungen werden separat geführt. Wesentliche Änderungen an Zweck, System, Anbieter, Datenarten oder Nutzerkreis erfordern eine erneute Prüfung.", color=MUTED, size=8.5)
        document.core_properties.title=title; document.core_properties.subject="Ausfüllbare Vorlage zur dokumentierten Freigabe einer KI-Nutzung"; document.core_properties.author="KAHLE-Vinci"
        out=io.BytesIO(); document.save(out); return out.getvalue()
    except Exception:
        return None

KI_POLICY_QUIZ_MC = [
    ("Welche KI-Systeme dürfen ohne Sonderfreigabe genutzt werden?", ["Nur Systeme der aktuellen KAHLE-Whitelist", "Jedes öffentlich erreichbare KI-System", "Alle Systeme mit kostenlosem Konto", "Nur privat beschaffte Systeme"]),
    ("Wer kann ein KI-Tool außerhalb der Whitelist freigeben?", ["KI-Beauftragter gemeinsam mit der Geschäftsführung", "Jede Führungskraft allein", "IT-Sicherheit allein", "Jeder Mitarbeitende für den eigenen Bereich"]),
    ("Was gilt bei kritischen Entscheidungen?", ["Ein Mensch muss aktiv prüfen und entscheiden (Human-in-the-Loop)", "Die KI-Entscheidung ist automatisch verbindlich", "Nur das Ergebnis muss archiviert werden", "Eine Stichprobe pro Jahr reicht aus"]),
    ("Welche Nutzung ist ohne ausdrückliche schriftliche Erlaubnis verboten?", ["HR-Entscheidungen, Leistungs-/Verhaltenskontrolle oder Kundenprofiling", "Entwurf einer internen Agenda", "Rechtschreibprüfung eines allgemeinen Textes", "Strukturierung nicht vertraulicher Notizen"]),
    ("Was gilt für KI-generierte Inhalte bei externer Nutzung?", ["Sie müssen entsprechend der Transparenzpflicht gekennzeichnet werden", "Sie dürfen nie verwendet werden", "Eine Kennzeichnung ist nur bei Bildern nötig", "Die Kennzeichnung ist freiwillig"]),
    ("Welche Aussage zu Daten ist richtig?", ["Vertrauliche oder personenbezogene Daten dürfen nicht in nicht freigegebene öffentliche KI-Systeme geladen werden", "Personenbezogene Daten sind immer erlaubt, wenn Namen entfernt wurden", "Interne Daten dürfen in jedes Consumer-Tool geladen werden", "Datenregeln gelten nur für Kundendaten"]),
    ("Was ist bei einem KI-Vorfall zu tun?", ["Unverzüglich über einen vorgesehenen Meldeweg melden", "Bis zur nächsten Schulung warten", "Nur die lokale Datei löschen", "Den Vorfall ausschließlich mit Kollegen besprechen"]),
    ("Wie häufig wird die Richtlinie überprüft?", ["Jährlich sowie bei regulatorischen Änderungen", "Nur nach einem schweren Vorfall", "Alle fünf Jahre", "Eine Überprüfung ist nicht vorgesehen"]),
    ("Welche Pflicht haben Mitarbeitende?", ["Richtlinie einhalten, Ergebnisse prüfen und Vorfälle melden", "Nur Prompts dokumentieren", "KI-Ergebnisse ungeprüft übernehmen", "Ausschließlich die IT über KI-Nutzung informieren"]),
]


def _quiz_question(document: Any, number: int, question: str, options: list[str] | None = None):
    p = document.add_paragraph(); p.paragraph_format.space_before = Pt(7); p.paragraph_format.keep_with_next = True
    _add_inline_docx(p, f"{number}. {question}", color=INK, size=10.5)
    if options:
        _fields(document, [{"label":"Antwort","tag":f"quiz_q{number}","type":"dropdown","placeholder":"Bitte auswählen","options":["Bitte auswählen", *options]}])
    else:
        _fields(document, [{"label":"Antwort / Begründung","tag":f"quiz_q{number}","type":"multiline","placeholder":"Antwort kurz und nachvollziehbar begründen"}])


def render_ki_policy_quiz_docx(template_path: Path, logo_path: Path, config_path: Path, generated_at: str) -> bytes | None:
    """Render a high-quality interactive knowledge check for KAHLE KI policy v1.4."""
    try:
        document = Document(str(template_path)) if template_path.exists() else Document()
        _reset(document); _style_document(document); _configure_section(document, _company(_config(config_path)), logo_path)
        kicker=document.add_paragraph(style="Kahle Kicker"); _add_inline_docx(kicker,"KAHLE INTERN · SCHULUNG & COMPLIANCE",color=BLUE,size=8.5)
        title="Wissenstest zur KAHLE-KI-Richtlinie"
        title_p=document.add_paragraph(style="Title"); _add_inline_docx(title_p,title,color=INK,size=27)
        _add_docx_callout(document,"Dieser Fragebogen prüft zentrale Regeln der KI-Richtlinie v1.4. Bitte alle Fragen selbstständig beantworten. Die fachliche Auswertung erfolgt durch die Schulungsleitung.")
        _heading(document,"1. Teilnehmer und Durchführung")
        _fields(document,[
            {"label":"Name","tag":"quiz_participant_name"},{"label":"Abteilung / Funktion","tag":"quiz_department"},
            {"label":"Standort","tag":"quiz_location"},{"label":"Datum","tag":"quiz_date","type":"date","placeholder":"Datum auswählen"},
            {"label":"Schulung / Anlass","tag":"quiz_training"},{"label":"Bearbeitungszeit","tag":"quiz_duration","placeholder":"Empfehlung: 15–20 Minuten"},
        ])
        _heading(document,"2. Multiple-Choice-Fragen")
        for index,(question,options) in enumerate(KI_POLICY_QUIZ_MC,1): _quiz_question(document,index,question,options)
        _heading(document,"3. Praxis- und Transferfragen")
        scenarios=[
            "Ein Kollege möchte Kundendaten in ein öffentliches, nicht freigegebenes KI-Tool laden. Wie reagieren Sie und warum?",
            "Ein KI-System liefert eine Empfehlung, die eine Personalentscheidung beeinflussen könnte. Welche Schritte sind erforderlich?",
            "Nennen Sie einen meldepflichtigen KI-Vorfall und beschreiben Sie den vorgesehenen Meldeweg.",
        ]
        for offset,question in enumerate(scenarios,len(KI_POLICY_QUIZ_MC)+1): _quiz_question(document,offset,question)
        _heading(document,"4. Bestätigung des Teilnehmenden")
        _checks(document,[("Ich habe den Fragebogen selbstständig und nach bestem Wissen beantwortet.","quiz_declaration"),("Mir ist bekannt, dass KI-Ergebnisse fachlich geprüft und Vorfälle gemeldet werden müssen.","quiz_acknowledgement")])
        _fields(document,[{"label":"Name","tag":"quiz_signature_name"},{"label":"Datum","tag":"quiz_signature_date","type":"date","placeholder":"Datum auswählen"},{"label":"Unterschrift / digitale Bestätigung","tag":"quiz_signature"}])
        _heading(document,"5. Auswertung durch die Schulungsleitung")
        _fields(document,[
            {"label":"Erreichte Punkte","tag":"quiz_score"},{"label":"Maximalpunkte","tag":"quiz_max_score","placeholder":"15"},
            {"label":"Ergebnis","tag":"quiz_result","type":"dropdown","placeholder":"Bitte auswählen","options":["Bitte auswählen","Bestanden","Nachschulung erforderlich","Noch nicht bewertet"]},
            {"label":"Feedback / Lernfelder","tag":"quiz_feedback","type":"multiline","placeholder":"Stärken, Fehler und konkrete Lernhinweise"},
            {"label":"Bewertet von","tag":"quiz_reviewer"},{"label":"Bewertungsdatum","tag":"quiz_review_date","type":"date","placeholder":"Datum auswählen"},
        ])
        document.core_properties.title=title; document.core_properties.subject="Interaktiver Wissenstest zur KAHLE-KI-Richtlinie v1.4"; document.core_properties.author="KAHLE-Vinci"
        out=io.BytesIO(); document.save(out); return out.getvalue()
    except Exception:
        return None

def _safe_form_tag(value: str, fallback: str) -> str:
    import re
    tag = re.sub(r"[^a-zA-Z0-9_]+", "_", str(value or "").strip()).strip("_").lower()
    return (tag or fallback)[:48]


def render_dynamic_form_docx(template_path: Path, logo_path: Path, config_path: Path, generated_at: str, schema: dict[str, Any]) -> bytes | None:
    """Render a validated, scenario-independent interactive Word form schema."""
    try:
        document = Document(str(template_path)) if template_path.exists() else Document()
        _reset(document); _style_document(document); _configure_section(document, _company(_config(config_path)), logo_path)
        title = str(schema.get("title") or "Interaktives KAHLE-Formular")[:120]
        kicker = document.add_paragraph(style="Kahle Kicker")
        _add_inline_docx(kicker, str(schema.get("kicker") or "KAHLE INTERN · INTERAKTIVES FORMULAR")[:100], color=BLUE, size=8.5)
        title_p = document.add_paragraph(style="Title"); _add_inline_docx(title_p, title, color=INK, size=27)
        instructions = str(schema.get("instructions") or "Bitte füllen Sie alle zutreffenden Felder vollständig aus.")[:700]
        _add_docx_callout(document, instructions)

        identity = schema.get("identity_fields") or []
        if identity:
            _heading(document, "Allgemeine Angaben")
            normalized=[]
            for index,item in enumerate(identity[:10],1):
                normalized.append({
                    "label":str(item.get("label") or f"Angabe {index}")[:100],
                    "tag":_safe_form_tag(item.get("id"),f"identity_{index}"),
                    "type":str(item.get("type") or "text"),
                    "placeholder":str(item.get("placeholder") or "Hier eingeben")[:140],
                    "options":[str(v)[:100] for v in (item.get("options") or [])[:12]],
                })
            _fields(document,normalized)

        field_no=0
        for section_no,section in enumerate((schema.get("sections") or [])[:12],1):
            _heading(document, f"{section_no}. {str(section.get('title') or 'Abschnitt')[:110]}")
            description=str(section.get("description") or "").strip()
            if description:
                p=document.add_paragraph(); _add_inline_docx(p,description[:700],color=MUTED,size=9)
            for item in (section.get("items") or [])[:16]:
                field_no+=1; kind=str(item.get("type") or "text").lower()
                label=str(item.get("label") or f"Feld {field_no}")[:300]
                tag=_safe_form_tag(item.get("id"),f"field_{field_no}")
                if kind == "checkbox":
                    _checks(document,[(label,tag)])
                else:
                    field={"label":label,"tag":tag,"type":kind if kind in {"text","multiline","dropdown","date"} else "text","placeholder":str(item.get("placeholder") or ("Bitte auswählen" if kind=="dropdown" else "Hier eingeben"))[:140]}
                    if kind=="dropdown": field["options"]=[str(v)[:120] for v in (item.get("options") or [])[:12]]
                    _fields(document,[field])

        declarations=[str(v)[:300] for v in (schema.get("declarations") or [])[:8] if str(v).strip()]
        if declarations:
            _heading(document,"Bestätigung")
            _checks(document,[(label,f"declaration_{i}") for i,label in enumerate(declarations,1)])
        signature=schema.get("signature_fields") or []
        if signature:
            normalized=[]
            for index,item in enumerate(signature[:8],1):
                normalized.append({"label":str(item.get("label") or f"Bestätigung {index}")[:100],"tag":_safe_form_tag(item.get("id"),f"signature_{index}"),"type":str(item.get("type") or "text"),"placeholder":str(item.get("placeholder") or "Hier eingeben")[:140]})
            _fields(document,normalized)
        document.core_properties.title=title; document.core_properties.subject="Interaktives, RAG-basiertes KAHLE-Formular"; document.core_properties.author="KAHLE-Vinci"
        out=io.BytesIO(); document.save(out); return out.getvalue()
    except Exception:
        return None