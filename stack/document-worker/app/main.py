from __future__ import annotations

import io
import json
import logging
import os
import re
import tempfile
import time
from typing import Any, Dict, List, Optional, Tuple

from fastapi import FastAPI, UploadFile, File, Form, Header, HTTPException, Request
from fastapi.exceptions import RequestValidationError
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel, Field
from pythonjsonlogger import jsonlogger

from .security import require_api_key, sha256_bytes
from .storage import save_temp_file, cleanup_expired, STORAGE_DIR, ensure_storage
from .docx_ops import docx_simple_replace
from .xlsx_ops import xlsx_update_cells
from .pdf_ops import pdf_remove_pages


# ----------------------------
# App & Logging
# ----------------------------

APP_VERSION = "1.1.0"

app = FastAPI(title="KAHLE-Vinci Document Worker", version=APP_VERSION)

log = logging.getLogger("doc-worker")
log.setLevel(logging.INFO)
log.propagate = False

if not log.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(jsonlogger.JsonFormatter("%(asctime)s %(levelname)s %(message)s"))
    log.addHandler(handler)


# ----------------------------
# Config
# ----------------------------

MAX_FILE_MB = int(os.getenv("MAX_FILE_MB", "30"))
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024
USE_MARKITDOWN = os.getenv("USE_MARKITDOWN", "true").lower() == "true"


# ----------------------------
# Models
# ----------------------------

class CleanupOldFilesRequest(BaseModel):
    days: int = Field(15, ge=1, le=3650, description="Delete files older than this many days")
    dry_run: bool = Field(False, description="If true, list files only, do not delete")
    max_files: int = Field(100000, ge=1, le=1000000, description="Safety limit for processed files")


# ----------------------------
# Helpers
# ----------------------------

def _auth(x_api_key: Optional[str]) -> None:
    """
    require_api_key(...) should be implemented so that:
    - if no API key is configured, it allows access (dev/default),
    - if configured, it enforces it.
    """
    try:
        require_api_key(x_api_key)
    except PermissionError:
        raise HTTPException(status_code=401, detail="Unauthorized")


async def _read_upload_limited(file: UploadFile) -> bytes:
    b = await file.read()
    if len(b) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Max {MAX_FILE_MB} MB."
        )
    return b


def _loads_json_str(s: str) -> Any:
    # robust against BOM (common on Windows)
    return json.loads((s or "").lstrip("\ufeff"))


def _decode_text(b: bytes) -> str:
    # Try UTF-8, fallback latin-1
    try:
        return b.decode("utf-8")
    except UnicodeDecodeError:
        return b.decode("latin-1", errors="replace")


_MOJIBAKE_MARKERS = ("Ã", "Â", "â€", "â€“", "â€”", "â€¢", "ðŸ")


def _mojibake_score(text: str) -> int:
    return sum((text or "").count(marker) for marker in _MOJIBAKE_MARKERS)


def _repair_mojibake(text: str) -> str:
    """
    Fix common UTF-8 text that was decoded as Windows-1252, e.g. "fÃ¼r" -> "für".
    PDF extractors can emit this when source PDFs carry malformed font encodings.
    """
    if not text or _mojibake_score(text) == 0:
        return text or ""

    best = text
    best_score = _mojibake_score(text)
    for encoding in ("cp1252", "latin-1"):
        try:
            candidate = text.encode(encoding).decode("utf-8")
        except Exception:
            continue
        score = _mojibake_score(candidate)
        if score < best_score:
            best = candidate
            best_score = score

    return best


def _fix_letter_spaced_words(text: str) -> str:
    letters = r"A-Za-zÄÖÜäöüß"

    def join_letter_spaced(match: re.Match[str]) -> str:
        return re.sub(r"\s+", "", match.group(0))

    # "A r b e i t s a n w e i s u n g" -> "Arbeitsanweisung"
    text = re.sub(rf"(?<![\w])(?:[{letters}][ \t]+){{3,}}[{letters}](?![\w])", join_letter_spaced, text)

    # "E -Mails", "CRM -Manager" -> "E-Mails", "CRM-Manager"
    text = re.sub(rf"([{letters}0-9])[ \t]+-[ \t]*([{letters}0-9])", r"\1-\2", text)

    # "e uch", "f ür", "K AHLE" -> "euch", "für", "KAHLE".
    # The left side must be an isolated single-letter fragment; never join
    # normal word pairs such as "für euch".
    text = re.sub(r"(?<!\S)([^\W\d_])[ \t]+([^\W\d_]{2,})\b", r"\1\2", text)

    def join_short_left(match: re.Match[str]) -> str:
        left, right = match.group(1), match.group(2)
        if right[:1].islower() and len(left) <= 5:
            return left + right
        return f"{left} {right}"

    # "Auf  gaben" -> "Aufgaben", but keep "After  Sales".
    text = re.sub(rf"\b([{letters}]{{2,5}})[ \t]{{2,}}([{letters}]{{2,}})\b", join_short_left, text)

    def join_short_suffix(match: re.Match[str]) -> str:
        left, right = match.group(1), match.group(2)
        if right in {"t", "gt", "d", "st", "et", "en", "er"}:
            return left + right
        return f"{left} {right}"

    # "Le gt", "Denk t", "wend et" -> "Legt", "Denkt", "wendet".
    text = re.sub(rf"\b([{letters}]{{2,12}})[ \t]+([a-zäöüß]{{1,2}})\b", join_short_suffix, text)
    return text


def _normalize_extracted_text(text: str, *, paragraphize: bool = False) -> str:
    text = _repair_mojibake(text or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _fix_letter_spaced_words(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" +([,.;:!?])", r"\1", text)
    text = re.sub(r"(?m)^(\d+)\.(?=\S)", r"\1. ", text)
    text = re.sub(r"(?m)^[•]\s*", "- ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not paragraphize:
        return text
    return _paragraphize_extracted_text(text)


def _is_markdown_structural_line(line: str) -> bool:
    if not line:
        return True
    if line.startswith("#"):
        return True
    if _is_markdown_list_line(line):
        return True
    if re.match(r"^\|.*\|$", line):
        return True
    if len(line) <= 90 and line.endswith(":"):
        return True
    return False


def _is_markdown_list_line(line: str) -> bool:
    return bool(re.match(r"^([-*+]|\d+[.)])\s+", line or ""))


def _line_ends_sentence(line: str) -> bool:
    return bool(re.search(r"[.!?;:)\"”]$", (line or "").strip()))


def _paragraphize_extracted_text(text: str) -> str:
    """
    PDF text extraction often returns visual line wraps instead of paragraphs.
    This keeps lists/headings intact and folds consecutive prose lines into paragraphs.
    """
    out: List[str] = []
    buffer: List[str] = []
    list_item: Optional[str] = None
    list_item_had_blank = False

    def flush() -> None:
        if not buffer:
            return
        out.append(" ".join(buffer).strip())
        buffer.clear()

    def flush_list_item() -> None:
        nonlocal list_item, list_item_had_blank
        if list_item:
            out.append(re.sub(r"\s+", " ", list_item).strip())
            list_item = None
        list_item_had_blank = False

    for raw in (text or "").splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            if list_item:
                list_item_had_blank = True
                continue
            if buffer and not _line_ends_sentence(buffer[-1]):
                continue
            flush()
            if out and out[-1] != "":
                out.append("")
            continue

        if _is_markdown_list_line(line):
            flush()
            flush_list_item()
            list_item = line
            list_item_had_blank = False
            continue

        if list_item:
            if _is_markdown_structural_line(line):
                flush_list_item()
                out.append(line)
            elif list_item_had_blank and _line_ends_sentence(list_item):
                flush_list_item()
                buffer.append(line)
            else:
                list_item = f"{list_item} {line}"
                list_item_had_blank = False
            continue

        if buffer and not _line_ends_sentence(buffer[-1]) and not (
            line.startswith("#") or _is_markdown_list_line(line) or re.match(r"^\|.*\|$", line)
        ):
            buffer.append(line)
            continue

        if _is_markdown_structural_line(line):
            flush()
            out.append(line)
            continue

        buffer.append(line)

    flush_list_item()
    flush()
    cleaned = "\n".join(out)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _extract_text_markitdown(filename: str, content: bytes) -> Optional[str]:
    if not USE_MARKITDOWN:
        return None

    ext = _guess_ext(filename)
    if ext not in ("pdf", "docx", "xlsx", "xlsm"):
        return None

    try:
        from markitdown import MarkItDown  # type: ignore
    except Exception:
        return None

    suffix = f".{ext}" if ext else ""
    tmp_name = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_name = tmp.name
        result = MarkItDown().convert(tmp_name)
        converted = getattr(result, "text_content", "") or ""
        converted = converted.strip()
        return converted or None
    except Exception:
        return None
    finally:
        if tmp_name:
            try:
                os.unlink(tmp_name)
            except Exception:
                pass


def _extract_text_docx(docx_bytes: bytes) -> str:
    # Extract plain text from DOCX using python-docx
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx not available: {e}")

    doc = Document(io.BytesIO(docx_bytes))
    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    # tables
    for t in doc.tables:
        for row in t.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = "\n".join([p.text for p in cell.paragraphs if p.text]).strip()
                row_cells.append(cell_text)
            if any(row_cells):
                parts.append(" | ".join(row_cells))

    return _normalize_extracted_text("\n".join(parts).strip(), paragraphize=False)


def _extract_text_pdf(pdf_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"pypdf not available: {e}")

    reader = PdfReader(io.BytesIO(pdf_bytes))
    out: List[str] = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            page_text = _normalize_extracted_text(txt, paragraphize=True)
            out.append(f"## Seite {i+1}\n{page_text}")
    return "\n\n".join(out).strip()


def _extract_text_xlsx(xlsx_bytes: bytes, max_rows: int = 200, max_cols: int = 50) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"openpyxl not available: {e}")

    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True, read_only=True)
    out: List[str] = []

    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_count += 1
            if row_count > max_rows:
                out.append(f"... (abgeschnitten nach {max_rows} Zeilen)")
                break
            vals = []
            for c, v in enumerate(row, start=1):
                if c > max_cols:
                    vals.append("...(cols truncated)")
                    break
                if v is None:
                    vals.append("")
                else:
                    vals.append(str(v))
            out.append("\t".join(vals))
        out.append("")  # blank line between sheets

    return _normalize_extracted_text("\n".join(out).strip(), paragraphize=False)


def _guess_ext(filename: str) -> str:
    fn = (filename or "").lower().strip()
    if "." in fn:
        return fn.rsplit(".", 1)[-1]
    return ""


def _extract_text_by_ext(filename: str, content: bytes) -> Tuple[str, str]:
    ext = _guess_ext(filename)

    markitdown_text = _extract_text_markitdown(filename, content)
    if markitdown_text is not None:
        return ext or "unknown", _normalize_extracted_text(markitdown_text, paragraphize=(ext == "pdf"))

    if ext in ("txt", "md", "csv"):
        return ext, _normalize_extracted_text(_decode_text(content).strip(), paragraphize=False)

    if ext == "docx":
        return ext, _extract_text_docx(content)

    if ext == "pdf":
        return ext, _extract_text_pdf(content)

    if ext in ("xlsx", "xlsm"):
        return ext, _extract_text_xlsx(content)

    # unknown: attempt as text
    return ext or "unknown", _decode_text(content).strip()


def _render_text_to_pdf_bytes(content: str) -> bytes:
    try:
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.utils import simpleSplit  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"reportlab not available: {e}")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    # Prefer a Unicode-capable font if available to avoid missing-glyph squares.
    font_name = "Helvetica"
    for candidate in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    ):
        if os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont("DejaVuSans", candidate))
                font_name = "DejaVuSans"
                break
            except Exception:
                pass

    font_size = 11
    leading = 14
    left_margin = 40
    right_margin = 40
    top_margin = 50
    bottom_margin = 60
    usable_width = max(100, width - left_margin - right_margin)

    textobject = c.beginText(left_margin, height - top_margin)
    textobject.setFont(font_name, font_size)
    textobject.setLeading(leading)

    raw_lines = (content or "").splitlines() or [""]
    for line in raw_lines:
        wrapped_lines = simpleSplit(line, font_name, font_size, usable_width) or [""]
        for wrapped in wrapped_lines:
            textobject.textLine(wrapped)

            if textobject.getY() < bottom_margin:
                c.drawText(textobject)
                c.showPage()
                textobject = c.beginText(left_margin, height - top_margin)
                textobject.setFont(font_name, font_size)
                textobject.setLeading(leading)

    c.drawText(textobject)
    c.save()
    return buf.getvalue()


def _docx_delete_last_paragraphs(docx_bytes: bytes, n: int, allow_empty_output: bool = False) -> bytes:
    """
    Delete the last n non-empty paragraphs from a DOCX.
    Uses python-docx. Keeps tables etc. intact.
    """
    if n <= 0:
        return docx_bytes

    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx not available: {e}")

    doc = Document(io.BytesIO(docx_bytes))

    # Collect body paragraphs (python-docx includes empty paragraphs as well)
    paras = list(doc.paragraphs)
    non_empty = [p for p in paras if (p.text or "").strip() != ""]

    if not non_empty:
        raise HTTPException(
            status_code=400,
            detail={
                "error": "docx_has_no_non_empty_paragraphs",
                "non_empty_paragraphs": 0,
                "requested_n": n,
            },
        )

    # Guardrail: avoid silently producing a blank document unless explicitly allowed.
    if (not allow_empty_output) and n >= len(non_empty):
        raise HTTPException(
            status_code=400,
            detail={
                "error": "operation_would_remove_all_paragraphs_set_allow_empty_output_true_to_override",
                "non_empty_paragraphs": len(non_empty),
                "requested_n": n,
                "note": "count_uses_word_paragraphs_not_visual_line_breaks",
            },
        )

    removed = 0
    # Remove from end; prefer removing non-empty paragraphs first
    for p in reversed(paras):
        txt = (p.text or "").strip()
        if txt == "":
            # Optional: you can also remove empty paragraphs if you want.
            continue
        # Remove paragraph XML element
        try:
            p._element.getparent().remove(p._element)  # type: ignore[attr-defined]
            removed += 1
        except Exception:
            # If deletion fails, ignore and continue
            continue
        if removed >= n:
            break

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue()


def _apply_text_ops(text: str, ops: List[Dict[str, Any]]) -> str:
    """
    Apply deterministic text operations sequentially.
    Supported ops:
      - delete_last_lines: {op, n}
      - replace_all: {op, from, to}
      - append: {op, text}
      - prepend: {op, text}
    """
    out = text or ""

    for i, op in enumerate(ops):
        if not isinstance(op, dict):
            raise HTTPException(status_code=400, detail=f"ops[{i}] must be an object")

        kind = op.get("op")
        if kind == "delete_last_lines":
            n = int(op.get("n", 0))
            if n < 0:
                raise HTTPException(status_code=400, detail="delete_last_lines.n must be >= 0")
            lines = out.splitlines()
            if n >= len(lines):
                out = ""
            else:
                out = "\n".join(lines[:-n])

        elif kind == "replace_all":
            frm = op.get("from")
            to = op.get("to")
            if not isinstance(frm, str) or not isinstance(to, str):
                raise HTTPException(status_code=400, detail="replace_all requires string fields 'from' and 'to'")
            out = out.replace(frm, to)

        elif kind == "append":
            t = op.get("text", "")
            if not isinstance(t, str):
                raise HTTPException(status_code=400, detail="append.text must be a string")
            if out and not out.endswith("\n"):
                out += "\n"
            out += t

        elif kind == "prepend":
            t = op.get("text", "")
            if not isinstance(t, str):
                raise HTTPException(status_code=400, detail="prepend.text must be a string")
            if t and not t.endswith("\n"):
                t += "\n"
            out = t + out

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported op: {kind}")

    return out

def _validate_replacements(replacements: Any) -> List[Dict[str, str]]:
    if not isinstance(replacements, list):
        raise HTTPException(status_code=400, detail="replacements_json must be a JSON list.")
    out: List[Dict[str, str]] = []
    for i, r in enumerate(replacements):
        if not isinstance(r, dict):
            raise HTTPException(status_code=400, detail=f"replacements[{i}] must be an object.")
        frm = r.get("from")
        to = r.get("to")
        if not isinstance(frm, str) or not isinstance(to, str):
            raise HTTPException(status_code=400, detail=f"replacements[{i}] must contain string fields 'from' and 'to'.")
        out.append({"from": frm, "to": to})
    return out


# ----------------------------
# Error handling (better diagnostics)
# ----------------------------

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    # No raw content logging (PII), only path/method and field errors
    log.info("request_validation_error", extra={
        "path": request.url.path,
        "method": request.method,
        "errors": exc.errors(),
    })
    return JSONResponse(status_code=422, content={"detail": exc.errors()})


# ----------------------------
# Health
# ----------------------------

@app.get("/health")
def health():
    deleted = cleanup_expired()
    return {"ok": True, "cleanup_deleted": deleted, "version": APP_VERSION}


@app.post("/maintenance/cleanup_old_files", include_in_schema=False)
def maintenance_cleanup_old_files(
    payload: CleanupOldFilesRequest,
    x_api_key: Optional[str] = Header(default=None),
):
    _auth(x_api_key)
    ensure_storage()

    cutoff_ts = time.time() - (payload.days * 86400)
    checked_files = 0
    matched_files = 0
    deleted_files = 0
    deleted_bytes = 0
    samples: list[str] = []
    truncated = False

    processed = 0
    for p in STORAGE_DIR.rglob("*"):
        if not p.is_file():
            continue
        processed += 1
        if processed > payload.max_files:
            truncated = True
            break

        checked_files += 1
        try:
            st = p.stat()
        except Exception:
            continue
        if float(st.st_mtime) >= cutoff_ts:
            continue

        matched_files += 1
        if len(samples) < 200:
            samples.append(p.name)
        if payload.dry_run:
            continue

        try:
            size = int(st.st_size)
            p.unlink(missing_ok=True)
            deleted_files += 1
            deleted_bytes += size
        except Exception:
            continue

    return {
        "ok": True,
        "dry_run": payload.dry_run,
        "days": payload.days,
        "storage_dir": str(STORAGE_DIR),
        "cutoff_unix": int(cutoff_ts),
        "checked_files": checked_files,
        "matched_files": matched_files,
        "deleted_files": deleted_files,
        "deleted_bytes": deleted_bytes,
        "samples": samples,
        "truncated": truncated,
    }


# ----------------------------
# DOCX
# ----------------------------

@app.post("/docx/replace_one")
async def docx_replace_one(
    file: UploadFile = File(...),
    from_text: str = Form(...),
    to_text: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        replacements = [{"from": from_text, "to": to_text}]
        out = docx_simple_replace(b, replacements)
        save_temp_file("output.docx", out)

        log.info("docx_replace_one_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,   # DO NOT use "filename" (reserved)
        })

        return Response(
            content=out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=edited.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("docx_replace_one_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/docx/replace")
async def docx_replace(
    file: UploadFile = File(...),
    replacements_json: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        replacements_raw = _loads_json_str(replacements_json)
        replacements = _validate_replacements(replacements_raw)

        out = docx_simple_replace(b, replacements)
        save_temp_file("output.docx", out)

        log.info("docx_replace_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
            "repl_count": len(replacements),
        })

        return Response(
            content=out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=edited.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("docx_replace_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/docx/delete_last_paragraphs")
async def docx_delete_last_paragraphs(
    file: UploadFile = File(...),
    n: int = Form(3),
    allow_empty_output: bool = Form(False),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    if n < 1 or n > 500:
        raise HTTPException(status_code=400, detail="n must be between 1 and 500")

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        out = _docx_delete_last_paragraphs(b, n=n, allow_empty_output=allow_empty_output)
        save_temp_file("output.docx", out)

        log.info("docx_delete_last_paragraphs_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
            "n": n,
            "allow_empty_output": allow_empty_output,
        })

        return Response(
            content=out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=edited.docx"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("docx_delete_last_paragraphs_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/docx/to_pdf")
async def docx_to_pdf(
    file: UploadFile = File(...),
    filename: str = Form("converted.pdf"),
    x_api_key: str | None = Header(default=None),
):
    """
    Convert DOCX to a simple PDF by extracting text and rendering it.
    Note: preserves textual content, not full DOCX layout fidelity.
    """
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        text = _extract_text_docx(b)
        out = _render_text_to_pdf_bytes(text)

        out_name = (filename or "converted.pdf").strip() or "converted.pdf"
        if not out_name.lower().endswith(".pdf"):
            out_name = f"{out_name}.pdf"
        save_temp_file(out_name, out)

        log.info("docx_to_pdf_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
            "out_name": out_name,
        })

        return Response(
            content=out,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={out_name}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("docx_to_pdf_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))


# ----------------------------
# XLSX
# ----------------------------

@app.post("/xlsx/update_cells")
async def xlsx_update(
    file: UploadFile = File(...),
    updates_json: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        updates = _loads_json_str(updates_json)
        out = xlsx_update_cells(b, updates)
    except HTTPException:
        raise
    except Exception as e:
        log.info("xlsx_update_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))

    sha_out = sha256_bytes(out)
    save_temp_file("output.xlsx", out)

    log.info("xlsx_update_ok", extra={
        "sha_in": sha_in,
        "sha_out": sha_out,
        "in_bytes": len(b),
        "out_bytes": len(out),
        "src_filename": file.filename,
    })

    return Response(
        content=out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=edited.xlsx"},
    )


# ----------------------------
# PDF
# ----------------------------

@app.post("/pdf/remove_pages")
async def pdf_remove(
    file: UploadFile = File(...),
    remove_pages_json: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        remove_pages = _loads_json_str(remove_pages_json)  # e.g. [1,2] (1-based)
        out = pdf_remove_pages(b, remove_pages)
    except HTTPException:
        raise
    except Exception as e:
        log.info("pdf_remove_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))

    sha_out = sha256_bytes(out)
    save_temp_file("output.pdf", out)

    log.info("pdf_remove_ok", extra={
        "sha_in": sha_in,
        "sha_out": sha_out,
        "in_bytes": len(b),
        "out_bytes": len(out),
        "src_filename": file.filename,
    })

    return Response(
        content=out,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=edited.pdf"},
    )


@app.post("/pdf/merge")
async def pdf_merge(
    files: List[UploadFile] = File(...),
    debug_meta: bool = Form(False),
    x_api_key: str | None = Header(default=None),
):
    """
    Merge multiple PDFs into one PDF.
    If debug_meta=true: returns JSON with page counts instead of merged PDF.
    """
    _auth(x_api_key)

    try:
        from pypdf import PdfReader, PdfWriter  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"pypdf not available: {e}")

    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Provide at least 2 PDF files to merge.")

    writer = PdfWriter()
    total_in = 0

    inputs_meta = []
    sha_inputs: List[str] = []
    out_pages = 0

    for f in files:
        b = await _read_upload_limited(f)
        total_in += len(b)
        sha = sha256_bytes(b)
        sha_inputs.append(sha)

        try:
            reader = PdfReader(io.BytesIO(b))
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to open PDF '{f.filename}': {e}")

        # encryption info (pypdf varies by version)
        is_encrypted = getattr(reader, "is_encrypted", False)

        # if encrypted and not decryptable, fail fast
        if is_encrypted:
            try:
                # try empty password
                reader.decrypt("")  # type: ignore[attr-defined]
            except Exception:
                raise HTTPException(status_code=400, detail=f"PDF '{f.filename}' is encrypted and cannot be merged without password.")

        pages_in = len(reader.pages)
        inputs_meta.append({
            "filename": f.filename,
            "sha256": sha,
            "bytes": len(b),
            "pages": pages_in,
            "encrypted": bool(is_encrypted),
        })

        if pages_in == 0:
            raise HTTPException(status_code=400, detail=f"PDF '{f.filename}' has 0 pages (cannot merge).")

        for page in reader.pages:
            writer.add_page(page)
            out_pages += 1

    if debug_meta:
        return JSONResponse(content={
            "files": inputs_meta,
            "out_pages": out_pages,
            "total_in_bytes": total_in,
        })

    out_buf = io.BytesIO()
    writer.write(out_buf)
    out = out_buf.getvalue()

    save_temp_file("merged.pdf", out)

    log.info("pdf_merge_ok", extra={
        "files_count": len(files),
        "sha_inputs": sha_inputs,
        "in_bytes": total_in,
        "out_bytes": len(out),
        "sha_out": sha256_bytes(out),
        "out_pages": out_pages,
        "inputs_meta": inputs_meta,  # ok: no PII, only filenames/hashes
    })

    return Response(
        content=out,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=merged.pdf"},
    )


# ----------------------------
# TXT / MD / CSV (Text tools)
# ----------------------------

@app.post("/text/replace_one")
async def text_replace_one(
    file: UploadFile = File(...),
    from_text: str = Form(...),
    to_text: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        text = _decode_text(b)
        out_text = text.replace(from_text, to_text)
        out = out_text.encode("utf-8")
        save_temp_file("output.txt", out)

        log.info("text_replace_one_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
        })

        # keep extension if provided
        ext = _guess_ext(file.filename or "")
        out_name = f"edited.{ext}" if ext in ("txt", "md", "csv") else "edited.txt"

        return Response(
            content=out,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={out_name}"},
        )
    except Exception as e:
        log.info("text_replace_one_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/text/replace")
async def text_replace(
    file: UploadFile = File(...),
    replacements_json: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    try:
        replacements_raw = _loads_json_str(replacements_json)
        replacements = _validate_replacements(replacements_raw)

        text = _decode_text(b)
        out_text = text
        for r in replacements:
            out_text = out_text.replace(r["from"], r["to"])

        out = out_text.encode("utf-8")
        save_temp_file("output.txt", out)

        log.info("text_replace_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
            "repl_count": len(replacements),
        })

        ext = _guess_ext(file.filename or "")
        out_name = f"edited.{ext}" if ext in ("txt", "md", "csv") else "edited.txt"

        return Response(
            content=out,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={out_name}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("text_replace_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/text/apply_ops")
async def text_apply_ops(
    file: UploadFile = File(...),
    ops_json: str = Form(...),
    x_api_key: str | None = Header(default=None),
):
    """
    Apply deterministic edit operations to txt/md/csv.
    ops_json must be a JSON list of ops.
    """
    _auth(x_api_key)

    b = await _read_upload_limited(file)
    sha_in = sha256_bytes(b)

    ext = _guess_ext(file.filename or "")
    if ext not in ("txt", "md", "csv"):
        raise HTTPException(status_code=400, detail="Only txt/md/csv supported for text/apply_ops")

    try:
        ops_raw = _loads_json_str(ops_json)
        if not isinstance(ops_raw, list):
            raise HTTPException(status_code=400, detail="ops_json must be a JSON list")
        text = _decode_text(b)
        out_text = _apply_text_ops(text, ops_raw)
        out = out_text.encode("utf-8")
        save_temp_file(f"output.{ext}", out)

        log.info("text_apply_ops_ok", extra={
            "sha_in": sha_in,
            "sha_out": sha256_bytes(out),
            "in_bytes": len(b),
            "out_bytes": len(out),
            "src_filename": file.filename,
            "ops_count": len(ops_raw),
        })

        return Response(
            content=out,
            media_type="text/plain; charset=utf-8" if ext != "md" else "text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename=edited.{ext}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        log.info("text_apply_ops_failed", extra={"sha_in": sha_in, "err": str(e)})
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/text/create")
async def text_create(
    content: str = Form(...),
    filename: str = Form("output.md"),
    x_api_key: str | None = Header(default=None),
):
    """
    Create a text file (md/txt/csv) from content.
    Used after the LLM created a "Masterkontext" and we need a download.
    """
    _auth(x_api_key)

    out = (content or "").encode("utf-8")
    save_temp_file(filename, out)

    # naive mime
    ext = _guess_ext(filename)
    mime = "text/markdown; charset=utf-8" if ext == "md" else "text/plain; charset=utf-8"

    log.info("text_create_ok", extra={
        "sha_out": sha256_bytes(out),
        "out_bytes": len(out),
        "out_name": filename,
    })

    return Response(
        content=out,
        media_type=mime,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/text/render_pdf")
async def text_render_pdf(
    content: str = Form(...),
    filename: str = Form("output.pdf"),
    x_api_key: str | None = Header(default=None),
):
    """
    Optional: render plain text/markdown into a simple PDF.
    Requires reportlab installed in the worker.
    """
    _auth(x_api_key)

    out = _render_text_to_pdf_bytes(content or "")
    save_temp_file(filename, out)

    log.info("text_render_pdf_ok", extra={
        "sha_out": sha256_bytes(out),
        "out_bytes": len(out),
        "out_name": filename,
    })

    return Response(
        content=out,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ----------------------------
# Multi-file: extract & combine
# ----------------------------

@app.post("/bundle/extract_text")
async def bundle_extract_text(
    files: List[UploadFile] = File(...),
    x_api_key: str | None = Header(default=None),
):
    """
    Extract text from multiple files (docx/pdf/xlsx/txt/md/csv).
    Returns JSON: list of {filename, ext, sha, bytes, text}
    NOTE: No raw content logged.
    """
    _auth(x_api_key)

    if not files:
        raise HTTPException(status_code=400, detail="No files provided.")

    results: List[Dict[str, Any]] = []

    for f in files:
        b = await _read_upload_limited(f)
        sha = sha256_bytes(b)
        ext, text = _extract_text_by_ext(f.filename or "", b)

        results.append({
            "filename": f.filename,
            "ext": ext,
            "sha": sha,
            "bytes": len(b),
            "text": text,
        })

    log.info("bundle_extract_text_ok", extra={
        "files_count": len(files),
        "total_in_bytes": sum(r["bytes"] for r in results),
    })

    return JSONResponse(content={"files": results})


@app.post("/bundle/to_md")
async def bundle_to_md(
    files: List[UploadFile] = File(...),
    title: str = Form("Masterkontext"),
    mode: str = Form("raw"),
    x_api_key: str | None = Header(default=None),
):
    """
    Modes:
    - raw: simple merged markdown corpus (legacy behavior)
    - rag_mastercontext: RAG-friendly structure with inventory + heuristics + full text
    """
    _auth(x_api_key)

    if not files:
        raise HTTPException(status_code=400, detail="No files provided.")

    docs: List[Dict[str, Any]] = []
    total_in = 0

    for f in files:
        b = await _read_upload_limited(f)
        total_in += len(b)
        ext, text = _extract_text_by_ext(f.filename or "", b)

        docs.append(
            {
                "filename": f.filename or "unbekannt",
                "ext": ext,
                "bytes": len(b),
                "sha": sha256_bytes(b),
                "text": text if text else "",
            }
        )

    def _first_informative_lines(text: str, limit: int = 6) -> List[str]:
        out: List[str] = []
        for raw in (text or "").splitlines():
            line = re.sub(r"\s+", " ", raw).strip()
            if not line:
                continue
            if len(line) < 25:
                continue
            if line in out:
                continue
            out.append(line)
            if len(out) >= limit:
                break
        return out

    def _top_keywords(text: str, limit: int = 20) -> List[str]:
        stop = {
            "und", "oder", "der", "die", "das", "ein", "eine", "den", "dem", "des",
            "mit", "für", "von", "auf", "im", "in", "ist", "sind", "zu", "am",
            "an", "als", "bei", "auch", "nicht", "nur", "durch", "aus", "dass",
            "the", "and", "for", "with", "from", "this", "that", "you", "your",
        }
        freq: Dict[str, int] = {}
        for token in re.findall(r"[A-Za-zÄÖÜäöüß0-9][A-Za-zÄÖÜäöüß0-9_-]{2,}", text or ""):
            t = token.lower()
            if t in stop:
                continue
            freq[t] = freq.get(t, 0) + 1
        ranked = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return [k for k, _ in ranked[:limit]]

    mode_key = (mode or "raw").strip().lower()
    sections: List[str]

    if mode_key == "rag_mastercontext":
        all_text = "\n".join(d["text"] for d in docs if d["text"])
        global_keywords = _top_keywords(all_text, limit=20)

        sections = [f"# {title}".strip(), ""]
        sections.append("## Ziel")
        sections.append(
            "Dieses Dokument ist fuer RAG vorbereitet: Quellenuebersicht, heuristische Kernaussagen, "
            "Schluesselbegriffe und extrahierter Volltext pro Datei."
        )
        sections.append("")
        sections.append("## Quellen")
        sections.append("| Datei | Typ | Groesse (Bytes) | SHA256 |")
        sections.append("|---|---:|---:|---|")
        for d in docs:
            safe_name = str(d["filename"]).replace("|", "\\|")
            sections.append(f"| {safe_name} | {d['ext']} | {d['bytes']} | `{d['sha']}` |")
        sections.append("")

        sections.append("## Globale Schluesselbegriffe (heuristisch)")
        if global_keywords:
            sections.append(", ".join(f"`{k}`" for k in global_keywords))
        else:
            sections.append("_(keine Schluesselbegriffe erkannt)_")
        sections.append("")

        sections.append("## Kernaussagen je Quelle (heuristisch)")
        for d in docs:
            sections.append(f"### {d['filename']} ({d['ext']})")
            key_lines = _first_informative_lines(d["text"], limit=6)
            if key_lines:
                for line in key_lines:
                    sections.append(f"- {line}")
            else:
                sections.append("- _(keine verwertbaren Zeilen gefunden)_")
            sections.append("")

        sections.append("## Volltext je Quelle (extrahiert)")
        for d in docs:
            sections.append(f"### {d['filename']} ({d['ext']})")
            sections.append("")
            sections.append(d["text"] if d["text"] else "_(kein extrahierbarer Text gefunden)_")
            sections.append("")
    else:
        sections = [f"# {title}".strip(), ""]
        for d in docs:
            sections.append(f"## Datei: {d['filename']} ({d['ext']})")
            sections.append("")
            sections.append(d["text"] if d["text"] else "_(kein extrahierbarer Text gefunden)_")
            sections.append("")

    md = "\n".join(sections).strip() + "\n"
    out = md.encode("utf-8")
    save_temp_file("bundle.md", out)

    log.info("bundle_to_md_ok", extra={
        "files_count": len(files),
        "in_bytes": total_in,
        "out_bytes": len(out),
        "sha_out": sha256_bytes(out),
    })

    return Response(
        content=out,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=masterkontext.md"},
    )
