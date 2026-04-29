from io import BytesIO
from docx import Document

def docx_simple_replace(docx_bytes: bytes, replacements: list[dict]) -> bytes:
    """
    replacements: [{"from":"Autohaus", "to":"Autohaus KAHLE"}, ...]
    MVP: replace in paragraphs + table cells. (Header/Footer optional later)
    """
    doc = Document(BytesIO(docx_bytes))

    def repl_text(text: str) -> str:
        out = text
        for r in replacements:
            out = out.replace(r["from"], r["to"])
        return out

    for p in doc.paragraphs:
        if p.runs:
            full = "".join(run.text for run in p.runs)
            new = repl_text(full)
            if new != full:
                # rewrite runs conservatively: set first run, clear rest
                p.runs[0].text = new
                for i in range(1, len(p.runs)):
                    p.runs[i].text = ""

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.runs:
                        full = "".join(run.text for run in p.runs)
                        new = repl_text(full)
                        if new != full:
                            p.runs[0].text = new
                            for i in range(1, len(p.runs)):
                                p.runs[i].text = ""

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
