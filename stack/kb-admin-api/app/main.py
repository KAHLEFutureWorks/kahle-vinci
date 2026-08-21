from __future__ import annotations

import base64
import asyncio
import contextvars
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import shutil
import tempfile
import threading
import time
import unicodedata
import uuid
from dataclasses import asdict
from datetime import date, datetime, timedelta
from pathlib import Path, PurePosixPath
from typing import Any, Callable

import requests
from docx import Document
from fastapi import BackgroundTasks, Depends, FastAPI, File, Form, Header, HTTPException, Request, Response, UploadFile
from fastapi.responses import FileResponse, Response as FastAPIResponse
from pydantic import BaseModel, Field
from pypdf import PdfReader


UPLOAD_CONVERSION_PROGRESS: contextvars.ContextVar[Callable[[int, int], None] | None] = (
    contextvars.ContextVar("upload_conversion_progress", default=None)
)
UPLOAD_JOB_CONTEXT: contextvars.ContextVar[dict[str, Any] | None] = contextvars.ContextVar(
    "upload_job_context", default=None,
)


try:
    from .portal_governance import (
        GovernanceError,
        PortalGovernance,
        SQLiteGovernanceStore,
        serialize as serialize_governance,
    )
except ImportError:  # pragma: no cover - supports the existing file-based contract harness
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from portal_governance import (  # type: ignore[no-redef]
        GovernanceError,
        PortalGovernance,
        SQLiteGovernanceStore,
        serialize as serialize_governance,
    )

try:
    from .document_lifecycle import Analysis, DocumentLifecycle, LifecycleError, workdays_until
    from .audit_export import AuditExporter
    from .maintenance import MaintenanceError, MaintenanceService
    from .quality_cases import QualityCaseError, QualityCaseService
    from .legacy_migration import LegacyMigrationService
    from .markdown_correction import IonosMarkdownCorrector, MarkdownCorrectionError, MarkdownCorrectionService
    from .quality_dashboard import QualityDashboard
    from .document_changes import DocumentChangeError, DocumentChangeService
    from .ownership import OwnershipError, OwnershipService
    from .upload_jobs import UploadJobError, UploadJobService, UploadSpool
    from .decision_jobs import DecisionJobError, DecisionJobQueue
    from .document_authority import AuthorityError, DocumentAuthorityService
    from .rag_metadata import RAGMetadataWriter
    from .content_classification import ContentConfidentialityClassifier
    from .restricted_terms import RestrictedTermError, RestrictedTermService
    from .global_analysis import (
        CorpusDocument, GlobalAnalysisError, GlobalCorpus, GlobalDocumentAnalyzer, IonosEmbeddingProvider,
    )
    from .secure_ingest import (
        ClamAVScanner, DocumentWorkerAdapter, IngestError, PromptInjectionInspector,
        QuarantineStorage, ScreenshotInspector, SecureFileInspector, SecureIngestPipeline,
    )
except ImportError:  # pragma: no cover
    from document_lifecycle import Analysis, DocumentLifecycle, LifecycleError, workdays_until
    from audit_export import AuditExporter
    from maintenance import MaintenanceError, MaintenanceService
    from quality_cases import QualityCaseError, QualityCaseService
    from legacy_migration import LegacyMigrationService
    from markdown_correction import IonosMarkdownCorrector, MarkdownCorrectionError, MarkdownCorrectionService
    from quality_dashboard import QualityDashboard
    from document_changes import DocumentChangeError, DocumentChangeService
    from ownership import OwnershipError, OwnershipService
    from upload_jobs import UploadJobError, UploadJobService, UploadSpool
    from decision_jobs import DecisionJobError, DecisionJobQueue
    from document_authority import AuthorityError, DocumentAuthorityService
    from rag_metadata import RAGMetadataWriter
    from content_classification import ContentConfidentialityClassifier
    from restricted_terms import RestrictedTermError, RestrictedTermService
    from global_analysis import (
        CorpusDocument, GlobalAnalysisError, GlobalCorpus, GlobalDocumentAnalyzer, IonosEmbeddingProvider,
    )
    from secure_ingest import (
        ClamAVScanner, DocumentWorkerAdapter, IngestError, PromptInjectionInspector,
        QuarantineStorage, ScreenshotInspector, SecureFileInspector, SecureIngestPipeline,
    )

APP_VERSION = "0.2.0"
SUPPORTED_EXTENSIONS = {".md", ".txt", ".csv", ".pdf", ".docx"}
EDITABLE_EXTENSIONS = {".md", ".txt", ".csv"}
KB_ROOT = Path(os.getenv("KB_ROOT", "/knowledgebases")).resolve()
KB_STATE_PATH = Path(os.getenv("KB_STATE_PATH", "/state/kb-sync-state.json")).resolve()
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333").rstrip("/")
OPENWEBUI_URL = os.getenv("OPENWEBUI_URL", "http://open-webui:8080").rstrip("/")
KB_SYNC_URL = os.getenv("KB_SYNC_URL", "http://kb-sync:8093").rstrip("/")
IONOS_BASE_URL = os.getenv(
    "IONOS_OPENAI_BASE_URL", "https://openai.inference.de-txl.ionos.com/v1"
).rstrip("/")
IONOS_API_KEY = os.getenv("IONOS_API_KEY", "").strip()
EMBEDDING_MODEL = os.getenv("IONOS_EMBEDDING_MODEL", "BAAI/bge-m3").strip()
COLLECTIONS = tuple(
    item.strip()
    for item in os.getenv(
        "KB_ADMIN_COLLECTIONS", "kahleallgemein,kahlekontext,kahlerichtlinien"
    ).split(",")
    if item.strip()
)
COLLECTION_LABELS = {
    "kahleallgemein": "Allgemeines Wissen",
    "kahlekontext": "Standorte & Unternehmen",
    "kahlerichtlinien": "Richtlinien & Prozesse",
}
DEV_AUTH_BYPASS = os.getenv("KB_ADMIN_DEV_AUTH_BYPASS", "false").lower() == "true"
MAX_UPLOAD_BYTES = int(os.getenv("KB_ADMIN_MAX_UPLOAD_BYTES", str(25 * 1024 * 1024)))
TRASH_RETENTION_DAYS = max(1, int(os.getenv("KB_ADMIN_TRASH_RETENTION_DAYS", "30")))
MAINTENANCE_API_KEY = os.getenv("KB_ADMIN_MAINTENANCE_API_KEY", "").strip()
KB_SYNC_INTERNAL_API_KEY = os.getenv("KB_SYNC_INTERNAL_API_KEY", MAINTENANCE_API_KEY).strip()
PORTAL_ALLOWED_EMAIL_DOMAINS = {
    item.strip().casefold() for item in os.getenv("PORTAL_ALLOWED_EMAIL_DOMAINS", "kahle.de").split(",")
    if item.strip()
}
UNLOCK_CODE_HASH = os.getenv("KB_ADMIN_UNLOCK_CODE_HASH", "").strip()
UNLOCK_SESSION_SECRET = os.getenv("KB_ADMIN_UNLOCK_SESSION_SECRET", "").strip()
UNLOCK_TTL_SECONDS = max(300, int(os.getenv("KB_ADMIN_UNLOCK_TTL_SECONDS", "28800")))
UNLOCK_MAX_ATTEMPTS = max(1, int(os.getenv("KB_ADMIN_UNLOCK_MAX_ATTEMPTS", "5")))
UNLOCK_BLOCK_SECONDS = max(60, int(os.getenv("KB_ADMIN_UNLOCK_BLOCK_SECONDS", "900")))
UNLOCK_COOKIE = "kahle_vector_unlock"
UNLOCK_ENABLED = bool(
    UNLOCK_CODE_HASH.startswith("pbkdf2_sha256.") and len(UNLOCK_SESSION_SECRET) >= 43
)
_unlock_failures: dict[str, list[float]] = {}
_unlock_failures_lock = threading.Lock()
VERSIONS_ROOT = KB_ROOT / ".versions"
TRASH_ROOT = KB_ROOT / ".trash"
AUDIT_ROOT = KB_ROOT / ".admin"
AUDIT_PATH = AUDIT_ROOT / "audit.jsonl"
PORTAL_DB_PATH = Path(
    os.getenv("KB_PORTAL_DB_PATH", "/portal-data/wissensportal.sqlite3")
).resolve()
PORTAL_GOVERNANCE = PortalGovernance(SQLiteGovernanceStore(PORTAL_DB_PATH))
_AUTO_ACTIVATION_DEFAULT = os.getenv("KB_PORTAL_AUTO_ACTIVATION_ENABLED", "false").strip().lower() == "true"
DOCUMENT_LIFECYCLE = DocumentLifecycle(
    PORTAL_GOVERNANCE.store, PORTAL_GOVERNANCE,
    auto_activation_enabled=lambda: PORTAL_GOVERNANCE.setting_bool(
        "auto_activation_enabled", default=_AUTO_ACTIVATION_DEFAULT,
    ),
)
MAINTENANCE = MaintenanceService(PORTAL_GOVERNANCE.store)
AUDIT_EXPORTER = AuditExporter(PORTAL_GOVERNANCE.store)
QUALITY_CASES = QualityCaseService(PORTAL_GOVERNANCE.store)
RESTRICTED_TERMS = RestrictedTermService(PORTAL_GOVERNANCE)
DOCUMENT_CHANGES = DocumentChangeService(PORTAL_GOVERNANCE.store, PORTAL_GOVERNANCE)
OWNERSHIP = OwnershipService(PORTAL_GOVERNANCE.store, PORTAL_GOVERNANCE)
DECISION_JOBS = DecisionJobQueue(PORTAL_GOVERNANCE.store)
DOCUMENT_AUTHORITY = DocumentAuthorityService(PORTAL_GOVERNANCE.store, PORTAL_GOVERNANCE)
GLOBAL_CORPUS = GlobalCorpus(PORTAL_GOVERNANCE.store)
GLOBAL_ANALYZER = GlobalDocumentAnalyzer(
    GLOBAL_CORPUS,
    IonosEmbeddingProvider(IONOS_BASE_URL, IONOS_API_KEY, EMBEDDING_MODEL) if IONOS_API_KEY else None,
)
PORTAL_FILES_ROOT = Path(os.getenv("KB_PORTAL_FILES_ROOT", "/portal-data/files"))
UPLOAD_JOBS = UploadJobService(PORTAL_GOVERNANCE.store)
UPLOAD_SPOOL = UploadSpool(PORTAL_FILES_ROOT / ".upload-spool")
RAG_METADATA = RAGMetadataWriter(PORTAL_GOVERNANCE.store, PORTAL_FILES_ROOT)
CONFIDENTIALITY_CLASSIFIER = ContentConfidentialityClassifier()
SECURE_INGEST = SecureIngestPipeline(
    SecureFileInspector(max_bytes=MAX_UPLOAD_BYTES),
    ClamAVScanner(host=os.getenv("KB_CLAMAV_HOST", "clamav"), port=int(os.getenv("KB_CLAMAV_PORT", "3310"))),
    DocumentWorkerAdapter(os.getenv("DOCUMENT_WORKER_URL", "http://document-worker:8090"), os.getenv("DOCUMENT_WORKER_API_KEY", "")),
    QuarantineStorage(PORTAL_FILES_ROOT),
    PromptInjectionInspector(),
)
LEGACY_MIGRATION = LegacyMigrationService(
    PORTAL_GOVERNANCE, DOCUMENT_LIFECYCLE, GLOBAL_ANALYZER, GLOBAL_CORPUS,
    QuarantineStorage(PORTAL_FILES_ROOT), SECURE_INGEST,
    restricted_term_matcher=RESTRICTED_TERMS.matches,
)
MARKDOWN_CORRECTION = MarkdownCorrectionService(
    PORTAL_GOVERNANCE, DOCUMENT_LIFECYCLE, GLOBAL_ANALYZER, GLOBAL_CORPUS,
    QuarantineStorage(PORTAL_FILES_ROOT),
    IonosMarkdownCorrector(
        IONOS_BASE_URL, IONOS_API_KEY,
        os.getenv("IONOS_CHAT_MODEL_DEFAULT", "mistralai/Mistral-Small-24B-Instruct"),
    ) if IONOS_API_KEY else None,
    restricted_term_matcher=RESTRICTED_TERMS.matches,
)
QUALITY_DASHBOARD = QualityDashboard(PORTAL_GOVERNANCE.store, Path(os.getenv("KB_BACKUP_STATE_PATH", "/backups/primary/backup-state.json")))
app = FastAPI(
    title="KAHLE Vector Admin API",
    version=APP_VERSION,
    docs_url=None,
    redoc_url=None,
)


class SaveDocumentRequest(BaseModel):
    content: str = Field(..., max_length=2_000_000)


class MoveDocumentRequest(BaseModel):
    target_collection: str
    target_path: str


class RestoreVersionRequest(BaseModel):
    version_id: str


class CreateCollectionRequest(BaseModel):
    id: str = Field(..., min_length=2, max_length=48, pattern=r"^[a-z0-9][a-z0-9_-]+$")
    label: str = Field(..., min_length=2, max_length=80)


class UpdateCollectionRequest(BaseModel):
    label: str = Field(..., min_length=2, max_length=80)


class DeleteCollectionRequest(BaseModel):
    confirm_id: str


class PurgeCollectionRequest(BaseModel):
    confirm_id: str


class UnlockRequest(BaseModel):
    code: str = Field(..., min_length=8, max_length=256)


class PortalRoleRequest(BaseModel):
    role: str = Field(..., pattern=r"^(employee|manager|admin|portal_admin)$")
    confirmed: bool = False


class PortalActivationRequest(BaseModel):
    active: bool


class PortalManagerRequest(BaseModel):
    manager_user_id: str | None = None


class PortalDelegationRequest(BaseModel):
    manager_user_id: str = Field(..., min_length=1, max_length=100)
    delegate_user_id: str = Field(..., min_length=1, max_length=100)
    valid_from: str | None = None
    valid_until: str | None = None


class PortalAbsenceRequest(BaseModel):
    manager_user_id: str = Field(..., min_length=1, max_length=100)
    delegate_user_id: str | None = Field(default=None, min_length=1, max_length=100)
    absent_from: str | None = None
    absent_until: str | None = None
    reason: str = Field(..., max_length=1000)


class PortalAccessRequest(BaseModel):
    knowledgebase_id: str = Field(..., min_length=1, max_length=100)
    can_read: bool
    can_upload: bool


class PortalRagFeedbackRequest(BaseModel):
    reason: str = Field(..., max_length=60)
    comment: str = Field("", max_length=2000)
    question: str = Field(..., max_length=8000)
    answer: str = Field(..., max_length=16000)
    sources: list[dict[str, Any]] = Field(default_factory=list, max_length=30)
    passages: list[dict[str, Any]] = Field(default_factory=list, max_length=30)
    runtime: dict[str, Any] = Field(default_factory=dict)
    request_id: str = Field(..., min_length=1, max_length=200)
    document_ids: list[str] = Field(default_factory=list, max_length=30)
    knowledgebase_ids: list[str] = Field(default_factory=list, max_length=30)


class PortalRetrievalEventRequest(BaseModel):
    user_id: str = Field(..., min_length=1, max_length=100)
    query_hash: str = Field(..., pattern=r"^[0-9a-fA-F]{64}$")
    found: bool
    source_count: int = Field(..., ge=0, le=30)
    latency_ms: int = Field(..., ge=0, le=300_000)
    error_code: str | None = Field(default=None, max_length=100)


class PortalIncidentCommentRequest(BaseModel):
    comment: str = Field(..., min_length=3, max_length=2000)


class PortalQualityCaseMessageRequest(BaseModel):
    message: str = Field(..., min_length=3, max_length=2000)


class PortalQualityCaseResolveRequest(BaseModel):
    resolution: str = Field(..., min_length=3, max_length=2000)


class PortalMigrationStageRequest(BaseModel):
    path: str = Field(..., min_length=1, max_length=500)
    confirmed: bool = False


class PortalMigrationDispositionRequest(BaseModel):
    path: str = Field(..., min_length=1, max_length=500)
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalRestrictedTermRequest(BaseModel):
    term: str = Field(..., min_length=2, max_length=120)


class PortalAutoActivationRequest(BaseModel):
    enabled: bool
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalMigrationMetadataRequest(BaseModel):
    path: str = Field(..., min_length=1, max_length=500)
    owner_email: str = Field(..., min_length=3, max_length=320)
    confidentiality: str = Field(..., pattern=r"^(internal|restricted|confidential)$")
    authority_type: str = Field(..., min_length=3, max_length=80)
    authority_level: int = Field(..., ge=1, le=6)
    knowledgebase_id: str | None = Field(default=None, min_length=1, max_length=100)
    scope: dict[str, Any] = Field(default_factory=dict)


class PortalMarkdownRevisionRequest(BaseModel):
    instruction: str = Field("", max_length=4000)
    replacement_markdown: str = Field("", max_length=2_000_000)
    reason: str = Field(..., min_length=3, max_length=2000)
    confirmed: bool


class PortalRemovalRequest(BaseModel):
    document_id: str = Field(..., min_length=1, max_length=100)
    kind: str = Field(..., pattern=r"^(deactivate|delete)$")
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalRemovalDecisionRequest(BaseModel):
    approve: bool
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalRestoreRequest(BaseModel):
    reason: str = Field(..., min_length=3, max_length=2000)
    confirmed: bool = False


class PortalArchivedVersionRestoreRequest(BaseModel):
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalLegalHoldRequest(BaseModel):
    enabled: bool
    reason: str = Field(..., min_length=3, max_length=2000)
    review_at: str | None = None


class PortalRenewalRequest(BaseModel):
    document_id: str = Field(..., min_length=1, max_length=100)
    reason: str = Field(..., min_length=3, max_length=2000)
    confirmed: bool


class PortalConfidentialityRequest(BaseModel):
    document_id: str = Field(..., min_length=1, max_length=100)
    desired: str = Field(..., pattern=r"^(internal|restricted|confidential)$")
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalDocumentChangeDecisionRequest(BaseModel):
    approve: bool
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalOwnershipProposalRequest(BaseModel):
    proposed_owner_user_id: str = Field(..., min_length=1, max_length=200)
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalOwnerPermissionRequest(BaseModel):
    allowed: bool


class PortalAuthorityRequest(BaseModel):
    authority_type: str = Field(..., min_length=3, max_length=80)
    scope: dict[str, Any] = Field(default_factory=dict)
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalAuthorityRelationRequest(BaseModel):
    target_document_id: str = Field(..., min_length=1, max_length=100)
    relation_type: str = Field(..., pattern=r"^(supersedes|overrides|applies_only_if|related_to)$")
    condition_text: str = Field("", max_length=2000)
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalOwnershipConfirmationRequest(BaseModel):
    accept: bool
    reason: str = Field(..., min_length=3, max_length=2000)


class PortalKnowledgebaseChangeRequest(BaseModel):
    kind: str = Field(..., pattern=r"^(create|rename|archive|delete)$")
    knowledgebase_id: str | None = None
    payload: dict[str, Any] = Field(default_factory=dict)
    confirmed: bool = False

class PortalKnowledgebaseDecisionRequest(BaseModel):
    approve: bool
    reason: str = Field(..., min_length=3, max_length=1000)
    confirmed: bool = False

class PortalRetrievalScopeRequest(BaseModel):
    user_id: str = Field(..., min_length=1, max_length=200)


class PortalCaseActionRequest(BaseModel):
    action: str = Field(..., pattern=r"^(create|replace|publish_existing|discard)$")
    target_document_id: str | None = Field(None, max_length=100)


class PortalCaseDecisionRequest(BaseModel):
    decision: str = Field(..., pattern=r"^(approve|reject|escalate)$")
    reason: str = Field(default="", max_length=2000)


class PortalCaseInquiryRequest(BaseModel):
    recipient_user_id: str = Field(..., min_length=1, max_length=200)
    question: str = Field(..., min_length=3, max_length=2000)


class PortalNotificationReplyRequest(BaseModel):
    message: str = Field(..., min_length=3, max_length=2000)


class PortalCaseTargetRequest(BaseModel):
    knowledgebase_id: str = Field(..., min_length=1, max_length=200)


class PortalCaseTargetsRequest(BaseModel):
    knowledgebase_ids: list[str] = Field(..., min_length=1, max_length=100)


class PortalUploadResponse(BaseModel):
    case_id: str
    document_id: str
    version_id: str
    status: str
    owner_email: str
    prompt_injection_risk: str
    confidentiality: str
    confidentiality_reason: str
    requires_admin: bool
    owner_confirmation_required: bool = False
    conversion_quality: str = "good"
    conversion_issues: list[str] = Field(default_factory=list)
    restricted_terms: list[str] = Field(default_factory=list)
    exact_duplicate_document_id: str | None = None
    matches: list[dict[str, Any]] = Field(default_factory=list)


def _collection_names() -> tuple[str, ...]:
    discovered = {
        path.name
        for path in KB_ROOT.iterdir()
        if path.is_dir() and not path.name.startswith(".")
    } if KB_ROOT.exists() else set()
    return tuple(sorted(set(COLLECTIONS) | discovered))


def _collection_label(name: str) -> str:
    metadata_path = KB_ROOT / name / ".collection.json"
    try:
        payload = json.loads(metadata_path.read_text(encoding="utf-8"))
        label = str(payload.get("label") or "").strip()
        if label:
            return label
    except Exception:
        pass
    return COLLECTION_LABELS.get(name, name)


def _safe_collection(name: str) -> str:
    if name not in _collection_names():
        raise HTTPException(status_code=404, detail="unknown_collection")
    return name


def _safe_relative_path(raw: str, *, require_supported: bool = True) -> PurePosixPath:
    value = str(raw or "").replace("\\", "/").strip().lstrip("/")
    path = PurePosixPath(value)
    if not value or path.is_absolute() or any(part in {"", ".", ".."} for part in path.parts):
        raise HTTPException(status_code=400, detail="invalid_relative_path")
    if any(part.startswith(".") for part in path.parts):
        raise HTTPException(status_code=400, detail="hidden_paths_not_allowed")
    if require_supported and path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="unsupported_file_type")
    return path


def _document_path(collection: str, relative_path: str, *, must_exist: bool = True) -> Path:
    collection = _safe_collection(collection)
    rel = _safe_relative_path(relative_path)
    root = (KB_ROOT / collection).resolve()
    candidate = (root / Path(*rel.parts)).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="path_outside_collection") from exc
    if must_exist and (not candidate.exists() or not candidate.is_file()):
        raise HTTPException(status_code=404, detail="document_not_found")
    return candidate


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_preview(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in EDITABLE_EXTENSIONS:
        return _read_text(path)
    if suffix == ".pdf":
        pages: list[str] = []
        for index, page in enumerate(PdfReader(str(path)).pages[:8], start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"[Seite {index}]\n{text}")
        return "\n\n".join(pages)
    if suffix == ".docx":
        doc = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs[:300])
    return ""


def _frontmatter(text: str) -> dict[str, Any]:
    if not text.startswith("---"):
        return {}
    match = re.search(r"^---\s*$", text[3:], flags=re.MULTILINE)
    if not match:
        return {}
    result: dict[str, Any] = {}
    current_list: str | None = None
    for raw_line in text[3 : 3 + match.start()].splitlines():
        line = raw_line.rstrip()
        list_match = re.match(r"^\s*-\s*(.+?)\s*$", line)
        if list_match and current_list:
            result.setdefault(current_list, []).append(list_match.group(1).strip(" \"'"))
            continue
        field = re.match(r"^\s*([A-Za-z0-9_-]+)\s*:\s*(.*?)\s*$", line)
        if not field:
            current_list = None
            continue
        key = field.group(1).lower().replace("-", "_")
        value = field.group(2).strip()
        if not value:
            result[key] = []
            current_list = key
            continue
        current_list = None
        if value.startswith("[") and value.endswith("]"):
            result[key] = [
                item.strip().strip(" \"'") for item in value[1:-1].split(",") if item.strip()
            ]
        else:
            result[key] = value.strip(" \"'")
    return result


def _expiry_from_filename(name: str) -> str:
    folded = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode().lower()
    patterns = (
        r"(?:gueltig|gultig)[\s_-]*bis[\s_:-]*(\d{4}-\d{2}-\d{2})",
        r"(?:gueltig|gultig)[\s_-]*bis[\s_:-]*(\d{2}\.\d{2}\.\d{4})",
    )
    for pattern in patterns:
        match = re.search(pattern, folded)
        if not match:
            continue
        raw = match.group(1)
        try:
            return (
                datetime.strptime(raw, "%Y-%m-%d")
                if "-" in raw
                else datetime.strptime(raw, "%d.%m.%Y")
            ).date().isoformat()
        except ValueError:
            return ""
    return ""


def _normalise_date(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
        try:
            return datetime.strptime(raw, fmt).date().isoformat()
        except ValueError:
            continue
    return ""


def _metadata(path: Path) -> dict[str, Any]:
    text = _read_text(path) if path.suffix.lower() in EDITABLE_EXTENSIONS else ""
    meta = _frontmatter(text)
    valid_until = _normalise_date(
        meta.get("valid_until") or meta.get("gueltig_bis") or meta.get("gültig_bis")
    ) or _expiry_from_filename(path.name)
    days_remaining: int | None = None
    try:
        notify_before_days = max(0, min(365, int(meta.get("notify_before_days") or 14)))
    except (TypeError, ValueError):
        notify_before_days = 14
    expiry_status = "none"
    if valid_until:
        days_remaining = (date.fromisoformat(valid_until) - date.today()).days
        if days_remaining < 0:
            expiry_status = "expired"
        elif days_remaining <= notify_before_days:
            expiry_status = "critical"
        elif days_remaining <= 30:
            expiry_status = "warning"
        else:
            expiry_status = "valid"
    rag_value = str(meta.get("rag_index", "true")).strip().lower()
    return {
        "title": str(meta.get("title") or path.stem).strip(),
        "document_id": str(meta.get("document_id") or "").strip(),
        "owner": str(meta.get("owner") or meta.get("owner_name") or "").strip(),
        "valid_until": valid_until,
        "days_remaining": days_remaining,
        "notify_before_days": notify_before_days,
        "expiry_status": expiry_status,
        "locations": meta.get("standorte") or meta.get("locations") or [],
        "tags": meta.get("tags") or [],
        "rag_index": rag_value not in {"false", "0", "no"},
    }


def _load_state() -> dict[str, Any]:
    try:
        payload = json.loads(KB_STATE_PATH.read_text(encoding="utf-8"))
        return payload if isinstance(payload, dict) else {}
    except Exception:
        return {}


def _state_entry(state: dict[str, Any], collection: str, relative_path: str) -> dict[str, Any]:
    return (
        (((state.get("collections") or {}).get(collection) or {}).get("files") or {}).get(
            relative_path
        )
        or {}
    )


def _index_status(path: Path, entry: dict[str, Any], rag_index: bool) -> str:
    if not rag_index:
        return "excluded"
    if not entry:
        return "pending"
    try:
        return "current" if entry.get("sha256") == _sha256(path) else "pending"
    except OSError:
        return "error"


def _version_key(collection: str, relative_path: str) -> Path:
    encoded = hashlib.sha256(f"{collection}/{relative_path}".encode("utf-8")).hexdigest()[:20]
    return VERSIONS_ROOT / collection / encoded


def _create_version(path: Path, collection: str, relative_path: str, actor: dict[str, Any], action: str) -> str:
    if not path.exists():
        return ""
    version_dir = _version_key(collection, relative_path)
    version_dir.mkdir(parents=True, exist_ok=True)
    version_id = f"{datetime.now().strftime('%Y%m%d-%H%M%S-%f')}-{_sha256(path)[:10]}"
    snapshot = version_dir / f"{version_id}{path.suffix.lower()}"
    shutil.copy2(path, snapshot)
    (version_dir / f"{version_id}.json").write_text(
        json.dumps(
            {
                "id": version_id,
                "collection": collection,
                "path": relative_path,
                "action": action,
                "actor": actor.get("email") or actor.get("name") or "admin",
                "created_at": datetime.now().astimezone().isoformat(),
                "size": snapshot.stat().st_size,
                "sha256": _sha256(snapshot),
                "snapshot": snapshot.name,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return version_id


def _audit(actor: dict[str, Any], action: str, **details: Any) -> None:
    AUDIT_ROOT.mkdir(parents=True, exist_ok=True)
    entry = {
        "timestamp": datetime.now().astimezone().isoformat(),
        "actor": actor.get("email") or actor.get("name") or "admin",
        "action": action,
        **details,
    }
    with AUDIT_PATH.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(entry, ensure_ascii=False) + "\n")


def _atomic_write(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(dir=path.parent, prefix=f".{path.name}.", delete=False) as handle:
        tmp = Path(handle.name)
        handle.write(data)
        handle.flush()
        os.fsync(handle.fileno())
    os.replace(tmp, path)


def require_openwebui_user(request: Request) -> dict[str, Any]:
    if DEV_AUTH_BYPASS:
        return {"id": "local-admin", "name": "Lokale Vorschau", "email": "admin@local", "role": "admin"}
    forwarded: dict[str, str] = {}
    for header in ("authorization", "cookie", "user-agent"):
        value = request.headers.get(header)
        if value:
            forwarded[header] = value
    try:
        response = requests.get(
            f"{OPENWEBUI_URL}/api/v1/auths/",
            headers=forwarded,
            timeout=10,
            allow_redirects=False,
        )
    except requests.RequestException as exc:
        raise HTTPException(status_code=503, detail="auth_service_unavailable") from exc
    if response.status_code != 200:
        raise HTTPException(status_code=401, detail="openwebui_login_required")
    try:
        user = response.json()
    except ValueError as exc:
        raise HTTPException(status_code=502, detail="invalid_auth_response") from exc
    return user


def require_admin(request: Request) -> dict[str, Any]:
    user = require_openwebui_user(request)
    if str(user.get("role") or "").lower() != "admin":
        raise HTTPException(status_code=403, detail="admin_role_required")
    return user


def require_portal_identity(
    user: dict[str, Any] = Depends(require_openwebui_user),
) -> dict[str, Any]:
    user_id = str(user.get("id") or "").strip()
    email = str(user.get("email") or "").strip()
    display_name = str(user.get("name") or user.get("display_name") or email).strip()
    if not user_id or not email:
        raise HTTPException(status_code=502, detail="openwebui_identity_incomplete")
    email_domain = email.rsplit("@", 1)[-1].casefold() if "@" in email else ""
    if not DEV_AUTH_BYPASS and email_domain not in PORTAL_ALLOWED_EMAIL_DOMAINS:
        raise HTTPException(status_code=403, detail="kahle_microsoft_tenant_required")
    try:
        identity = PORTAL_GOVERNANCE.sync_identity(
            user_id=user_id,
            email=email,
            display_name=display_name,
            active=True,
            bootstrap_portal_admin=str(user.get("role") or "").lower() == "admin",
        )
    except GovernanceError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    if not identity.active:
        raise HTTPException(status_code=403, detail="portal_user_inactive")
    return {**serialize_governance(identity), "openwebui_role": user.get("role")}


def _portal_call(call: Callable[[], Any], *, forbidden_status: int = 403) -> Any:
    try:
        return call()
    except GovernanceError as exc:
        detail = str(exc)
        status = 404 if detail.startswith("unknown_") else 409 if detail.endswith("_required") else forbidden_status
        raise HTTPException(status_code=status, detail=detail) from exc

def _b64url_encode(value: bytes) -> str:
    return base64.urlsafe_b64encode(value).decode("ascii").rstrip("=")


def _b64url_decode(value: str) -> bytes:
    return base64.urlsafe_b64decode(value + "=" * (-len(value) % 4))


def _unlock_code_matches(code: str) -> bool:
    try:
        algorithm, iterations_raw, salt_raw, digest_raw = UNLOCK_CODE_HASH.split(".", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        iterations = int(iterations_raw)
        if iterations < 200_000 or iterations > 2_000_000:
            return False
        actual = hashlib.pbkdf2_hmac(
            "sha256", code.encode("utf-8"), _b64url_decode(salt_raw), iterations
        )
        return hmac.compare_digest(actual, _b64url_decode(digest_raw))
    except (TypeError, ValueError):
        return False


def _unlock_subject(admin: dict[str, Any]) -> str:
    return str(admin.get("id") or admin.get("email") or "admin")


def _issue_unlock_token(admin: dict[str, Any]) -> str:
    payload = {
        "sub": _unlock_subject(admin),
        "exp": int(time.time()) + UNLOCK_TTL_SECONDS,
        "nonce": secrets.token_urlsafe(12),
    }
    encoded = _b64url_encode(
        json.dumps(payload, separators=(",", ":"), sort_keys=True).encode("utf-8")
    )
    signature = hmac.new(
        UNLOCK_SESSION_SECRET.encode("utf-8"), encoded.encode("ascii"), hashlib.sha256
    ).digest()
    return f"{encoded}.{_b64url_encode(signature)}"


def _has_valid_unlock(request: Request, admin: dict[str, Any]) -> bool:
    if DEV_AUTH_BYPASS:
        return True
    if not UNLOCK_ENABLED:
        return False
    token = request.cookies.get(UNLOCK_COOKIE, "")
    try:
        encoded, signature_raw = token.split(".", 1)
        expected = hmac.new(
            UNLOCK_SESSION_SECRET.encode("utf-8"), encoded.encode("ascii"), hashlib.sha256
        ).digest()
        if not hmac.compare_digest(expected, _b64url_decode(signature_raw)):
            return False
        payload = json.loads(_b64url_decode(encoded))
        return (
            isinstance(payload, dict)
            and str(payload.get("sub") or "") == _unlock_subject(admin)
            and int(payload.get("exp") or 0) > int(time.time())
        )
    except (TypeError, ValueError, json.JSONDecodeError):
        return False


def _failure_key(request: Request, admin: dict[str, Any]) -> str:
    remote = request.client.host if request.client else "unknown"
    return f"{_unlock_subject(admin)}|{remote}"


def _attempt_state(request: Request, admin: dict[str, Any]) -> tuple[str, list[float]]:
    key = _failure_key(request, admin)
    cutoff = time.time() - UNLOCK_BLOCK_SECONDS
    with _unlock_failures_lock:
        recent = [stamp for stamp in _unlock_failures.get(key, []) if stamp >= cutoff]
        if recent:
            _unlock_failures[key] = recent
        else:
            _unlock_failures.pop(key, None)
    return key, recent


def require_unlocked_admin(
    request: Request, admin: dict[str, Any] = Depends(require_admin)
) -> dict[str, Any]:
    if not DEV_AUTH_BYPASS and not UNLOCK_ENABLED:
        raise HTTPException(status_code=503, detail="admin_unlock_not_configured")
    if not _has_valid_unlock(request, admin):
        raise HTTPException(status_code=423, detail="admin_unlock_required")
    return admin


def _qdrant(method: str, path: str, **kwargs: Any) -> dict[str, Any]:
    try:
        response = requests.request(method, f"{QDRANT_URL}{path}", timeout=45, **kwargs)
    except requests.RequestException as exc:
        raise HTTPException(status_code=503, detail="qdrant_unavailable") from exc
    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=f"qdrant_error_{response.status_code}")
    return response.json() if response.text else {}


def _trigger_reindex(collection: str, relative_path: str = "") -> dict[str, Any]:
    try:
        response = requests.post(
            f"{KB_SYNC_URL}/reindex",
            json={"collection": collection, "path": relative_path},
            headers={"X-API-Key": KB_SYNC_INTERNAL_API_KEY},
            timeout=180,
        )
        if response.status_code >= 400:
            return {"ok": False, "error": f"kb_sync_http_{response.status_code}"}
        payload = response.json()
        return payload if isinstance(payload, dict) else {"ok": True}
    except (requests.RequestException, ValueError) as exc:
        return {"ok": False, "error": str(exc)}


def _trigger_hybrid_reindex() -> dict[str, Any]:
    try:
        response = requests.post(
            f"{KB_SYNC_URL}/reindex-all",
            headers={"X-API-Key": KB_SYNC_INTERNAL_API_KEY},
            timeout=300,
        )
        if response.status_code >= 400:
            return {"ok": False, "error": f"kb_sync_http_{response.status_code}"}
        payload = response.json()
        return payload if isinstance(payload, dict) else {"ok": True}
    except (requests.RequestException, ValueError) as exc:
        return {"ok": False, "error": str(exc)}


def _trigger_hybrid_version_sync(version_id: str) -> dict[str, Any]:
    """Publish one canonical version; full rebuilds are reserved for maintenance."""
    try:
        response = requests.post(
            f"{KB_SYNC_URL}/hybrid/versions/sync",
            json={"version_id": version_id},
            headers={"X-API-Key": KB_SYNC_INTERNAL_API_KEY},
            timeout=180,
        )
        if response.status_code >= 400:
            return {"ok": False, "error": f"kb_sync_http_{response.status_code}"}
        payload = response.json()
        return payload if isinstance(payload, dict) else {"ok": True}
    except (requests.RequestException, ValueError) as exc:
        return {"ok": False, "error": str(exc)}


def _trigger_hybrid_document_sync(document_id: str) -> dict[str, Any]:
    try:
        response = requests.post(
            f"{KB_SYNC_URL}/hybrid/documents/sync", json={"document_id": document_id},
            headers={"X-API-Key": KB_SYNC_INTERNAL_API_KEY}, timeout=180,
        )
        if response.status_code >= 400:
            return {"ok": False, "error": f"kb_sync_http_{response.status_code}"}
        payload = response.json()
        return payload if isinstance(payload, dict) else {"ok": True}
    except (requests.RequestException, ValueError) as exc:
        return {"ok": False, "error": str(exc)}


def _archived_version_files(document_id: str, version_id: str) -> tuple[Path, list[Path], Path]:
    """Return retained archive paths only when they remain under portal storage."""
    version_root = (PORTAL_FILES_ROOT / document_id / version_id).resolve()
    try:
        version_root.relative_to(PORTAL_FILES_ROOT.resolve())
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="archived_version_not_available") from exc
    originals = [item for item in version_root.glob("original.*") if item.is_file()]
    return version_root, originals, version_root / "rag.md"


def _document_summary(path: Path, collection: str, state: dict[str, Any]) -> dict[str, Any]:
    relative_path = path.relative_to(KB_ROOT / collection).as_posix()
    meta = _metadata(path)
    entry = _state_entry(state, collection, relative_path)
    return {
        "collection": collection,
        "path": relative_path,
        "name": path.name,
        "extension": path.suffix.lower().lstrip("."),
        "size": path.stat().st_size,
        "modified_at": datetime.fromtimestamp(path.stat().st_mtime).astimezone().isoformat(),
        "chunks": int(entry.get("chunks") or 0),
        "indexed_at": entry.get("updated_at") or "",
        "index_status": _index_status(path, entry, bool(meta["rag_index"])),
        **meta,
    }


def _safe_archive_id(raw: str) -> str:
    archive_id = str(raw or "").strip()
    if not re.fullmatch(r"[0-9]{8}-[0-9]{6}-[a-z0-9][a-z0-9_-]+", archive_id):
        raise HTTPException(status_code=400, detail="invalid_archive_id")
    return archive_id


def _trash_manifest(archive: Path) -> dict[str, Any]:
    manifest_path = archive / ".trash.json"
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if not isinstance(manifest, dict):
            manifest = {}
    except Exception:
        manifest = {}
    match = re.fullmatch(r"(?P<stamp>[0-9]{8}-[0-9]{6})-(?P<collection>[a-z0-9][a-z0-9_-]+)", archive.name)
    if not match:
        raise ValueError("invalid_archive_name")
    archived_at = datetime.strptime(match.group("stamp"), "%Y%m%d-%H%M%S").astimezone()
    purge_at = archived_at + timedelta(days=TRASH_RETENTION_DAYS)
    collection = str(manifest.get("collection") or match.group("collection"))
    files = sum(
        1
        for path in archive.rglob("*")
        if path.is_file()
        and not path.name.startswith(".")
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    return {
        "archive_id": archive.name,
        "collection": collection,
        "label": str(manifest.get("label") or _collection_label_from_root(archive, collection)),
        "files": int(manifest.get("files") or files),
        "archived_at": str(manifest.get("archived_at") or archived_at.isoformat()),
        "purge_at": str(manifest.get("purge_at") or purge_at.isoformat()),
        "retention_days": TRASH_RETENTION_DAYS,
    }


def _collection_label_from_root(root: Path, fallback: str) -> str:
    try:
        payload = json.loads((root / ".collection.json").read_text(encoding="utf-8"))
        label = str(payload.get("label") or "").strip()
        if label:
            return label
    except Exception:
        pass
    return COLLECTION_LABELS.get(fallback, fallback)


def _trash_collections() -> list[dict[str, Any]]:
    root = TRASH_ROOT / "collections"
    if not root.exists():
        return []
    items: list[dict[str, Any]] = []
    for archive in root.iterdir():
        if not archive.is_dir():
            continue
        try:
            items.append(_trash_manifest(archive))
        except ValueError:
            continue
    return sorted(items, key=lambda item: item["archived_at"], reverse=True)


def require_maintenance(request: Request) -> dict[str, Any]:
    authorization = request.headers.get("authorization", "")
    provided = authorization.removeprefix("Bearer ").strip()
    if not MAINTENANCE_API_KEY or not hmac.compare_digest(provided, MAINTENANCE_API_KEY):
        raise HTTPException(status_code=401, detail="maintenance_auth_required")
    return {"id": "system:trash-cleanup", "name": "Tägliche Papierkorb-Bereinigung"}


def _purge_expired_collections(*, dry_run: bool, actor: dict[str, Any]) -> dict[str, Any]:
    now = datetime.now().astimezone()
    expired = [
        item
        for item in _trash_collections()
        if datetime.fromisoformat(item["purge_at"]) <= now
    ]
    purged: list[str] = []
    if not dry_run:
        for item in expired:
            archive = TRASH_ROOT / "collections" / item["archive_id"]
            if archive.exists():
                shutil.rmtree(archive)
                purged.append(item["archive_id"])
                _audit(actor, "collection_purged", **item)
    return {
        "ok": True,
        "dry_run": dry_run,
        "retention_days": TRASH_RETENTION_DAYS,
        "eligible": [item["archive_id"] for item in expired],
        "purged": purged,
    }

@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "ok": True,
        "version": APP_VERSION,
        "kb_root": str(KB_ROOT),
        "collections": list(_collection_names()),
    }


@app.get("/session")
def session(admin: dict[str, Any] = Depends(require_admin)) -> dict[str, Any]:
    return {
        "id": admin.get("id"),
        "name": admin.get("name"),
        "email": admin.get("email"),
        "role": admin.get("role"),
    }



@app.post("/portal/internal/retrieval-scope")
def portal_internal_retrieval_scope(
    payload: PortalRetrievalScopeRequest,
    x_api_key: str = Header(default="", alias="X-API-Key"),
) -> dict[str, Any]:
    if not MAINTENANCE_API_KEY or not hmac.compare_digest(x_api_key, MAINTENANCE_API_KEY):
        raise HTTPException(status_code=401, detail="internal_api_key_required")
    try:
        identity = PORTAL_GOVERNANCE.identity(payload.user_id)
        if not identity.active:
            raise GovernanceError("portal_user_inactive")
        knowledgebase_ids = PORTAL_GOVERNANCE.allowed_knowledgebases(payload.user_id, "read")
    except GovernanceError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    if not knowledgebase_ids:
        raise HTTPException(status_code=403, detail="no_readable_knowledgebases")
    with PORTAL_GOVERNANCE.store.connect() as db:
        placeholders = ",".join("?" for _ in knowledgebase_ids)
        active_version_ids = [row["active_version_id"] for row in db.execute(
            f"""SELECT DISTINCT d.active_version_id FROM canonical_documents d
                 JOIN document_versions v ON v.version_id=d.active_version_id AND v.status='active'
                 JOIN document_publications p ON p.document_id=d.document_id AND p.status='active'
                 WHERE p.knowledgebase_id IN ({placeholders}) AND v.valid_from <= ? AND v.valid_until >= ?""",
            (*knowledgebase_ids, date.today().isoformat(), date.today().isoformat()),
        ).fetchall()]
    if not active_version_ids:
        raise HTTPException(status_code=403, detail="no_active_readable_versions")
    return {"user_id": identity.user_id, "knowledgebase_ids": knowledgebase_ids,
            "active_version_ids": active_version_ids}


@app.post("/portal/internal/retrieval-events")
def portal_internal_retrieval_event(
    payload: PortalRetrievalEventRequest,
    x_api_key: str = Header(default="", alias="X-API-Key"),
) -> dict[str, bool]:
    if not MAINTENANCE_API_KEY or not hmac.compare_digest(x_api_key, MAINTENANCE_API_KEY):
        raise HTTPException(status_code=401, detail="internal_api_key_required")
    QUALITY_DASHBOARD.record_retrieval(
        user_id=payload.user_id, query_hash=payload.query_hash, found=payload.found,
        source_count=payload.source_count, latency_ms=payload.latency_ms,
        error_code=payload.error_code,
    )
    if payload.error_code:
        QUALITY_CASES.system_incident("retrieval", {"error_type": payload.error_code})
    return {"ok": True}


@app.get("/portal/session")
def portal_session(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    return identity


@app.get("/portal/knowledgebases")
def portal_knowledgebases(
    access: str = "read",
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if access not in {"read", "upload"}:
        raise HTTPException(status_code=400, detail="invalid_access_kind")
    try:
        knowledgebases = PORTAL_GOVERNANCE.list_knowledgebases(identity["user_id"], access)
    except GovernanceError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    return {
        "access": access,
        "knowledgebases": serialize_governance(knowledgebases),
    }

@app.get("/portal/documents")
def portal_list_documents(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    readable = PORTAL_GOVERNANCE.allowed_knowledgebases(identity["user_id"], "read")
    with PORTAL_GOVERNANCE.store.connect() as db:
        rows = db.execute(
            """SELECT d.document_id, d.title, d.owner_user_id, d.active_version_id,
                              v.status, v.valid_until,
                              (SELECT first_case.target_knowledgebase_id
                               FROM document_cases first_case
                               WHERE first_case.document_id=d.document_id
                               ORDER BY first_case.created_at,first_case.case_id LIMIT 1)
                              AS primary_knowledgebase_id
                              ,(SELECT latest_case.case_id
                                FROM document_cases latest_case
                                WHERE latest_case.document_id=d.document_id
                                ORDER BY latest_case.created_at DESC,latest_case.case_id DESC LIMIT 1)
                               AS latest_case_id
               FROM canonical_documents d
               LEFT JOIN document_versions v ON v.version_id=d.active_version_id
               WHERE (d.owner_user_id=? OR EXISTS (
                   SELECT 1 FROM document_publications readable_publication
                   WHERE readable_publication.document_id=d.document_id
                     AND readable_publication.status!='inactive'
                     AND readable_publication.knowledgebase_id IN ({})))
                 AND d.document_id NOT IN (
                     SELECT document_id FROM document_trash
                     WHERE physically_deleted_at IS NULL)
                 -- Zurueckgezogene und abgelehnte Vorgaenge sind erledigt. Ohne
                 -- diesen Filter blieben sie als Entwuerfe im Bestand stehen und
                 -- sahen aus, als warteten sie noch auf etwas. Der Status haengt
                 -- an der Version, nicht an active_version_id: ein verworfener
                 -- Entwurf hat gar keine aktive Version.
                 AND EXISTS (
                     SELECT 1 FROM document_versions live
                     WHERE live.document_id = d.document_id
                       AND live.status NOT IN ('withdrawn', 'rejected'))
               ORDER BY d.updated_at DESC""".format(",".join("?" for _ in readable) or "NULL"),
            (identity["user_id"], *readable),
        ).fetchall()
        document_ids = [row["document_id"] for row in rows]
        publications = []
        if document_ids:
            placeholders = ",".join("?" for _ in document_ids)
            publications = db.execute(
                f"""SELECT p.document_id,p.knowledgebase_id,k.label,p.created_at
                       FROM document_publications p
                       JOIN knowledgebases k ON k.knowledgebase_id=p.knowledgebase_id
                       WHERE p.document_id IN ({placeholders}) AND p.status!='inactive'
                       ORDER BY p.created_at,k.label""",
                document_ids,
            ).fetchall()
    by_document: dict[str, list[dict[str, str]]] = {}
    for publication in publications:
        by_document.setdefault(publication["document_id"], []).append({
            "knowledgebase_id": publication["knowledgebase_id"],
            "label": publication["label"],
        })
    documents = []
    for row in rows:
        item = dict(row)
        publication_items = by_document.get(row["document_id"], [])
        primary_id = item.pop("primary_knowledgebase_id", None)
        primary = next(
            (base for base in publication_items if base["knowledgebase_id"] == primary_id),
            publication_items[0] if publication_items else None,
        )
        item["primary_knowledgebase"] = primary
        item["additional_knowledgebases"] = [
            base for base in publication_items if not primary
            or base["knowledgebase_id"] != primary["knowledgebase_id"]
        ]
        if item.get("active_version_id"):
            item["original_url"] = f"/wissen/api/portal/sources/{item['active_version_id']}"
        elif item.get("latest_case_id"):
            item["original_url"] = f"/wissen/api/portal/cases/{item['latest_case_id']}/original"
        else:
            item["original_url"] = None
        item.pop("latest_case_id", None)
        documents.append(item)
    return {"documents": documents}


@app.get("/portal/admin/knowledgebase-overview")
def portal_admin_knowledgebase_overview(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    """
    Verwaltungssicht auf Wissensbereiche und ihre Dokumente.

    Bewusst getrennt von /portal/documents: Dort filtern die Leserechte, hier
    verwaltet ein Admin die Struktur und braucht auch Bereiche, in die er selbst
    nicht hineinlesen darf. Zurueckgegeben werden ausschliesslich Metadaten,
    niemals Dokumentinhalte, und Owner erscheinen mit Anzeigenamen statt roher
    Benutzer-ID.
    """
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    with PORTAL_GOVERNANCE.store.connect() as db:
        bases = db.execute(
            # Geloeschte Bereiche verschwinden aus der Verwaltung; archivierte
            # bleiben sichtbar, weil nur sie endgueltig entfernt werden koennen.
            "SELECT knowledgebase_id, slug, label, purpose, status FROM knowledgebases"
            " WHERE status != 'deleted' ORDER BY label"
        ).fetchall()
        documents = db.execute(
            """SELECT p.knowledgebase_id, d.document_id, d.title, d.owner_user_id,
                      COALESCE(u.display_name, d.owner_user_id) AS owner_name,
                      COALESCE(v.status, 'draft') AS status, v.valid_until
               FROM document_publications p
               JOIN canonical_documents d ON d.document_id = p.document_id
               LEFT JOIN document_versions v ON v.version_id = d.active_version_id
               LEFT JOIN portal_users u ON u.user_id = d.owner_user_id
               WHERE p.status != 'inactive'
                 AND d.document_id NOT IN (
                   SELECT document_id FROM document_trash
                   WHERE physically_deleted_at IS NULL)
               ORDER BY d.title"""
        ).fetchall()

    by_base: dict[str, list[dict[str, Any]]] = {}
    for row in documents:
        by_base.setdefault(row["knowledgebase_id"], []).append({
            key: row[key] for key in
            ("document_id", "title", "owner_name", "status", "valid_until")
        })

    overview = []
    for base in bases:
        entries = by_base.get(base["knowledgebase_id"], [])
        counts: dict[str, int] = {}
        for entry in entries:
            counts[entry["status"]] = counts.get(entry["status"], 0) + 1
        overview.append({
            **{key: base[key] for key in ("knowledgebase_id", "slug", "label", "purpose", "status")},
            "document_count": len(entries),
            "status_counts": counts,
            "documents": entries,
        })
    unassigned = sum(1 for row in documents if row["knowledgebase_id"] is None)
    return {"knowledgebases": overview, "unpublished_documents": unassigned}


class PortalPublicationRequest(BaseModel):
    knowledgebase_id: str = Field(..., min_length=1, max_length=100)
    active: bool = True
    reason: str = Field(..., min_length=3, max_length=2000)


@app.get("/portal/admin/documents/{document_id}/publications")
def portal_admin_publications(
    document_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"publications": DOCUMENT_LIFECYCLE.publications_of(document_id)}


@app.put("/portal/admin/documents/{document_id}/publications")
def portal_admin_set_publication(
    document_id: str, payload: PortalPublicationRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    """
    Ordnet ein bestehendes Dokument einem Wissensbereich zu oder loest die
    Zuordnung (PRD 9.3). Ohne diesen Weg blieb ein Dokument nach dem Entfernen
    seines Bereichs ohne Zuordnung und liess sich keinem anderen zuweisen.
    """
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        DOCUMENT_LIFECYCLE.set_publication(
            document_id=document_id, knowledgebase_id=payload.knowledgebase_id,
            active=payload.active, actor_user_id=identity["user_id"], reason=payload.reason,
        )
    except LifecycleError as exc:
        status = 404 if str(exc).startswith(("document_not_found", "unknown_")) else 409
        raise HTTPException(status_code=status, detail=str(exc)) from exc
    PORTAL_GOVERNANCE.record_audit(
        identity["user_id"], "document_publication_changed", "document", document_id,
        {"knowledgebase_id": payload.knowledgebase_id, "active": payload.active},
    )
    indexing = _trigger_hybrid_document_sync(document_id)
    return {"publications": DOCUMENT_LIFECYCLE.publications_of(document_id), "indexing": indexing}


def _resolve_valid_workdays(valid_workdays: int | None, valid_until: str | None) -> int:
    """
    PRD 17.1 laesst zwei gleichwertige Eingaben zu: Arbeitstage oder ein
    geprueftes Datum. Genau eine davon muss gesetzt sein; das Datum wird
    serverseitig umgerechnet, damit die niedersaechsischen Feiertage und die
    Grenze von 60 Arbeitstagen verbindlich bleiben.
    """
    if (valid_workdays is None) == (valid_until is None):
        raise HTTPException(status_code=422, detail="valid_workdays_or_valid_until_required")
    if valid_workdays is not None:
        return valid_workdays
    try:
        target = date.fromisoformat(valid_until or "")
    except ValueError as exc:
        raise HTTPException(status_code=422, detail="invalid_valid_until") from exc
    try:
        return workdays_until(date.today(), target)
    except LifecycleError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc


def _upload_knowledgebase_ids(primary_id: str, encoded_ids: Any) -> tuple[str, ...]:
    if not isinstance(encoded_ids, str):
        encoded_ids = ""
    try:
        decoded = json.loads(encoded_ids) if encoded_ids else []
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=422, detail="invalid_knowledgebase_selection") from exc
    if not isinstance(decoded, list) or any(not isinstance(item, str) for item in decoded):
        raise HTTPException(status_code=422, detail="invalid_knowledgebase_selection")
    ids = tuple(dict.fromkeys(item.strip() for item in decoded if item.strip()))
    if not ids:
        ids = (primary_id.strip(),) if primary_id.strip() else ()
    if not ids:
        raise HTTPException(status_code=422, detail="knowledgebase_required")
    return ids


@app.post("/portal/documents", response_model=PortalUploadResponse, status_code=201)
async def portal_upload_document(
    file: UploadFile = File(...),
    knowledgebase_id: str = Form(...),
    knowledgebase_ids_json: str = Form(""),
    title: str = Form(..., min_length=2, max_length=300),
    valid_workdays: int | None = Form(None, ge=1, le=60),
    valid_until: str | None = Form(None),
    confidentiality: str = Form("internal"),
    owner_user_id: str | None = Form(None),
    security_review_requested: bool = Form(False),
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> PortalUploadResponse:
    valid_workdays = _resolve_valid_workdays(valid_workdays, valid_until)
    knowledgebase_ids = _upload_knowledgebase_ids(knowledgebase_id, knowledgebase_ids_json)
    knowledgebase_id = knowledgebase_ids[0]
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="file_too_large")
    filename = file.filename or ""
    intended_owner_user_id = owner_user_id or identity["user_id"]
    try:
        if intended_owner_user_id != identity["user_id"]:
            if not OWNERSHIP.may_propose_other(identity["user_id"]):
                raise GovernanceError("owner_proposal_permission_required")
            if not PORTAL_GOVERNANCE.identity(intended_owner_user_id).active:
                raise GovernanceError("owner_inactive")
        security_review_finding = ""
        try:
            inspected = SECURE_INGEST.inspector.inspect(filename, data)
        except IngestError as exc:
            if str(exc) != "embedded_executable_content_not_allowed" or not security_review_requested:
                raise
            data = SECURE_INGEST.inspector.sanitize_pdf_for_admin_review(filename, data)
            inspected = SECURE_INGEST.inspector.inspect(filename, data)
            security_review_finding = "embedded_executable_content_not_allowed"
        document_id = str(uuid.uuid4())
        submission = DOCUMENT_LIFECYCLE.submit(
            uploaded_by_user_id=identity["user_id"], owner_user_id=identity["user_id"],
            target_knowledgebase_id=knowledgebase_id, title=title, original_filename=filename,
            original_file_id=f"portal://documents/{document_id}", original_sha256=inspected.sha256,
            valid_workdays=valid_workdays, confidentiality=confidentiality, document_id=document_id,
            target_knowledgebase_ids=knowledgebase_ids,
        )
        result = SECURE_INGEST.ingest(
            submission.document_id, submission.version_id, filename, data, title,
            conversion_progress=UPLOAD_CONVERSION_PROGRESS.get(),
        )
        markdown = result.markdown_path.read_text(encoding="utf-8")
        confidentiality_suggestion = CONFIDENTIALITY_CLASSIFIER.classify(markdown)
        restricted_terms = RESTRICTED_TERMS.matches(markdown)
        submission = DOCUMENT_LIFECYCLE.apply_automatic_confidentiality(
            case_id=submission.case_id, level=confidentiality_suggestion.level,
            reason=confidentiality_suggestion.reason, signals=confidentiality_suggestion.signals,
        )
        global_result = GLOBAL_ANALYZER.analyze(
            version_id=submission.version_id, title=title,
            markdown=markdown,
        )
        material_matches = tuple(match for match in global_result.matches if match.level in {"identical", "very_high", "medium"})
        cross_kb_matches = tuple(
            match.document_id for match in material_matches
            if not set(knowledgebase_ids).intersection(match.knowledgebase_ids)
        )
        same_kb_levels = [
            match.level for match in material_matches
            if set(knowledgebase_ids).intersection(match.knowledgebase_ids)
        ]
        same_kb_similarity = next((level for level in ("very_high", "medium") if level in same_kb_levels), "none")
        submission = DOCUMENT_LIFECYCLE.record_analysis(
            case_id=submission.case_id, normalized_sha256=global_result.normalized_sha256,
            markdown_sha256=result.markdown_sha256,
            analysis=Analysis(
                exact_duplicate_document_id=global_result.exact_document_id,
                same_kb_similarity=same_kb_similarity, cross_kb_matches=cross_kb_matches,
                contradiction_document_ids=global_result.contradiction_document_ids,
                version_candidate_document_ids=tuple(
                    match.document_id for match in material_matches if match.version_candidate
                ),
                prompt_injection_risk=result.injection.risk,
                conversion_quality=result.conversion_quality,
                notes=result.conversion_issues,
                restricted_terms=restricted_terms,
            ),
        )
        if security_review_finding:
            submission = DOCUMENT_LIFECYCLE.route_sanitized_security_review(
                case_id=submission.case_id,
                actor_user_id=identity["user_id"],
                finding=security_review_finding,
            )
        RAG_METADATA.write(submission.version_id, result.markdown_path)
        GLOBAL_CORPUS.upsert(CorpusDocument(
            submission.document_id, submission.version_id, title,
            markdown, knowledgebase_ids,
            "pending" if submission.status in {"pending_employee_decision", "ready_to_activate"} else "quarantine",
        ))
        if intended_owner_user_id != identity["user_id"]:
            OWNERSHIP.create_initial_proposal(
                submission.document_id, submission.case_id, identity["user_id"], intended_owner_user_id,
            )
            submission = DOCUMENT_LIFECYCLE.submission(submission.case_id)
        elif submission.status == "ready_to_activate":
            previous_version_id = DOCUMENT_LIFECYCLE.active_version(submission.document_id)
            submission = DOCUMENT_LIFECYCLE.activate(case_id=submission.case_id, actor_user_id="auto_activation")
            RAG_METADATA.write(submission.version_id)
            indexing = _trigger_hybrid_version_sync(submission.version_id)
            if not indexing.get("ok"):
                submission = DOCUMENT_LIFECYCLE.rollback_activation(
                    case_id=submission.case_id, previous_version_id=previous_version_id,
                    reason=str(indexing.get("error") or "hybrid_reindex_failed"),
                    actor_user_id="auto_activation",
                )
                if previous_version_id:
                    _trigger_hybrid_version_sync(previous_version_id)
                _refresh_global_corpus_version(submission.version_id, "pending")
                if previous_version_id:
                    _refresh_global_corpus_version(previous_version_id, "active")
                raise HTTPException(status_code=503, detail="activation_index_failed_previous_version_restored")
            _refresh_global_corpus_version(submission.version_id, "active")
            if previous_version_id and previous_version_id != submission.version_id:
                _refresh_global_corpus_version(previous_version_id, "superseded")
            _notify_case_status(
                submission, "Automatische Sicherheits- und Qualitätsprüfung bestanden.",
            )
    except GovernanceError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    except LifecycleError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except IngestError as exc:
        code = str(exc)
        if code in {"malware_scanner_unavailable", "malware_scan_failed", "document_conversion_empty"}:
            diagnostic = _upload_job_diagnostic(code)
            incident_id = _notify_system_error(
                "required_ingest_check", diagnostic,
                fingerprint=_upload_job_fingerprint(code),
            )
            raise HTTPException(status_code=503, detail=f"required_check_unavailable:{incident_id}") from exc
        if code.startswith("document_conversion_failed"):
            page_range = code.partition(":pages=")[2] or None
            diagnostic = _upload_job_diagnostic(
                "document_conversion_failed", page_range=page_range,
            )
            incident_id = _notify_system_error(
                "required_ingest_check", diagnostic,
                fingerprint=_upload_job_fingerprint("document_conversion_failed"),
            )
            raise HTTPException(status_code=503, detail=f"document_conversion_unavailable:{incident_id}") from exc
        raise HTTPException(status_code=422, detail=code) from exc
    except GlobalAnalysisError as exc:
        incident_id = _notify_system_error("global_document_analysis", {
            "error_code": str(exc),
            "document_id": getattr(locals().get("submission"), "document_id", None),
            "case_id": getattr(locals().get("submission"), "case_id", None),
            "title": title,
            "original_filename": filename,
            "uploaded_by_user_id": identity["user_id"],
            "file_size_bytes": len(data),
        })
        raise HTTPException(status_code=503, detail=f"required_check_unavailable:{incident_id}") from exc
    except Exception as exc:
        incident_id = _notify_system_error("portal_upload", {"error_type": type(exc).__name__})
        print(f"portal_upload_failed incident={incident_id} error={exc}", flush=True)
        raise HTTPException(status_code=500, detail=f"system_error:{incident_id}") from exc
    return PortalUploadResponse(
        case_id=submission.case_id, document_id=submission.document_id,
        version_id=submission.version_id, status=submission.status,
        owner_email=PORTAL_GOVERNANCE.identity(intended_owner_user_id).email,
        owner_confirmation_required=intended_owner_user_id != identity["user_id"],
        prompt_injection_risk=result.injection.risk,
        conversion_quality=result.conversion_quality,
        conversion_issues=list(result.conversion_issues),
        restricted_terms=list(restricted_terms),
        confidentiality=DOCUMENT_LIFECYCLE.submission(submission.case_id).confidentiality,
        confidentiality_reason=(
            confidentiality_suggestion.reason
            + (f" Gesperrte Begriffe gefunden: {', '.join(restricted_terms)}."
               if restricted_terms else "")
        ),
        requires_admin=submission.requires_admin,
        exact_duplicate_document_id=global_result.exact_document_id,
        matches=[{
            "document_id": match.document_id, "title": match.title, "level": match.level,
            "knowledgebase_ids": list(match.knowledgebase_ids), "version_candidate": match.version_candidate,
            "has_conflict": bool(match.conflicting_passages),
            # Ohne diese Werte sieht der Entscheider nur eine Stufe wie "medium"
            # und kann nicht einschaetzen, wie stark der Treffer wirklich ist.
            "match_percent": round(match.combined_score * 100),
            "conflict_count": len(match.conflicting_passages),
        } for match in material_matches],
    )


def _run_portal_upload_job(
    job_id: str, data: bytes, filename: str, knowledgebase_ids: tuple[str, ...], title: str,
    valid_workdays: int, confidentiality: str, owner_user_id: str | None,
    identity: dict[str, Any], security_review_requested: bool = False,
) -> None:
    try:
        UPLOAD_JOBS.heartbeat(job_id, "security", 20)
        UPLOAD_JOBS.heartbeat(job_id, "conversion", 30)
        def conversion_progress(completed_chunks: int, total_chunks: int) -> None:
            # 30-75% is reserved for conversion. This makes long PDFs visibly
            # advance while every block is processed independently.
            progress = 30 + round(45 * completed_chunks / max(1, total_chunks))
            UPLOAD_JOBS.heartbeat(job_id, "conversion", progress)
        upload = UploadFile(file=io.BytesIO(data), filename=filename)
        progress_token = UPLOAD_CONVERSION_PROGRESS.set(conversion_progress)
        try:
            result = asyncio.run(portal_upload_document(
                file=upload, knowledgebase_id=knowledgebase_ids[0],
                knowledgebase_ids_json=json.dumps(knowledgebase_ids), title=title,
                # Die Gueltigkeit ist beim Anlegen des Jobs bereits in Arbeitstage
                # aufgeloest. valid_until muss trotzdem ausdruecklich None sein:
                # Beim direkten Funktionsaufruf greift sonst der Form(None)-Default,
                # und das ist ein FieldInfo-Objekt, nicht None.
                valid_workdays=valid_workdays, valid_until=None,
                confidentiality=confidentiality,
                owner_user_id=owner_user_id, identity=identity,
                security_review_requested=security_review_requested,
            ))
        finally:
            UPLOAD_CONVERSION_PROGRESS.reset(progress_token)
        UPLOAD_JOBS.heartbeat(job_id, "comparison", 90)
        UPLOAD_JOBS.complete(job_id, result.model_dump())
    except HTTPException as exc:
        incident_id = str(exc.detail).rsplit(":", 1)[-1] if ":" in str(exc.detail) else None
        if incident_id:
            UPLOAD_JOBS.set_incident(job_id, incident_id)
        UPLOAD_JOBS.fail(job_id, str(exc.detail))
        _notify_upload_failure(UPLOAD_JOB_CONTEXT.get(), incident_id, str(exc.detail).split(":", 1)[0])
    except Exception as exc:  # pragma: no cover - final fail-safe for worker failures
        diagnostic = _upload_job_diagnostic("system_error", error_type=type(exc).__name__)
        incident_id = _notify_system_error(
            "portal_upload_job", diagnostic, fingerprint=_upload_job_fingerprint("system_error"),
        )
        UPLOAD_JOBS.set_incident(job_id, incident_id)
        UPLOAD_JOBS.fail(job_id, f"system_error:{incident_id}")
        _notify_upload_failure(UPLOAD_JOB_CONTEXT.get(), incident_id, "system_error")


def drain_one_upload_job() -> bool:
    job, interrupted = UPLOAD_JOBS.recover_and_claim_next()
    _finalize_interrupted_upload_jobs(interrupted)
    if not job:
        return False
    context_token = UPLOAD_JOB_CONTEXT.set(job)
    try:
        account = PORTAL_GOVERNANCE.identity(job["user_id"])
        identity = {
            "user_id": account.user_id, "email": account.email,
            "display_name": account.display_name, "role": account.role,
        }
        _run_portal_upload_job(
            job["job_id"], UPLOAD_SPOOL.read(job["job_id"]), job["original_filename"],
            tuple(job["knowledgebase_ids"]), job["title"], int(job["valid_workdays"]),
            job["confidentiality"], job["owner_user_id"], identity,
            bool(job["security_review_requested"]),
        )
    except Exception as exc:
        diagnostic = _upload_job_diagnostic("system_error", error_type=type(exc).__name__)
        incident_id = _notify_system_error(
            "portal_upload_job", diagnostic, fingerprint=_upload_job_fingerprint("system_error"),
        )
        UPLOAD_JOBS.set_incident(job["job_id"], incident_id)
        UPLOAD_JOBS.fail(job["job_id"], f"system_error:{incident_id}")
        _notify_upload_failure(job, incident_id, "system_error")
    finally:
        terminal = UPLOAD_JOBS.get(job["job_id"], job["user_id"], is_admin=True)["status"]
        if terminal in {"completed", "failed"}:
            UPLOAD_SPOOL.remove(job["job_id"])
        UPLOAD_JOB_CONTEXT.reset(context_token)
    return True


def recover_interrupted_upload_jobs() -> list[dict[str, Any]]:
    expired = UPLOAD_JOBS.expire_interrupted()
    _finalize_interrupted_upload_jobs(expired)
    return expired


def _finalize_interrupted_upload_jobs(expired: list[dict[str, Any]]) -> None:
    for job in expired:
        context_token = UPLOAD_JOB_CONTEXT.set(job)
        try:
            try:
                incident_id = _notify_system_error(
                    "portal_upload_job", _upload_job_diagnostic("upload_worker_interrupted"),
                    fingerprint=_upload_job_fingerprint("upload_worker_interrupted"),
                )
                UPLOAD_JOBS.set_incident(job["job_id"], incident_id)
                _notify_upload_failure(job, incident_id, "upload_worker_interrupted")
            except Exception as exc:
                print(
                    f"upload_interruption_notification_failed job={job['job_id']} "
                    f"error={type(exc).__name__}", flush=True,
                )
            finally:
                UPLOAD_SPOOL.remove(job["job_id"])
        finally:
            UPLOAD_JOB_CONTEXT.reset(context_token)


@app.post("/portal/upload-jobs", status_code=202)
async def portal_create_upload_job(
    file: UploadFile = File(...), knowledgebase_id: str = Form(...),
    knowledgebase_ids_json: str = Form(""),
    title: str = Form(..., min_length=2, max_length=300),
    valid_workdays: int | None = Form(None, ge=1, le=60),
    valid_until: str | None = Form(None),
    confidentiality: str = Form("internal"),
    owner_user_id: str | None = Form(None),
    security_review_requested: bool = Form(False),
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    valid_workdays = _resolve_valid_workdays(valid_workdays, valid_until)
    knowledgebase_ids = _upload_knowledgebase_ids(knowledgebase_id, knowledgebase_ids_json)
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="file_too_large")
    try:
        for selected_id in knowledgebase_ids:
            PORTAL_GOVERNANCE.require_access(identity["user_id"], selected_id, "upload")
    except GovernanceError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    job_id = str(uuid.uuid4())
    intended_owner_user_id = owner_user_id or identity["user_id"]
    UPLOAD_SPOOL.stage(job_id, data)
    try:
        return UPLOAD_JOBS.enqueue(
            job_id=job_id, user_id=identity["user_id"], original_filename=file.filename or "",
            title=title, knowledgebase_ids=knowledgebase_ids, valid_workdays=valid_workdays,
            confidentiality=confidentiality, owner_user_id=intended_owner_user_id,
            security_review_requested=security_review_requested, staged_path=f"{job_id}.upload",
            file_size_bytes=len(data),
        )
    except Exception:
        UPLOAD_SPOOL.remove(job_id)
        raise


@app.get("/portal/upload-jobs")
def portal_active_upload_jobs(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    return {"jobs": UPLOAD_JOBS.list_active(
        identity["user_id"], identity["role"] in {"admin", "portal_admin"},
    )}


@app.get("/portal/upload-jobs/{job_id}")
def portal_get_upload_job(
    job_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        return UPLOAD_JOBS.get(
            job_id, identity["user_id"], identity["role"] in {"admin", "portal_admin"}
        )
    except UploadJobError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


def _feedback_reference_options(user_id: str) -> dict[str, Any]:
    """Liefert nur Quellen, die der aktuelle Nutzer tatsaechlich lesen darf."""
    readable = PORTAL_GOVERNANCE.allowed_knowledgebases(user_id, "read")
    with PORTAL_GOVERNANCE.store.connect() as db:
        if not readable:
            return {"knowledgebases": [], "documents": []}
        placeholders = ",".join("?" for _ in readable)
        bases = db.execute(
            f"SELECT knowledgebase_id,label FROM knowledgebases "
            f"WHERE status='active' AND knowledgebase_id IN ({placeholders}) ORDER BY label",
            readable,
        ).fetchall()
        rows = db.execute(
            f"""SELECT DISTINCT d.document_id,d.title,p.knowledgebase_id
                   FROM canonical_documents d
                   JOIN document_publications p ON p.document_id=d.document_id
                   JOIN document_versions v ON v.version_id=d.active_version_id
                  WHERE p.status='active' AND v.status='active'
                    AND p.knowledgebase_id IN ({placeholders})
                  ORDER BY d.title""",
            readable,
        ).fetchall()
    documents: dict[str, dict[str, Any]] = {}
    for row in rows:
        item = documents.setdefault(row["document_id"], {
            "document_id": row["document_id"], "title": row["title"], "knowledgebase_ids": [],
        })
        item["knowledgebase_ids"].append(row["knowledgebase_id"])
    return {
        "knowledgebases": [dict(row) for row in bases],
        "documents": list(documents.values()),
    }


def _feedback_source_document_ids(*collections: Any) -> list[str]:
    """Extrahiert Dokumentreferenzen aus den variierenden OpenWebUI-Quellenformaten."""
    found: list[str] = []

    def visit(value: Any) -> None:
        if isinstance(value, dict):
            document_id = value.get("document_id")
            if isinstance(document_id, str) and document_id and document_id not in found:
                found.append(document_id)
            for child in value.values():
                visit(child)
        elif isinstance(value, list):
            for child in value:
                visit(child)

    for collection in collections:
        visit(collection)
    return found[:30]


def _feedback_source_knowledgebase_ids(*collections: Any) -> list[str]:
    found: list[str] = []

    def visit(value: Any) -> None:
        if isinstance(value, dict):
            raw_ids = value.get("knowledgebase_ids") or []
            if isinstance(raw_ids, list):
                for knowledgebase_id in raw_ids:
                    if isinstance(knowledgebase_id, str) and knowledgebase_id not in found:
                        found.append(knowledgebase_id)
            for child in value.values():
                visit(child)
        elif isinstance(value, list):
            for child in value:
                visit(child)

    for collection in collections:
        visit(collection)
    return found[:30]


@app.get("/portal/feedback/options")
def portal_feedback_options(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    return _feedback_reference_options(identity["user_id"])


@app.get("/portal/feedback/context")
def portal_feedback_context(
    chat_id: str, message_id: str, request: Request,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if not re.fullmatch(r"[A-Za-z0-9_-]{1,100}", chat_id) or not re.fullmatch(r"[A-Za-z0-9_-]{1,100}", message_id):
        raise HTTPException(status_code=400, detail="invalid_feedback_reference")
    headers = {}
    if request.headers.get("Authorization"):
        headers["Authorization"] = request.headers["Authorization"]
    if request.headers.get("Cookie"):
        headers["Cookie"] = request.headers["Cookie"]
    try:
        response = requests.get(f"{OPENWEBUI_URL}/api/v1/chats/{chat_id}", headers=headers, timeout=20)
    except requests.RequestException as exc:
        raise HTTPException(status_code=503, detail="chat_context_unavailable") from exc
    if response.status_code != 200:
        raise HTTPException(status_code=404, detail="chat_context_not_available")
    chat = response.json().get("chat") or response.json()
    messages = chat.get("messages") or {}
    if isinstance(messages, list):
        messages = {str(item.get("id")): item for item in messages}
    answer = messages.get(message_id) or {}
    parent = messages.get(str(answer.get("parentId") or "")) or {}
    sources = (answer.get("sources") or [])[:30]
    passages = (answer.get("metadata") or {}).get("sources", [])[:30]
    options = _feedback_reference_options(identity["user_id"])
    allowed_documents = {item["document_id"]: item for item in options["documents"]}
    selected_document_ids = [
        document_id for document_id in _feedback_source_document_ids(sources, passages)
        if document_id in allowed_documents
    ]
    allowed_knowledgebase_ids = {item["knowledgebase_id"] for item in options["knowledgebases"]}
    selected_knowledgebase_ids = [
        value for value in _feedback_source_knowledgebase_ids(sources, passages)
        if value in allowed_knowledgebase_ids
    ]
    # Aeltere Chatantworten enthalten noch keine Bereichs-IDs. Fuer sie ist
    # die Dokumentzuordnung die bestmoegliche sichere Rueckfallebene.
    if not selected_knowledgebase_ids:
        for document_id in selected_document_ids:
            for knowledgebase_id in allowed_documents[document_id]["knowledgebase_ids"]:
                if knowledgebase_id not in selected_knowledgebase_ids:
                    selected_knowledgebase_ids.append(knowledgebase_id)
    return {
        "question": str(parent.get("content") or "")[:8000],
        "answer": str(answer.get("content") or "")[:16000],
        "sources": sources,
        "passages": passages,
        "runtime": {
            "model": answer.get("model"), "model_id": answer.get("modelId"),
            "prompt_version": "kahle-vinci-current", "retrieval_version": "hybrid-v2",
            "chat_id": chat_id, "message_id": message_id,
        },
        "request_id": str((answer.get("metadata") or {}).get("request_id") or message_id),
        "document_ids": selected_document_ids,
        "knowledgebase_ids": selected_knowledgebase_ids,
        **options,
    }


@app.post("/portal/feedback/rag", status_code=201)
def portal_report_rag_feedback(
    payload: PortalRagFeedbackRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    rights = PORTAL_GOVERNANCE.allowed_knowledgebases(identity["user_id"], "read")
    options = _feedback_reference_options(identity["user_id"])
    valid_documents = {item["document_id"] for item in options["documents"]}
    valid_knowledgebases = {item["knowledgebase_id"] for item in options["knowledgebases"]}
    document_ids = list(dict.fromkeys(payload.document_ids))
    knowledgebase_ids = list(dict.fromkeys(payload.knowledgebase_ids))
    if not set(document_ids).issubset(valid_documents) or not set(knowledgebase_ids).issubset(valid_knowledgebases):
        raise HTTPException(status_code=422, detail="invalid_feedback_reference_selection")
    try:
        feedback_id = QUALITY_CASES.report_rag(
            user_id=identity["user_id"], reason=payload.reason, comment=payload.comment,
            question=payload.question, answer=payload.answer, sources=payload.sources,
            passages=payload.passages, rights=rights, runtime=payload.runtime, request_id=payload.request_id,
            document_ids=document_ids, knowledgebase_ids=knowledgebase_ids,
        )
    except QualityCaseError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    recipients = _admin_emails() if payload.reason == "suspected_permission_issue" else []
    for recipient in recipients:
        MAINTENANCE.enqueue_notification(
            recipient, "critical_rag_feedback", "KAHLE-Vinci: M?glicher Berechtigungsfehler",
            f"Ein kritischer Wissensfehler wurde gemeldet. Referenz: {feedback_id}", dedupe_key=feedback_id,
        )
    return {"feedback_id": feedback_id, "status": "open"}


SCREENSHOT_INSPECTOR = ScreenshotInspector()
FEEDBACK_FILE_INSPECTOR = SecureFileInspector(max_bytes=5 * 1024 * 1024)


def _feedback_screenshot_dir(feedback_id: str) -> Path:
    safe = PurePosixPath(feedback_id).name
    if not safe or safe != feedback_id:
        raise HTTPException(status_code=422, detail="invalid_feedback_reference")
    return (PORTAL_FILES_ROOT / "feedback-screenshots" / safe).resolve()


@app.post("/portal/feedback/{feedback_id}/screenshot", status_code=201)
async def portal_attach_feedback_screenshot(
    feedback_id: str, file: UploadFile = File(...),
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    """
    Bis zu fünf sichere Datei- oder Bildanhänge zu einer Wissensfehlermeldung.

    Ein Screenshot zeigt dem Admin oft in Sekunden, was eine Beschreibung nur
    umstaendlich erklaert. Der Anhang durchlaeuft dieselbe Kette wie ein
    Dokument: Typpruefung am Inhalt statt an der Endung, Groessengrenze und
    Malwarescan. Nur der Meldende darf anhaengen, nur Admins duerfen abrufen.
    """
    data = await file.read(5 * 1024 * 1024 + 1)
    original_filename = PurePosixPath(file.filename or "").name
    try:
        suffix = Path(original_filename).suffix.lower()
        if suffix in {".png", ".jpg", ".jpeg"}:
            extension, media_type = SCREENSHOT_INSPECTOR.inspect(data)
        else:
            inspected = FEEDBACK_FILE_INSPECTOR.inspect(original_filename, data)
            extension, media_type = inspected.extension, inspected.media_type
    except IngestError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    try:
        SECURE_INGEST.scanner.scan(original_filename, data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="malware_scan_failed") from exc

    attachment_id = uuid.uuid4().hex
    filename = f"{attachment_id}.{extension}"
    target = _feedback_screenshot_dir(feedback_id)
    target.mkdir(parents=True, exist_ok=True)
    path = target / filename
    path.write_bytes(data)
    try:
        QUALITY_CASES.add_attachment(
            feedback_id, identity["user_id"], attachment_id=attachment_id,
            original_filename=original_filename, stored_filename=filename,
            media_type=media_type, size_bytes=len(data),
        )
    except QualityCaseError as exc:
        path.unlink(missing_ok=True)
        status = 404 if str(exc) == "feedback_not_found" else 403
        if str(exc) == "feedback_attachment_limit_reached":
            status = 422
        raise HTTPException(status_code=status, detail=str(exc)) from exc
    PORTAL_GOVERNANCE.record_audit(
        identity["user_id"], "feedback_attachment_added", "rag_feedback", feedback_id,
        {"filename": original_filename},
    )
    return {"feedback_id": feedback_id, "attachment_id": attachment_id, "filename": original_filename}


@app.get("/portal/admin/feedback/{feedback_id}/attachments/{attachment_id}")
def portal_admin_feedback_attachment(
    feedback_id: str, attachment_id: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        attachment = QUALITY_CASES.attachment(feedback_id, attachment_id)
    except QualityCaseError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    path = _feedback_screenshot_dir(feedback_id) / attachment["stored_filename"]
    if not path.is_file():
        raise HTTPException(status_code=404, detail="attachment_not_available")
    return FileResponse(
        path, media_type=attachment["media_type"], filename=attachment["original_filename"],
    )


@app.get("/portal/admin/feedback/{feedback_id}/screenshot")
def portal_admin_feedback_screenshot(
    feedback_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        filename = QUALITY_CASES.screenshot_of(feedback_id)
    except QualityCaseError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    if not filename:
        image = next((item for item in QUALITY_CASES.attachments_of(feedback_id)
                      if item["media_type"].startswith("image/")), None)
        if image:
            attachment = QUALITY_CASES.attachment(feedback_id, image["attachment_id"])
            filename = attachment["stored_filename"]
        else:
            raise HTTPException(status_code=404, detail="screenshot_not_available")
    path = _feedback_screenshot_dir(feedback_id) / filename
    if not path.is_file():
        raise HTTPException(status_code=404, detail="screenshot_not_available")
    media = "image/png" if filename.endswith(".png") else "image/jpeg"
    return FileResponse(path, media_type=media)


@app.post("/portal/incidents/{incident_id}/comment")
def portal_comment_incident(
    incident_id: str, payload: PortalIncidentCommentRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    del identity
    try:
        QUALITY_CASES.add_incident_comment(incident_id, payload.comment)
    except QualityCaseError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return {"ok": True}


@app.get("/portal/admin/quality-cases")
def portal_admin_quality_cases(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    cases = QUALITY_CASES.open_cases()
    with PORTAL_GOVERNANCE.store.connect() as db:
        base_labels = {row["knowledgebase_id"]: row["label"] for row in db.execute(
            "SELECT knowledgebase_id,label FROM knowledgebases"
        )}
        document_titles = {row["document_id"]: row["title"] for row in db.execute(
            "SELECT document_id,title FROM canonical_documents"
        )}
        user_names = {row["user_id"]: row["display_name"] for row in db.execute(
            "SELECT user_id,display_name FROM portal_users"
        )}
    for item in cases["feedback"]:
        document_ids = json.loads(item.pop("document_ids_json", "[]") or "[]")
        knowledgebase_ids = json.loads(item.pop("knowledgebase_ids_json", "[]") or "[]")
        item["documents"] = [
            {"document_id": value, "title": document_titles.get(value, value)} for value in document_ids
        ]
        item["knowledgebases"] = [
            {"knowledgebase_id": value, "label": base_labels.get(value, value)} for value in knowledgebase_ids
        ]
        item["reported_by_name"] = user_names.get(
            item["reported_by_user_id"], item["reported_by_user_id"]
        )
        item["attachments"] = QUALITY_CASES.attachments_of(item["feedback_id"])
    return cases


def _quality_feedback_recipient(feedback_id: str) -> tuple[Any, str]:
    try:
        reporter_id = QUALITY_CASES.feedback_reporter(feedback_id)
        reporter = PORTAL_GOVERNANCE.identity(reporter_id)
    except (QualityCaseError, GovernanceError) as exc:
        raise HTTPException(status_code=404, detail="feedback_not_found") from exc
    return reporter, reporter_id


def _notify_quality_reporter(feedback_id: str, message: str, *, status: str) -> str:
    reporter, reporter_id = _quality_feedback_recipient(feedback_id)
    notification_id = uuid.uuid4().hex
    stamp = datetime.now().astimezone().isoformat()
    with PORTAL_GOVERNANCE.store.connect() as db:
        db.execute(
            "INSERT INTO portal_notifications "
            "(notification_id,recipient_user_id,subject_type,subject_id,subject_title,status,message,reason,created_at) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (notification_id, reporter_id, "rag_feedback", feedback_id,
             "Deine Wissensfehlermeldung", status,
             "Ein Admin hat deine Wissensfehlermeldung bearbeitet.", message, stamp),
        )
    MAINTENANCE.enqueue_notification(
        reporter.email, f"quality_case_{status}",
        "KAHLE-Vinci: Rückmeldung zu deiner Wissensfehlermeldung",
        f"Ein Admin hat deine Meldung bearbeitet.\n\nRückmeldung:\n{message}\n\n"
        "Im Wissensportal öffnen: /wissen/?notifications=1",
        dedupe_key=notification_id,
    )
    return notification_id


@app.post("/portal/admin/quality-cases/feedback/{feedback_id}/message", status_code=201)
def portal_admin_message_quality_reporter(
    feedback_id: str, payload: PortalQualityCaseMessageRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    notification_id = _notify_quality_reporter(
        feedback_id, payload.message.strip(), status="quality_case_message",
    )
    PORTAL_GOVERNANCE.record_audit(
        identity["user_id"], "quality_case_message_sent", "rag_feedback", feedback_id, {},
    )
    return {"notification_id": notification_id, "status": "sent"}


@app.post("/portal/admin/quality-cases/{case_type}/{case_id}/resolve")
def portal_admin_resolve_quality_case(
    case_type: str, case_id: str, payload: PortalQualityCaseResolveRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        QUALITY_CASES.resolve(case_type, case_id, payload.resolution)
    except QualityCaseError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    if case_type == "feedback":
        _notify_quality_reporter(
            case_id, payload.resolution.strip(), status="quality_case_resolved",
        )
    PORTAL_GOVERNANCE.record_audit(
        identity["user_id"], "quality_case_resolved", case_type, case_id,
        {"resolution": payload.resolution.strip()},
    )
    return {"case_id": case_id, "status": "resolved"}


@app.get("/portal/sources/{version_id}")
def portal_source(
    version_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    try:
        record = DOCUMENT_LIFECYCLE.source_record(version_id, identity["user_id"])
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=404, detail="source_not_available") from exc
    version_root = (PORTAL_FILES_ROOT / record["document_id"] / version_id).resolve()
    try:
        version_root.relative_to(PORTAL_FILES_ROOT.resolve())
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="source_not_available") from exc
    originals = list(version_root.glob("original.*"))
    if len(originals) != 1 or not originals[0].is_file():
        raise HTTPException(status_code=404, detail="source_not_available")
    inline_extensions = {".pdf", ".txt", ".md"}
    disposition = "inline" if originals[0].suffix.lower() in inline_extensions else "attachment"
    return FileResponse(
        originals[0], filename=record["original_filename"],
        content_disposition_type=disposition,
        headers={"Cache-Control": "private, no-store", "X-Content-Type-Options": "nosniff"},
    )


@app.get("/portal/cases/{case_id}/review")
def portal_case_review(
    case_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        review = MARKDOWN_CORRECTION.review(case_id, identity["user_id"])
    except MarkdownCorrectionError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    with PORTAL_GOVERNANCE.store.connect() as db:
        knowledgebases = [dict(row) for row in db.execute(
            "SELECT p.knowledgebase_id,k.label FROM document_publications p "
            "JOIN knowledgebases k ON k.knowledgebase_id=p.knowledgebase_id "
            "WHERE p.document_id=? AND p.status!='inactive' ORDER BY k.label COLLATE NOCASE",
            (review["case"].document_id,),
        ).fetchall()]
    return {"case": asdict(review["case"]), "markdown": review["markdown"],
            "original_url": review["original_url"], "knowledgebases": knowledgebases}


@app.get("/portal/cases/{case_id}/original")
def portal_case_original(
    case_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    try:
        review = MARKDOWN_CORRECTION.review(case_id, identity["user_id"])
    except MarkdownCorrectionError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    case = review["case"]; originals = list((PORTAL_FILES_ROOT / case.document_id / case.version_id).glob("original.*"))
    if len(originals) != 1:
        raise HTTPException(status_code=404, detail="original_not_available")
    return FileResponse(originals[0], filename=case.original_filename, content_disposition_type="inline",
                        headers={"Cache-Control": "private, no-store", "X-Content-Type-Options": "nosniff"})


@app.post("/portal/cases/{case_id}/revision", status_code=201)
def portal_case_revision(
    case_id: str, payload: PortalMarkdownRevisionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        revised = MARKDOWN_CORRECTION.revise(
            case_id, identity["user_id"], instruction=payload.instruction,
            replacement_markdown=payload.replacement_markdown, reason=payload.reason,
            confirmed=payload.confirmed,
        )
    except (MarkdownCorrectionError, LifecycleError, requests.RequestException) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    RAG_METADATA.write(revised.version_id)
    return {"case": asdict(revised)}


@app.get("/portal/tasks")
def portal_tasks(identity: dict[str, Any] = Depends(require_portal_identity)) -> dict[str, Any]:
    try:
        tasks = DOCUMENT_LIFECYCLE.tasks_for(identity["user_id"])
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    payload = []
    with PORTAL_GOVERNANCE.store.connect() as db:
        knowledgebase_labels = {
            row["knowledgebase_id"]: row["label"] for row in db.execute(
                "SELECT knowledgebase_id,label FROM knowledgebases"
            ).fetchall()
        }
        review_knowledgebases = [dict(row) for row in db.execute(
            "SELECT knowledgebase_id,label FROM knowledgebases WHERE status='active' "
            "ORDER BY label COLLATE NOCASE"
        ).fetchall()]
        for task in tasks:
            item = asdict(task)
            row = db.execute(
                "SELECT analysis_json,created_at FROM document_cases WHERE case_id=?",
                (task.case_id,),
            ).fetchone()
            analysis = json.loads(row["analysis_json"] or "{}") if row else {}
            item["restricted_terms"] = analysis.get("restricted_terms", [])
            item["security_review_finding"] = analysis.get("security_review_finding", "")
            item["sanitized_for_review"] = bool(analysis.get("sanitized_for_review"))
            duplicate_ids = list(dict.fromkeys(filter(None, [
                analysis.get("exact_duplicate_document_id"),
                analysis.get("normalized_duplicate_document_id"),
                *(analysis.get("version_candidate_document_ids") or []),
            ])))
            duplicate_matches = []
            for duplicate_id in duplicate_ids:
                duplicate = db.execute(
                    "SELECT document_id,title,active_version_id FROM canonical_documents "
                    "WHERE document_id=?", (duplicate_id,),
                ).fetchone()
                if not duplicate or not duplicate["active_version_id"]:
                    continue
                relation = (
                    "exact_duplicate"
                    if duplicate_id in {
                        analysis.get("exact_duplicate_document_id"),
                        analysis.get("normalized_duplicate_document_id"),
                    }
                    else "version_candidate"
                )
                duplicate_matches.append({
                    "document_id": duplicate["document_id"],
                    "title": duplicate["title"],
                    "active_version_id": duplicate["active_version_id"],
                    "relation": relation,
                    "has_conflict": duplicate_id in (analysis.get("contradiction_document_ids") or []),
                    "original_url": (
                        f"/wissen/api/portal/cases/{task.case_id}/comparison/"
                        f"{duplicate['document_id']}/original"
                    ),
                })
            item["duplicate_matches"] = duplicate_matches
            rollback = db.execute(
                "SELECT details_json FROM document_events WHERE case_id=? "
                "AND event_type IN ('activation_rolled_back','existing_publication_rolled_back') "
                "ORDER BY sequence DESC LIMIT 1",
                (task.case_id,),
            ).fetchone()
            rollback_details = json.loads(rollback["details_json"] or "{}") if rollback else {}
            item["publication_error"] = (
                str(rollback_details.get("reason") or "")
                if task.status == "ready_to_activate" else ""
            )
            uploader = db.execute(
                "SELECT display_name FROM portal_users WHERE user_id=?",
                (task.uploaded_by_user_id,),
            ).fetchone()
            item["contact_name"] = (
                uploader["display_name"] if uploader else task.uploaded_by_user_id
            )
            item["target_knowledgebase_label"] = next(
                (base["label"] for base in review_knowledgebases
                 if base["knowledgebase_id"] == task.target_knowledgebase_id),
                task.target_knowledgebase_id,
            )
            selected_ids = [
                publication["knowledgebase_id"] for publication in db.execute(
                    "SELECT knowledgebase_id FROM document_publications "
                    "WHERE document_id=? AND status='pending' ORDER BY created_at,knowledgebase_id",
                    (task.document_id,),
                ).fetchall()
            ] or [task.target_knowledgebase_id]
            item["target_knowledgebase_ids"] = selected_ids
            item["target_knowledgebase_labels"] = [
                knowledgebase_labels.get(base_id, base_id) for base_id in selected_ids
            ]
            changes = []
            for change in db.execute(
                "SELECT e.actor_user_id,e.details_json,e.created_at,"
                "COALESCE(u.display_name,e.actor_user_id) AS actor_name,u.role AS actor_role "
                "FROM document_events e LEFT JOIN portal_users u ON u.user_id=e.actor_user_id "
                "WHERE e.case_id=? AND e.event_type IN "
                "('target_knowledgebase_changed','target_knowledgebases_changed') "
                "ORDER BY e.sequence",
                (task.case_id,),
            ).fetchall():
                details = json.loads(change["details_json"] or "{}")
                changed_ids = details.get("knowledgebase_ids") or [details.get("knowledgebase_id", "")]
                previous_ids = details.get("previous_knowledgebase_ids") or [
                    details.get("previous_knowledgebase_id", "")
                ]
                changes.append({
                    "knowledgebase_ids": [item for item in changed_ids if item],
                    "knowledgebase_id": next((item for item in changed_ids if item), ""),
                    "selected_by_user_id": change["actor_user_id"],
                    "selected_by_name": change["actor_name"],
                    "selected_by_role": change["actor_role"] or "",
                    "selected_at": change["created_at"],
                    "previous_knowledgebase_ids": [item for item in previous_ids if item],
                })
            submitted_event = db.execute(
                "SELECT details_json FROM document_events WHERE case_id=? AND event_type='submitted' "
                "ORDER BY sequence LIMIT 1", (task.case_id,),
            ).fetchone()
            submitted_details = json.loads(submitted_event["details_json"] or "{}") if submitted_event else {}
            initial_knowledgebase_ids = submitted_details.get("knowledgebase_ids") or (
                changes[0]["previous_knowledgebase_ids"] if changes else selected_ids
            )
            uploader_identity = db.execute(
                "SELECT display_name,role FROM portal_users WHERE user_id=?",
                (task.uploaded_by_user_id,),
            ).fetchone()
            history = [{
                "knowledgebase_ids": initial_knowledgebase_ids,
                "knowledgebase_labels": [
                    knowledgebase_labels.get(base_id, base_id)
                    for base_id in initial_knowledgebase_ids
                ],
                "knowledgebase_id": initial_knowledgebase_ids[0],
                "knowledgebase_label": knowledgebase_labels.get(
                    initial_knowledgebase_ids[0], initial_knowledgebase_ids[0],
                ),
                "selected_by_user_id": task.uploaded_by_user_id,
                "selected_by_name": (
                    uploader_identity["display_name"] if uploader_identity
                    else task.uploaded_by_user_id
                ),
                "selected_by_role": "employee",
                "selected_at": row["created_at"] if row else "",
            }]
            history.extend({
                **change,
                "knowledgebase_labels": [
                    knowledgebase_labels.get(base_id, base_id)
                    for base_id in change["knowledgebase_ids"]
                ],
                "knowledgebase_label": knowledgebase_labels.get(
                    change["knowledgebase_id"], change["knowledgebase_id"],
                ),
            } for change in changes)
            item["target_knowledgebase_history"] = history
            payload.append(item)
    return {"tasks": payload, "knowledgebases": review_knowledgebases}


@app.get("/portal/cases/{case_id}/comparison/{document_id}/original")
def portal_case_comparison_original(
    case_id: str, document_id: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    case = DOCUMENT_LIFECYCLE.submission(case_id)
    may_review = identity["user_id"] == case.uploaded_by_user_id or any(
        item.case_id == case_id for item in DOCUMENT_LIFECYCLE.tasks_for(identity["user_id"])
    )
    if not may_review:
        raise HTTPException(status_code=403, detail="reviewer_required")
    with PORTAL_GOVERNANCE.store.connect() as db:
        row = db.execute(
            "SELECT analysis_json FROM document_cases WHERE case_id=?", (case_id,),
        ).fetchone()
        analysis = json.loads(row["analysis_json"] or "{}") if row else {}
        allowed_ids = {
            analysis.get("exact_duplicate_document_id"),
            analysis.get("normalized_duplicate_document_id"),
            *(analysis.get("version_candidate_document_ids") or []),
        }
        if document_id not in allowed_ids:
            raise HTTPException(status_code=404, detail="comparison_not_available")
        version = db.execute(
            "SELECT active_version_id FROM canonical_documents WHERE document_id=?",
            (document_id,),
        ).fetchone()
        if not version or not version["active_version_id"]:
            raise HTTPException(status_code=404, detail="comparison_not_available")
        record = db.execute(
            "SELECT original_filename FROM document_versions WHERE version_id=? AND status='active'",
            (version["active_version_id"],),
        ).fetchone()
    if not record:
        raise HTTPException(status_code=404, detail="comparison_not_available")
    root = (PORTAL_FILES_ROOT / document_id / version["active_version_id"]).resolve()
    try:
        root.relative_to(PORTAL_FILES_ROOT.resolve())
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="comparison_not_available") from exc
    originals = list(root.glob("original.*"))
    if len(originals) != 1 or not originals[0].is_file():
        raise HTTPException(status_code=404, detail="comparison_not_available")
    return FileResponse(
        originals[0], filename=record["original_filename"],
        content_disposition_type=(
            "inline" if originals[0].suffix.lower() in {".pdf", ".txt", ".md"} else "attachment"
        ),
        headers={"Cache-Control": "private, no-store", "X-Content-Type-Options": "nosniff"},
    )


@app.patch("/portal/cases/{case_id}/target-knowledgebase")
def portal_case_target_knowledgebase(
    case_id: str, payload: PortalCaseTargetRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        case = DOCUMENT_LIFECYCLE.change_target_knowledgebase(
            case_id=case_id, actor_user_id=identity["user_id"],
            knowledgebase_id=payload.knowledgebase_id,
        )
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"case": asdict(case)}


@app.patch("/portal/cases/{case_id}/target-knowledgebases")
def portal_case_target_knowledgebases(
    case_id: str, payload: PortalCaseTargetsRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        case = DOCUMENT_LIFECYCLE.change_target_knowledgebases(
            case_id=case_id, actor_user_id=identity["user_id"],
            knowledgebase_ids=tuple(payload.knowledgebase_ids),
        )
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"case": asdict(case)}


def _case_inquiry_participants(case_id: str, actor_user_id: str) -> tuple[Any, list[dict[str, str]]]:
    try:
        visible_case_ids = {
            item.case_id for item in DOCUMENT_LIFECYCLE.tasks_for(actor_user_id)
        }
        case = DOCUMENT_LIFECYCLE.submission(case_id)
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=403, detail="case_not_available") from exc
    if case_id not in visible_case_ids:
        raise HTTPException(status_code=403, detail="case_not_available")
    participant_ids = [case.uploaded_by_user_id]
    if case.manager_user_id:
        participant_ids.append(case.manager_user_id)
    participant_ids = [item for item in dict.fromkeys(participant_ids) if item != actor_user_id]
    if not participant_ids:
        return case, []
    placeholders = ",".join("?" for _ in participant_ids)
    with PORTAL_GOVERNANCE.store.connect() as db:
        rows = db.execute(
            f"SELECT user_id,display_name,email FROM portal_users "
            f"WHERE active=1 AND user_id IN ({placeholders}) ORDER BY display_name COLLATE NOCASE",
            participant_ids,
        ).fetchall()
    return case, [dict(row) for row in rows]


@app.get("/portal/cases/{case_id}/inquiry-participants")
def portal_case_inquiry_participants(
    case_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    case, participants = _case_inquiry_participants(case_id, identity["user_id"])
    return {"case_id": case.case_id, "participants": participants}


@app.post("/portal/cases/{case_id}/inquiries", status_code=201)
def portal_case_inquiry(
    case_id: str, payload: PortalCaseInquiryRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    case, participants = _case_inquiry_participants(case_id, identity["user_id"])
    recipient = next(
        (item for item in participants if item["user_id"] == payload.recipient_user_id),
        None,
    )
    if not recipient:
        raise HTTPException(status_code=422, detail="inquiry_recipient_not_involved")
    sender = PORTAL_GOVERNANCE.identity(identity["user_id"])
    question = payload.question.strip()
    notification_id = uuid.uuid4().hex
    message = f"{sender.display_name} hat eine Rückfrage zu diesem Dokument."
    with PORTAL_GOVERNANCE.store.connect() as db:
        db.execute(
            "INSERT INTO portal_case_notifications "
            "(notification_id,recipient_user_id,case_id,status,message,reason,created_at,"
            "sender_user_id,thread_id) VALUES (?,?,?,?,?,?,?,?,?)",
            (notification_id, recipient["user_id"], case_id, "clarification_requested",
             message, question, datetime.now().astimezone().isoformat(),
             sender.user_id, notification_id),
        )
    MAINTENANCE.enqueue_notification(
        recipient["email"], "case_inquiry",
        f"KAHLE-Vinci: Rückfrage zu {case.title}",
        f"{message}\n\nRückfrage:\n{question}\n\n"
        f"Im Wissensportal öffnen: /wissen/?notifications=1",
        dedupe_key=notification_id,
    )
    return {"notification_id": notification_id, "status": "sent"}


@app.post("/portal/notifications/{notification_id}/reply", status_code=201)
def portal_notification_reply(
    notification_id: str, payload: PortalNotificationReplyRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        original = db.execute(
            "SELECT n.case_id,n.sender_user_id,n.thread_id,d.title AS document_title,"
            "sender.email AS sender_email "
            "FROM portal_case_notifications n "
            "JOIN document_cases c ON c.case_id=n.case_id "
            "JOIN canonical_documents d ON d.document_id=c.document_id "
            "LEFT JOIN portal_users sender ON sender.user_id=n.sender_user_id AND sender.active=1 "
            "WHERE n.notification_id=? AND n.recipient_user_id=?",
            (notification_id, identity["user_id"]),
        ).fetchone()
        if not original:
            raise HTTPException(status_code=404, detail="notification_not_found")
        if not original["sender_user_id"] or not original["sender_email"]:
            raise HTTPException(status_code=422, detail="notification_reply_not_available")
        sender = PORTAL_GOVERNANCE.identity(identity["user_id"])
        reply_id = uuid.uuid4().hex
        answer = payload.message.strip()
        message = f"{sender.display_name} hat auf deine Rückfrage geantwortet."
        db.execute(
            "INSERT INTO portal_case_notifications "
            "(notification_id,recipient_user_id,case_id,status,message,reason,created_at,"
            "sender_user_id,thread_id) VALUES (?,?,?,?,?,?,?,?,?)",
            (reply_id, original["sender_user_id"], original["case_id"],
             "clarification_reply", message, answer,
             datetime.now().astimezone().isoformat(), sender.user_id,
             original["thread_id"] or notification_id),
        )
    MAINTENANCE.enqueue_notification(
        original["sender_email"], "case_inquiry_reply",
        f"KAHLE-Vinci: Antwort zu {original['document_title']}",
        f"{message}\n\nAntwort:\n{answer}\n\n"
        f"Im Wissensportal öffnen: /wissen/?notifications=1",
        dedupe_key=reply_id,
    )
    return {"notification_id": reply_id, "status": "sent"}


@app.get("/portal/notifications/{notification_id}/thread")
def portal_notification_thread(
    notification_id: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        selected = db.execute(
            "SELECT COALESCE(thread_id,notification_id) AS thread_id "
            "FROM portal_case_notifications WHERE notification_id=?",
            (notification_id,),
        ).fetchone()
        if not selected:
            raise HTTPException(status_code=404, detail="notification_not_found")
        allowed = db.execute(
            "SELECT 1 FROM portal_case_notifications "
            "WHERE COALESCE(thread_id,notification_id)=? "
            "AND (recipient_user_id=? OR sender_user_id=?) LIMIT 1",
            (selected["thread_id"], identity["user_id"], identity["user_id"]),
        ).fetchone()
        if not allowed:
            raise HTTPException(status_code=404, detail="notification_not_found")
        rows = db.execute(
            "SELECT n.notification_id,n.status,n.reason AS message,n.created_at,"
            "n.sender_user_id,sender.display_name AS sender_name,"
            "n.recipient_user_id,recipient.display_name AS recipient_name "
            "FROM portal_case_notifications n "
            "LEFT JOIN portal_users sender ON sender.user_id=n.sender_user_id "
            "LEFT JOIN portal_users recipient ON recipient.user_id=n.recipient_user_id "
            "WHERE COALESCE(n.thread_id,n.notification_id)=? "
            "ORDER BY n.created_at,n.notification_id",
            (selected["thread_id"],),
        ).fetchall()
    return {"thread_id": selected["thread_id"], "messages": [dict(row) for row in rows]}


@app.get("/portal/notifications")
def portal_notifications(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        rows = db.execute(
            "SELECT * FROM ("
            "SELECT n.notification_id,n.case_id,n.status,n.message,n.reason,n.created_at,n.read_at,"
            "d.title AS document_title,n.sender_user_id,n.thread_id FROM portal_case_notifications n "
            "JOIN document_cases c ON c.case_id=n.case_id "
            "JOIN canonical_documents d ON d.document_id=c.document_id "
            "WHERE n.recipient_user_id=? UNION ALL "
            "SELECT n.notification_id,NULL AS case_id,n.status,n.message,n.reason,n.created_at,n.read_at,"
            "n.subject_title AS document_title,NULL AS sender_user_id,NULL AS thread_id FROM portal_notifications n "
            "WHERE n.recipient_user_id=?"
            ") ORDER BY created_at DESC,notification_id DESC LIMIT 100",
            (identity["user_id"], identity["user_id"]),
        ).fetchall()
    notifications = []
    for row in rows:
        item = dict(row)
        item["can_reply"] = bool(item["case_id"] and item["sender_user_id"])
        notifications.append(item)
    return {"notifications": notifications}


@app.post("/portal/notifications/{notification_id}/read")
def portal_mark_notification_read(
    notification_id: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, int]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        cursor = db.execute(
            "UPDATE portal_case_notifications SET read_at=? "
            "WHERE notification_id=? AND recipient_user_id=? AND read_at IS NULL",
            (datetime.now().astimezone().isoformat(), notification_id, identity["user_id"]),
        )
        exists = db.execute(
            "SELECT 1 FROM portal_case_notifications "
            "WHERE notification_id=? AND recipient_user_id=?",
            (notification_id, identity["user_id"]),
        ).fetchone()
        if not exists:
            cursor = db.execute(
                "UPDATE portal_notifications SET read_at=? "
                "WHERE notification_id=? AND recipient_user_id=? AND read_at IS NULL",
                (datetime.now().astimezone().isoformat(), notification_id, identity["user_id"]),
            )
            exists = db.execute(
                "SELECT 1 FROM portal_notifications "
                "WHERE notification_id=? AND recipient_user_id=?",
                (notification_id, identity["user_id"]),
            ).fetchone()
    if not exists:
        raise HTTPException(status_code=404, detail="notification_not_found")
    return {"marked_read": cursor.rowcount}


def _admin_emails() -> list[str]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        return [row["email"] for row in db.execute(
            "SELECT email FROM portal_users WHERE active = 1 AND role IN ('admin','portal_admin')"
        ).fetchall()]


def _knowledgebase_readers(knowledgebase_ids: tuple[str, ...]) -> list[dict[str, str]]:
    if not knowledgebase_ids:
        return []
    placeholders = ",".join("?" for _ in knowledgebase_ids)
    with PORTAL_GOVERNANCE.store.connect() as db:
        rows = db.execute(
            f"""SELECT DISTINCT u.user_id,u.email FROM portal_users u
                LEFT JOIN knowledgebase_access a ON a.user_id=u.user_id
                WHERE u.active=1 AND (
                  u.role IN ('admin','portal_admin') OR
                  (a.can_read=1 AND a.knowledgebase_id IN ({placeholders}))
                )""",
            knowledgebase_ids,
        ).fetchall()
    return [dict(row) for row in rows]


def _document_access_snapshot(document_id: str) -> tuple[str, list[dict[str, str]]]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        document = db.execute(
            "SELECT title FROM canonical_documents WHERE document_id=?", (document_id,),
        ).fetchone()
        knowledgebase_ids = tuple(row["knowledgebase_id"] for row in db.execute(
            "SELECT knowledgebase_id FROM document_publications "
            "WHERE document_id=? AND status='active'", (document_id,),
        ).fetchall())
    return (document["title"] if document else "Dokument", _knowledgebase_readers(knowledgebase_ids))


def _knowledgebase_access_snapshot(knowledgebase_id: str) -> tuple[str, list[dict[str, str]]]:
    knowledgebase = PORTAL_GOVERNANCE.knowledgebase(knowledgebase_id)
    return knowledgebase.label, _knowledgebase_readers((knowledgebase_id,))


def _notify_access_removed(
    recipients: list[dict[str, str]], *, subject_type: str, subject_id: str,
    subject_title: str, status: str, message: str, reason: str,
) -> None:
    stamp = datetime.now().astimezone().isoformat()
    clean_reason = reason.strip()
    queued_emails: list[tuple[str, str]] = []
    with PORTAL_GOVERNANCE.store.connect() as db:
        for recipient in recipients:
            notification_id = hashlib.sha256(
                f"{subject_type}|{subject_id}|{status}|{recipient['user_id']}|{clean_reason}".encode()
            ).hexdigest()[:32]
            db.execute(
                "INSERT OR IGNORE INTO portal_notifications "
                "(notification_id,recipient_user_id,subject_type,subject_id,subject_title,status,message,reason,created_at) "
                "VALUES (?,?,?,?,?,?,?,?,?)",
                (notification_id, recipient["user_id"], subject_type, subject_id, subject_title,
                 status, message, clean_reason, stamp),
            )
            queued_emails.append((recipient["email"], notification_id))
    body = f"{message}\n"
    if clean_reason:
        body += f"Begründung: {clean_reason}\n"
    body += "/wissen/?notifications=1"
    for recipient_email, notification_id in queued_emails:
        MAINTENANCE.enqueue_notification(
            recipient_email, "access_removed", f"KAHLE-Vinci: {subject_title}", body,
            dedupe_key=notification_id,
        )


def _notify_case_status(
    case: Any, reason: str = "", *, send_publication_email: bool = False,
) -> None:
    labels = {
        "active": "Das Dokument wurde veröffentlicht und in Vinci abrufbar gemacht. Der Zugriff bleibt auf berechtigte Nutzer beschränkt.",
        "pending_manager_approval": "Das Dokument wurde noch nicht veröffentlicht. Die Führungskraft muss entscheiden.",
        "pending_admin_approval": "Das Dokument wurde noch nicht veröffentlicht. Eine zusätzliche Adminprüfung ist erforderlich.",
        "rejected": "Das Dokument wurde abgelehnt und ist nicht in Vinci abrufbar.",
        "withdrawn": "Der Upload wurde verworfen und ist nicht in Vinci abrufbar.",
        "needs_correction": "Das Dokument wurde nicht veröffentlicht und muss korrigiert werden.",
    }
    message = labels.get(case.status)
    if not message:
        return
    recipient_ids = [case.uploaded_by_user_id]
    if case.manager_user_id and case.status in {
        "pending_manager_approval", "pending_admin_approval", "active", "rejected",
    }:
        recipient_ids.append(case.manager_user_id)
    if case.status == "pending_admin_approval":
        with PORTAL_GOVERNANCE.store.connect() as db:
            recipient_ids.extend(row["user_id"] for row in db.execute(
                "SELECT user_id FROM portal_users WHERE active=1 AND role IN ('admin','portal_admin')"
            ).fetchall())
    clean_reason = reason.strip()
    with PORTAL_GOVERNANCE.store.connect() as db:
        for recipient_user_id in dict.fromkeys(recipient_ids):
            notification_id = hashlib.sha256(
                f"{case.case_id}|{case.status}|{recipient_user_id}|{clean_reason}".encode()
            ).hexdigest()[:32]
            db.execute(
                "INSERT OR IGNORE INTO portal_case_notifications "
                "(notification_id,recipient_user_id,case_id,status,message,reason,created_at) "
                "VALUES (?,?,?,?,?,?,?)",
                (notification_id, recipient_user_id, case.case_id, case.status, message,
                 clean_reason, datetime.now().astimezone().isoformat()),
            )
    # Die übrigen Status sind im Portal direkt sichtbar. Eine E-Mail wird nur
    # nach einer tatsächlichen Freigabe und erfolgreicher Indexierung versendet.
    if not (case.status == "active" and send_publication_email):
        return
    uploader = PORTAL_GOVERNANCE.identity(case.uploaded_by_user_id)
    if not uploader.active:
        return
    MAINTENANCE.enqueue_notification(
        uploader.email,
        "approved_document_published",
        f"KAHLE-Vinci: {case.title} · Veröffentlicht",
        "Dein Dokument wurde freigegeben, erfolgreich indexiert und ist jetzt in Vinci abrufbar.\n"
        f"/wissen/?document={case.document_id}",
        dedupe_key=f"{case.case_id}:approved_document_published:{uploader.user_id}",
    )


def _refresh_global_corpus_version(version_id: str, status: str) -> None:
    with PORTAL_GOVERNANCE.store.connect() as db:
        row = db.execute(
            """SELECT v.version_id,v.document_id,d.title FROM document_versions v
               JOIN canonical_documents d ON d.document_id=v.document_id WHERE v.version_id=?""",
            (version_id,),
        ).fetchone()
        if not row:
            return
        knowledgebase_ids = tuple(item["knowledgebase_id"] for item in db.execute(
            "SELECT knowledgebase_id FROM document_publications WHERE document_id=? AND status!='inactive'",
            (row["document_id"],),
        ).fetchall())
    markdown_path = PORTAL_FILES_ROOT / row["document_id"] / version_id / "rag.md"
    if markdown_path.exists():
        GLOBAL_CORPUS.upsert(CorpusDocument(
            row["document_id"], version_id, row["title"],
            markdown_path.read_text(encoding="utf-8"), knowledgebase_ids, status,
        ))
    else:
        GLOBAL_CORPUS.set_status(version_id, status)


def _upload_job_diagnostic(error_code: str, **extra: Any) -> dict[str, Any]:
    job = UPLOAD_JOB_CONTEXT.get()
    diagnostic: dict[str, Any] = {"error_code": error_code}
    if job:
        diagnostic.update({
            "file_size_bytes": job["file_size_bytes"],
            "intended_owner_user_id": job.get("owner_user_id"),
            "job_id": job["job_id"],
            "knowledgebase_ids": job["knowledgebase_ids"],
            "original_filename": job["original_filename"],
            "title": job["title"],
            "uploaded_by_user_id": job["user_id"],
        })
    diagnostic.update(extra)
    return diagnostic


def _upload_job_fingerprint(error_code: str) -> str | None:
    job = UPLOAD_JOB_CONTEXT.get()
    return f"upload_job:{job['job_id']}:{error_code}" if job else None


def _notify_upload_failure(
    job: dict[str, Any] | None, incident_id: str | None, error_code: str,
) -> None:
    if not job:
        return
    recipient_ids = {job["user_id"]}
    if job.get("owner_user_id"):
        recipient_ids.add(job["owner_user_id"])
    stamp = datetime.now().astimezone().isoformat()
    for recipient_id in sorted(recipient_ids):
        try:
            recipient = PORTAL_GOVERNANCE.identity(recipient_id)
        except GovernanceError:
            continue
        if not recipient.active:
            continue
        notification_id = hashlib.sha256(
            f"upload_failed:{job['job_id']}:{recipient_id}".encode("utf-8")
        ).hexdigest()
        subject_title = job.get("title") or job.get("original_filename") or "Älterer unterbrochener Upload"
        reason_text = {
            "upload_worker_interrupted": "Die Verarbeitung wurde unterbrochen. Bitte lade das Dokument erneut hoch.",
            "document_conversion_unavailable": "Das Dokument konnte technisch nicht vollständig aufbereitet werden.",
            "required_check_unavailable": "Eine erforderliche technische Prüfung war nicht verfügbar.",
            "system_error": "Bei der Verarbeitung ist ein technischer Fehler aufgetreten.",
        }.get(error_code, "Das Dokument konnte nicht vollständig verarbeitet werden.")
        user_reason = f"Betroffenes Dokument: {subject_title}. {reason_text}"
        with PORTAL_GOVERNANCE.store.connect() as db:
            db.execute(
                "INSERT OR IGNORE INTO portal_notifications "
                "(notification_id,recipient_user_id,subject_type,subject_id,subject_title,status,message,reason,created_at) "
                "VALUES (?,?,?,?,?,?,?,?,?)",
                (notification_id, recipient_id, "upload_job", job["job_id"], subject_title,
                 "failed", "Das Dokument konnte nicht vollständig aufbereitet werden.",
                 user_reason, stamp),
            )


def _notify_system_error(
    step: str, diagnostic: dict[str, Any], *, fingerprint: str | None = None,
) -> str:
    incident_id = QUALITY_CASES.system_incident(step, diagnostic, fingerprint=fingerprint)
    for recipient in _admin_emails():
        MAINTENANCE.enqueue_notification(
            recipient, "system_error", "KAHLE-Vinci: Systemfehler",
            f"Ein Systemfehler ist aufgetreten. Referenz: {incident_id}",
            dedupe_key=incident_id,
        )
    return incident_id


@app.post("/portal/cases/{case_id}/action")
def portal_case_action(
    case_id: str, payload: PortalCaseActionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        if payload.action == "replace":
            if not payload.target_document_id:
                raise LifecycleError("replacement_target_required")
            before = DOCUMENT_LIFECYCLE.submission(case_id)
            source_dir = (PORTAL_FILES_ROOT / before.document_id / before.version_id).resolve()
            target_dir = (PORTAL_FILES_ROOT / payload.target_document_id / before.version_id).resolve()
            source_dir.relative_to(PORTAL_FILES_ROOT.resolve())
            target_dir.relative_to(PORTAL_FILES_ROOT.resolve())
            if not source_dir.is_dir() or target_dir.exists():
                raise LifecycleError("replacement_storage_not_ready")
            target_dir.parent.mkdir(parents=True, exist_ok=True)
            source_dir.rename(target_dir)
            bound_complete = False
            try:
                bound = DOCUMENT_LIFECYCLE.bind_replacement(
                    case_id=case_id, target_document_id=payload.target_document_id,
                    actor_user_id=identity["user_id"],
                )
                bound_complete = True
                RAG_METADATA.write(bound.version_id, target_dir / "rag.md")
                with PORTAL_GOVERNANCE.store.connect() as db:
                    kb_ids = tuple(row["knowledgebase_id"] for row in db.execute(
                        "SELECT knowledgebase_id FROM document_publications WHERE document_id=?",
                        (bound.document_id,),
                    ).fetchall())
                GLOBAL_CORPUS.upsert(CorpusDocument(
                    bound.document_id, bound.version_id, bound.title,
                    (target_dir / "rag.md").read_text(encoding="utf-8"), kb_ids, "pending",
                ))
                try:
                    source_dir.parent.rmdir()
                except OSError:
                    pass
            except Exception:
                if not bound_complete and target_dir.exists() and not source_dir.exists():
                    source_dir.parent.mkdir(parents=True, exist_ok=True)
                    target_dir.rename(source_dir)
                raise
        case = DOCUMENT_LIFECYCLE.choose_action(
            case_id=case_id, actor_user_id=identity["user_id"], action=payload.action,
        )
    except LifecycleError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    except Exception as exc:
        incident_id = _notify_system_error("replacement_binding", {"error_type": type(exc).__name__})
        raise HTTPException(status_code=500, detail=f"system_error:{incident_id}") from exc
    if case.status == "withdrawn":
        GLOBAL_CORPUS.set_status(case.version_id, "withdrawn")
    action_labels = {
        "create": "Als eigenständiges Dokument vorschlagen",
        "replace": "Als neue Version veröffentlichen",
        "publish_existing": "Vorhandenes Dokument zusätzlich veröffentlichen",
        "discard": "Upload verwerfen",
    }
    _notify_case_status(
        case, f"Gewählte Aktion: {action_labels.get(payload.action, payload.action)}",
    )
    return {"case": asdict(case)}


def _execute_portal_case_decision(
    case_id: str, payload: PortalCaseDecisionRequest, identity: dict[str, Any],
) -> dict[str, Any]:
    try:
        case = DOCUMENT_LIFECYCLE.submission(case_id)
        if case.status == "ready_to_activate":
            actor = PORTAL_GOVERNANCE.identity(identity["user_id"])
            may_retry = (
                payload.decision == "approve"
                and (
                    actor.role in {"admin", "portal_admin"}
                    or identity["user_id"] == case.manager_user_id
                    or PORTAL_GOVERNANCE.may_approve_for_manager(
                        identity["user_id"], case.manager_user_id,
                    )
                )
            )
            if not may_retry:
                raise LifecycleError("activation_retry_not_allowed")
        else:
            case = DOCUMENT_LIFECYCLE.decide(
                case_id=case_id, actor_user_id=identity["user_id"],
                decision=payload.decision, reason=payload.reason,
            )
        if case.status == "ready_to_activate" and case.requested_action == "publish_existing":
            case, target_version_id, previous_publication_status = DOCUMENT_LIFECYCLE.publish_existing(case_id=case_id)
            RAG_METADATA.write(target_version_id)
            indexing = _trigger_hybrid_version_sync(target_version_id)
            if not indexing.get("ok"):
                case = DOCUMENT_LIFECYCLE.rollback_existing_publication(
                    case_id=case_id, previous_status=previous_publication_status,
                    reason=str(indexing.get("error") or "hybrid_reindex_failed"),
                )
                _trigger_hybrid_version_sync(target_version_id)
                _refresh_global_corpus_version(case.version_id, "pending")
                _refresh_global_corpus_version(target_version_id, "active")
                raise HTTPException(status_code=503, detail="publication_index_failed_previous_scope_restored")
            _refresh_global_corpus_version(case.version_id, "withdrawn_duplicate")
            _refresh_global_corpus_version(target_version_id, "active")
        elif case.status == "ready_to_activate":
            previous_version_id = DOCUMENT_LIFECYCLE.active_version(case.document_id)
            case = DOCUMENT_LIFECYCLE.activate(case_id=case_id)
            RAG_METADATA.write(case.version_id)
            indexing = _trigger_hybrid_version_sync(case.version_id)
            if not indexing.get("ok"):
                case = DOCUMENT_LIFECYCLE.rollback_activation(
                    case_id=case_id, previous_version_id=previous_version_id,
                    reason=str(indexing.get("error") or "hybrid_reindex_failed"),
                )
                # Incremental sync is fail-closed and restores the old visible points itself.
                if previous_version_id:
                    _trigger_hybrid_version_sync(previous_version_id)
                _refresh_global_corpus_version(case.version_id, "pending")
                if previous_version_id:
                    _refresh_global_corpus_version(previous_version_id, "active")
                raise HTTPException(status_code=503, detail="activation_index_failed_previous_version_restored")
            _refresh_global_corpus_version(case.version_id, "active")
            if previous_version_id and previous_version_id != case.version_id:
                _refresh_global_corpus_version(previous_version_id, "superseded")
        elif case.status == "rejected":
            GLOBAL_CORPUS.set_status(case.version_id, "rejected")
    except HTTPException:
        raise
    except (LifecycleError, GovernanceError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    _notify_case_status(case, payload.reason, send_publication_email=True)
    return {"case": asdict(case)}


def _drain_decision_jobs() -> None:
    while job := DECISION_JOBS.claim_next():
        try:
            result = _execute_portal_case_decision(
                job["case_id"],
                PortalCaseDecisionRequest(decision=job["decision"], reason=job["reason"]),
                {"user_id": job["user_id"]},
            )
            DECISION_JOBS.complete(job["job_id"], result)
        except HTTPException as exc:
            DECISION_JOBS.fail(job["job_id"], str(exc.detail))
        except Exception as exc:  # pragma: no cover - last-resort persistent failure record
            incident_id = _notify_system_error("decision_queue", {
                "job_id": job["job_id"], "error": type(exc).__name__,
            })
            DECISION_JOBS.fail(job["job_id"], f"system_error:{incident_id}")


@app.post("/portal/cases/{case_id}/decision", status_code=202)
def portal_case_decision(
    case_id: str, payload: PortalCaseDecisionRequest, background_tasks: BackgroundTasks,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    job = DECISION_JOBS.enqueue(
        case_id, identity["user_id"], payload.decision, payload.reason,
    )
    background_tasks.add_task(_drain_decision_jobs)
    return job


@app.get("/portal/decision-jobs/{job_id}")
def portal_decision_job(
    job_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        return DECISION_JOBS.get(
            job_id, identity["user_id"], identity["role"] in {"admin", "portal_admin"},
        )
    except DecisionJobError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.get("/portal/decision-jobs")
def portal_decision_jobs(
    active: bool = True,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    # The first UI iteration deliberately exposes only active work. Completed
    # outcomes already arrive through the normal notification stream.
    if not active:
        return {"jobs": []}
    return {"jobs": DECISION_JOBS.list_active(
        identity["user_id"], identity["role"] in {"admin", "portal_admin"},
    )}


@app.get("/portal/documents/{document_id}/authority")
def portal_document_authority(
    document_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        return DOCUMENT_AUTHORITY.view(identity["user_id"], document_id)
    except AuthorityError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.put("/portal/documents/{document_id}/authority")
def portal_update_document_authority(
    document_id: str, payload: PortalAuthorityRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        result = DOCUMENT_AUTHORITY.update(
            identity["user_id"], document_id, payload.authority_type, payload.scope, payload.reason,
        )
        with PORTAL_GOVERNANCE.store.connect() as db:
            version = db.execute(
                "SELECT version_id FROM document_versions WHERE document_id=? ORDER BY created_at DESC LIMIT 1",
                (document_id,),
            ).fetchone()
        if version:
            RAG_METADATA.write(version["version_id"])
            indexing = _trigger_hybrid_document_sync(document_id)
            if not indexing.get("ok"):
                incident_id = _notify_system_error(
                    "authority_reindex", {"document_id": document_id, "error": indexing.get("error")}
                )
                raise HTTPException(status_code=503, detail=f"authority_reindex_failed:{incident_id}")
        return {**result, "reindex": {"ok": True}}
    except AuthorityError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc


@app.post("/portal/documents/{document_id}/authority-relations", status_code=201)
def portal_create_authority_relation(
    document_id: str, payload: PortalAuthorityRelationRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        relation_id = DOCUMENT_AUTHORITY.relate(
            identity["user_id"], document_id, payload.target_document_id,
            payload.relation_type, payload.condition_text, payload.reason,
        )
        return {"relation_id": relation_id}
    except AuthorityError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc


@app.get("/portal/ownership-tasks")
def portal_ownership_tasks(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    return {"tasks": OWNERSHIP.tasks_for(identity["user_id"])}


@app.get("/portal/owner-candidates")
def portal_owner_candidates(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    return {
        "can_propose_other": OWNERSHIP.may_propose_other(identity["user_id"]),
        "users": OWNERSHIP.active_candidates(identity["user_id"]),
    }


@app.get("/portal/admin/users/{target_user_id}/owner-proposal-permission")
def portal_admin_owner_permission(
    target_user_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"allowed": OWNERSHIP.may_propose_other(target_user_id)}


@app.put("/portal/admin/users/{target_user_id}/owner-proposal-permission")
def portal_admin_set_owner_permission(
    target_user_id: str, payload: PortalOwnerPermissionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        OWNERSHIP.set_proposal_permission(identity["user_id"], target_user_id, payload.allowed)
    except OwnershipError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    return {"allowed": payload.allowed}


@app.post("/portal/ownership-tasks/{task_id}/proposal")
def portal_propose_owner(
    task_id: str, payload: PortalOwnershipProposalRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        OWNERSHIP.propose(task_id, identity["user_id"], payload.proposed_owner_user_id, payload.reason)
    except OwnershipError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"status": "pending_owner_confirmation"}


@app.post("/portal/ownership-tasks/{task_id}/confirmation")
def portal_confirm_owner(
    task_id: str, payload: PortalOwnershipConfirmationRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        status = OWNERSHIP.confirm(task_id, identity["user_id"], payload.accept, payload.reason)
        if status == "completed" and payload.accept:
            with PORTAL_GOVERNANCE.store.connect() as db:
                version = db.execute(
                    """SELECT v.version_id, t.document_id FROM owner_reassignment_tasks t
                       JOIN canonical_documents d ON d.document_id=t.document_id
                       JOIN document_versions v ON v.document_id=d.document_id
                       WHERE t.task_id=? ORDER BY v.created_at DESC LIMIT 1""", (task_id,),
                ).fetchone()
            if version:
                RAG_METADATA.write(version["version_id"])
                indexing = _trigger_hybrid_document_sync(version["document_id"])
                if not indexing.get("ok"):
                    incident_id = _notify_system_error(
                        "owner_reindex", {"task_id": task_id, "error": indexing.get("error")}
                    )
                    raise HTTPException(status_code=503, detail=f"owner_reindex_failed:{incident_id}")
    except OwnershipError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"status": status}


@app.post("/portal/document-changes/renewal", status_code=201)
def portal_request_renewal(
    payload: PortalRenewalRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try: request_id=DOCUMENT_CHANGES.request_renewal(payload.document_id,identity["user_id"],payload.reason,payload.confirmed)
    except DocumentChangeError as exc: raise HTTPException(status_code=409,detail=str(exc)) from exc
    return {"request_id":request_id}


@app.post("/portal/document-changes/confidentiality", status_code=201)
def portal_request_confidentiality(
    payload: PortalConfidentialityRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try: request_id=DOCUMENT_CHANGES.request_confidentiality(payload.document_id,identity["user_id"],payload.desired,payload.reason)
    except DocumentChangeError as exc: raise HTTPException(status_code=409,detail=str(exc)) from exc
    return {"request_id":request_id}


@app.get("/portal/document-changes")
def portal_document_changes(identity: dict[str, Any] = Depends(require_portal_identity)) -> dict[str, Any]:
    return {"changes":DOCUMENT_CHANGES.pending_for(identity["user_id"])}


@app.post("/portal/document-changes/{request_id}/decision")
def portal_decide_document_change(
    request_id: str,payload: PortalDocumentChangeDecisionRequest,identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try: status=DOCUMENT_CHANGES.decide(request_id,identity["user_id"],payload.approve,payload.reason)
    except DocumentChangeError as exc: raise HTTPException(status_code=409,detail=str(exc)) from exc
    if status=="approved":
        with PORTAL_GOVERNANCE.store.connect() as db:
            change = db.execute(
                "SELECT document_id FROM document_change_requests WHERE request_id=?", (request_id,),
            ).fetchone()
        if change: _trigger_hybrid_document_sync(change["document_id"])
    return {"status":status}


@app.post("/portal/removal-requests", status_code=201)
def portal_request_removal(
    payload: PortalRemovalRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    title, readers = _document_access_snapshot(payload.document_id)
    try:
        request_id = MAINTENANCE.request_removal(payload.document_id, identity["user_id"], payload.kind, payload.reason)
    except MaintenanceError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    if identity["role"] in {"admin", "portal_admin"}:
        indexing = _trigger_hybrid_document_sync(payload.document_id)
        if not indexing.get("ok"): raise HTTPException(status_code=503, detail="removal_reindex_failed")
        _notify_access_removed(
            readers, subject_type="document", subject_id=payload.document_id,
            subject_title=title, status="removed",
            message="Das Dokument wurde aus Vinci entfernt und ist nicht mehr abrufbar.",
            reason=payload.reason,
        )
    return {"request_id": request_id}


@app.get("/portal/admin/removals")
def portal_admin_removals(identity: dict[str, Any] = Depends(require_portal_identity)) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}: raise HTTPException(status_code=403, detail="admin_required")
    return MAINTENANCE.list_removals(identity["user_id"])


@app.get("/portal/admin/archive")
def portal_admin_archived_versions(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    with PORTAL_GOVERNANCE.store.connect() as db:
        rows = db.execute(
            """SELECT v.version_id, v.document_id, v.title, v.original_filename, v.status,
                      v.superseded_at, v.purged_at, v.valid_from, v.valid_until,
                      d.active_version_id, active.title active_version_title,
                      active.original_filename active_original_filename,
                      (SELECT COUNT(*) FROM document_versions all_versions
                       WHERE all_versions.document_id=v.document_id) version_count
               FROM document_versions v
               JOIN canonical_documents d ON d.document_id=v.document_id
               LEFT JOIN document_versions active ON active.version_id=d.active_version_id
               WHERE v.status IN ('superseded', 'purged')
               ORDER BY COALESCE(v.purged_at, v.superseded_at, v.created_at) DESC, v.version_id DESC""",
        ).fetchall()
    versions = []
    for row in rows:
        item = dict(row)
        _, originals, markdown = _archived_version_files(item["document_id"], item["version_id"])
        item["has_original"] = len(originals) == 1
        item["can_restore"] = item["status"] == "superseded" and item["has_original"] and markdown.is_file()
        versions.append(item)
    return {"versions": versions}


@app.get("/portal/admin/archive/{version_id}/source")
def portal_admin_archived_version_source(
    version_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        record = DOCUMENT_LIFECYCLE.version_record(version_id)
    except LifecycleError as exc:
        raise HTTPException(status_code=404, detail="archived_version_not_available") from exc
    if record["status"] not in {"superseded", "purged"}:
        raise HTTPException(status_code=404, detail="archived_version_not_available")
    _, originals, _ = _archived_version_files(record["document_id"], version_id)
    if len(originals) != 1:
        raise HTTPException(status_code=404, detail="archived_original_not_available")
    disposition = "inline" if originals[0].suffix.lower() in {".pdf", ".txt", ".md"} else "attachment"
    return FileResponse(
        originals[0], filename=record["original_filename"], content_disposition_type=disposition,
        headers={"Cache-Control": "private, no-store", "X-Content-Type-Options": "nosniff"},
    )


@app.post("/portal/admin/archive/{version_id}/restore")
def portal_admin_restore_archived_version(
    version_id: str, payload: PortalArchivedVersionRestoreRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        record = DOCUMENT_LIFECYCLE.version_record(version_id)
    except LifecycleError as exc:
        raise HTTPException(status_code=404, detail="archived_version_not_available") from exc
    if record["status"] != "superseded":
        raise HTTPException(status_code=409, detail="archived_version_restore_not_allowed")
    _, originals, markdown = _archived_version_files(record["document_id"], version_id)
    if len(originals) != 1 or not markdown.is_file():
        raise HTTPException(status_code=409, detail="archived_version_payload_not_available")
    previous_version_id = DOCUMENT_LIFECYCLE.active_version(record["document_id"])
    try:
        restored = DOCUMENT_LIFECYCLE.restore_superseded_version(
            version_id=version_id, actor_user_id=identity["user_id"], reason=payload.reason,
        )
        RAG_METADATA.write(restored.version_id)
        indexing = _trigger_hybrid_version_sync(restored.version_id)
        if not indexing.get("ok"):
            DOCUMENT_LIFECYCLE.rollback_superseded_version_restore(
                restored_version_id=restored.version_id, previous_version_id=previous_version_id,
                actor_user_id="indexer", reason=str(indexing.get("error") or "hybrid_reindex_failed"),
            )
            _trigger_hybrid_version_sync(previous_version_id)
            _refresh_global_corpus_version(restored.version_id, "superseded")
            _refresh_global_corpus_version(previous_version_id, "active")
            raise HTTPException(status_code=503, detail="archive_restore_index_failed_previous_version_restored")
    except HTTPException:
        raise
    except LifecycleError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    _refresh_global_corpus_version(restored.version_id, "active")
    if previous_version_id and previous_version_id != restored.version_id:
        _refresh_global_corpus_version(previous_version_id, "superseded")
    return {"version": asdict(restored)}


@app.post("/portal/admin/trash/read")
def portal_admin_mark_trash_read(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, bool]:
    try:
        MAINTENANCE.mark_trash_read(identity["user_id"])
    except MaintenanceError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc
    return {"ok": True}


@app.post("/portal/admin/trash/{document_id}/delete")
def portal_admin_delete_from_trash(
    document_id: str, payload: PortalRestoreRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    """Endgueltige Loeschung durch Admins ab Tag 30 (PRD 23.3)."""
    if not payload.confirmed:
        raise HTTPException(status_code=409, detail="confirmation_required")
    try:
        MAINTENANCE.delete_now(
            document_id, identity["user_id"], payload.reason, file_root=PORTAL_FILES_ROOT,
        )
    except MaintenanceError as exc:
        status = 403 if str(exc) == "admin_required" else 409
        raise HTTPException(status_code=status, detail=str(exc)) from exc
    PORTAL_GOVERNANCE.record_audit(
        identity["user_id"], "document_deleted_from_trash", "document", document_id,
        {"reason": payload.reason},
    )
    return {"document_id": document_id, "status": "deleted"}


@app.post("/portal/admin/removal-requests/{request_id}/decision")
def portal_admin_decide_removal(
    request_id: str, payload: PortalRemovalDecisionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    with PORTAL_GOVERNANCE.store.connect() as db:
        removal = db.execute(
            "SELECT document_id FROM document_removal_requests WHERE request_id=?", (request_id,),
        ).fetchone()
    title, readers = _document_access_snapshot(removal["document_id"]) if removal else ("Dokument", [])
    try: MAINTENANCE.decide_removal(request_id, identity["user_id"], payload.approve, payload.reason)
    except MaintenanceError as exc: raise HTTPException(status_code=409, detail=str(exc)) from exc
    if payload.approve:
        indexing = _trigger_hybrid_document_sync(removal["document_id"] if removal else "")
        if not indexing.get("ok"): raise HTTPException(status_code=503, detail="removal_reindex_failed")
        if removal:
            _notify_access_removed(
                readers, subject_type="document", subject_id=removal["document_id"],
                subject_title=title, status="removed",
                message="Das Dokument wurde aus Vinci entfernt und ist nicht mehr abrufbar.",
                reason=payload.reason,
            )
    return {"ok": True}


@app.post("/portal/admin/trash/{document_id}/restore")
def portal_admin_restore_trash(
    document_id: str, payload: PortalRestoreRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try: MAINTENANCE.restore_from_trash(document_id, identity["user_id"], payload.reason)
    except MaintenanceError as exc: raise HTTPException(status_code=409, detail=str(exc)) from exc
    indexing = _trigger_hybrid_document_sync(document_id)
    if not indexing.get("ok"): raise HTTPException(status_code=503, detail="restore_reindex_failed")
    return {"ok": True}


@app.put("/portal/admin/trash/{document_id}/legal-hold")
def portal_admin_legal_hold(
    document_id: str, payload: PortalLegalHoldRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try: MAINTENANCE.set_legal_hold(document_id, identity["user_id"], payload.enabled, payload.reason, payload.review_at)
    except MaintenanceError as exc: raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"ok": True}


@app.get("/portal/admin/dashboard")
def portal_admin_dashboard(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    result = QUALITY_DASHBOARD.snapshot()
    try:
        response = requests.get(f"{KB_SYNC_URL}/health", timeout=10)
        result["index"] = response.json() if response.status_code == 200 else {"ok": False}
    except requests.RequestException:
        result["index"] = {"ok": False, "error": "unavailable"}
    return result


@app.post("/portal/admin/migration/inventory")
def portal_admin_migration_inventory(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"items": [asdict(item) for item in LEGACY_MIGRATION.inventory(KB_ROOT)]}


@app.get("/portal/admin/migration/inventory")
def portal_admin_migration_inventory_status(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"items": [asdict(item) for item in LEGACY_MIGRATION.inventory_items()]}


@app.get("/portal/admin/migration/tasks")
def portal_admin_migration_tasks(
    status: str = "open", identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        return {"tasks": LEGACY_MIGRATION.tasks(status)}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/portal/admin/migration/file")
def portal_admin_migration_file(
    path: str, kind: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> FileResponse:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    try:
        review_file = LEGACY_MIGRATION.review_file(KB_ROOT, path, kind)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    media_type = "text/markdown; charset=utf-8" if kind == "markdown" else None
    return FileResponse(
        review_file,
        filename=review_file.name,
        media_type=media_type,
        content_disposition_type="inline",
    )


@app.put("/portal/admin/migration/metadata")
def portal_admin_migration_metadata(
    payload: PortalMigrationMetadataRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        LEGACY_MIGRATION.resolve_metadata(
            payload.path, identity["user_id"], owner_email=payload.owner_email,
            confidentiality=payload.confidentiality, authority_type=payload.authority_type,
            authority_level=payload.authority_level, scope=payload.scope,
            knowledgebase_id=payload.knowledgebase_id,
        )
    except (GovernanceError, ValueError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"status": "metadata_resolved"}


@app.post("/portal/admin/migration/stage")
def portal_admin_migration_stage(
    payload: PortalMigrationStageRequest, request: Request,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    if not payload.confirmed:
        raise HTTPException(status_code=409, detail="confirmation_required")
    try:
        case_id = LEGACY_MIGRATION.stage(KB_ROOT, payload.path, identity["user_id"])
    except (ValueError, LifecycleError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"case_id": case_id, "status": "staged"}


@app.post("/portal/admin/migration/exclude")
def portal_admin_migration_exclude(
    payload: PortalMigrationDispositionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    try:
        LEGACY_MIGRATION.exclude(payload.path, identity["user_id"], payload.reason)
    except (GovernanceError, ValueError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"status": "excluded"}


@app.post("/portal/admin/migration/restore")
def portal_admin_migration_restore(
    payload: PortalMigrationDispositionRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, str]:
    try:
        LEGACY_MIGRATION.restore_excluded(payload.path, identity["user_id"], payload.reason)
    except (GovernanceError, ValueError) as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"status": "restored"}


@app.get("/portal/admin/audit")
def portal_admin_audit(
    limit: int = 250, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"entries": [asdict(item) for item in AUDIT_EXPORTER.entries(limit)]}


@app.get("/portal/admin/audit/export.{format}")
def portal_admin_audit_export(
    format: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> FastAPIResponse:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    if format == "csv":
        return FastAPIResponse(AUDIT_EXPORTER.csv_bytes(), media_type="text/csv; charset=utf-8",
            headers={"Content-Disposition": 'attachment; filename="kahle-vinci-audit.csv"'})
    if format == "pdf":
        return FastAPIResponse(AUDIT_EXPORTER.pdf_bytes(), media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="kahle-vinci-audit.pdf"'})
    raise HTTPException(status_code=404, detail="unknown_audit_format")


def _sync_openwebui_user_directory(request: Request) -> None:
    if DEV_AUTH_BYPASS:
        return
    headers: dict[str, str] = {}
    if request.headers.get("Authorization"):
        headers["Authorization"] = request.headers["Authorization"]
    if request.headers.get("Cookie"):
        headers["Cookie"] = request.headers["Cookie"]
    try:
        response = requests.get(f"{OPENWEBUI_URL}/api/v1/users/", headers=headers, timeout=20)
        response.raise_for_status()
        payload = response.json()
        users = payload.get("users", payload) if isinstance(payload, dict) else payload
        if not isinstance(users, list):
            raise ValueError("invalid_openwebui_user_directory")
        for item in users:
            user_id = str(item.get("id") or "").strip()
            email = str(item.get("email") or "").strip()
            email_domain = email.rsplit("@", 1)[-1].casefold() if "@" in email else ""
            if not user_id or not email or email_domain not in PORTAL_ALLOWED_EMAIL_DOMAINS:
                continue
            PORTAL_GOVERNANCE.sync_identity(
                user_id=user_id, email=email,
                display_name=str(item.get("name") or item.get("display_name") or email).strip(),
                active=True, bootstrap_portal_admin=False,
            )
    except Exception as exc:
        incident_id = _notify_system_error(
            "openwebui_user_directory_sync", {"error_type": type(exc).__name__}
        )
        raise HTTPException(status_code=503, detail=f"user_directory_sync_failed:{incident_id}") from exc


@app.get("/portal/admin/users")
def portal_admin_users(
    request: Request, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if str(identity.get("openwebui_role") or "").lower() == "admin":
        _sync_openwebui_user_directory(request)
    users = _portal_call(lambda: PORTAL_GOVERNANCE.list_identities(identity["user_id"]))
    return {"users": serialize_governance(users)}


@app.get("/portal/admin/restricted-terms")
def portal_admin_restricted_terms(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        return {"terms": RESTRICTED_TERMS.serialized(identity["user_id"])}
    except RestrictedTermError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc


@app.get("/portal/admin/settings/auto-activation")
def portal_admin_auto_activation_setting(
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, bool]:
    if identity["role"] not in {"admin", "portal_admin"}:
        raise HTTPException(status_code=403, detail="admin_required")
    return {"enabled": PORTAL_GOVERNANCE.setting_bool(
        "auto_activation_enabled", default=_AUTO_ACTIVATION_DEFAULT,
    )}


@app.put("/portal/admin/settings/auto-activation")
def portal_admin_set_auto_activation(
    payload: PortalAutoActivationRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, bool]:
    if identity["role"] != "portal_admin":
        raise HTTPException(status_code=403, detail="portal_admin_required")
    enabled = _portal_call(lambda: PORTAL_GOVERNANCE.set_setting_bool(
        identity["user_id"], "auto_activation_enabled", payload.enabled, payload.reason,
    ))
    return {"enabled": enabled}


@app.post("/portal/admin/restricted-terms", status_code=201)
def portal_admin_add_restricted_term(
    payload: PortalRestrictedTermRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    try:
        return {"term": asdict(RESTRICTED_TERMS.add(identity["user_id"], payload.term))}
    except RestrictedTermError as exc:
        status = 403 if str(exc) == "admin_required" else 422
        raise HTTPException(status_code=status, detail=str(exc)) from exc


@app.delete("/portal/admin/restricted-terms/{rule_id}")
def portal_admin_remove_restricted_term(
    rule_id: str,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, bool]:
    try:
        RESTRICTED_TERMS.remove(identity["user_id"], rule_id)
    except RestrictedTermError as exc:
        status = 403 if str(exc) == "admin_required" else 404
        raise HTTPException(status_code=status, detail=str(exc)) from exc
    return {"removed": True}


@app.patch("/portal/admin/users/{target_user_id}/role")
def portal_admin_set_role(
    target_user_id: str,
    payload: PortalRoleRequest,
    request: Request,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if not payload.confirmed:
        raise HTTPException(status_code=409, detail="confirmation_required")
    user = _portal_call(
        lambda: PORTAL_GOVERNANCE.set_role(identity["user_id"], target_user_id, payload.role)
    )
    return {"user": serialize_governance(user)}


@app.patch("/portal/admin/users/{target_user_id}/activation")
def portal_admin_set_activation(
    target_user_id: str,
    payload: PortalActivationRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    user = _portal_call(
        lambda: PORTAL_GOVERNANCE.set_active(identity["user_id"], target_user_id, payload.active)
    )
    reassignment_tasks = [] if payload.active else OWNERSHIP.create_for_deactivated_owner(
        target_user_id, identity["user_id"]
    )
    return {"user": serialize_governance(user), "owner_reassignment_task_ids": reassignment_tasks}


@app.patch("/portal/admin/users/{target_user_id}/manager")
def portal_admin_assign_manager(
    target_user_id: str,
    payload: PortalManagerRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    user = _portal_call(
        lambda: PORTAL_GOVERNANCE.assign_manager(
            identity["user_id"], target_user_id, payload.manager_user_id
        )
    )
    return {"user": serialize_governance(user)}


@app.get("/portal/admin/absences")
def portal_admin_absences(identity: dict[str, Any] = Depends(require_portal_identity)) -> dict[str, Any]:
    return {"absences": _portal_call(lambda: PORTAL_GOVERNANCE.list_absences(identity["user_id"]))}


@app.put("/portal/admin/absences")
def portal_admin_set_absence(
    payload: PortalAbsenceRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    _portal_call(lambda: PORTAL_GOVERNANCE.set_absence(
        identity["user_id"], payload.manager_user_id, payload.absent_from,
        payload.absent_until, payload.reason, payload.delegate_user_id,
    ))
    return {"ok": True}


@app.get("/portal/admin/delegations")
def portal_admin_delegations(identity: dict[str, Any] = Depends(require_portal_identity)) -> dict[str, Any]:
    return {"delegations": _portal_call(lambda: PORTAL_GOVERNANCE.list_delegations(identity["user_id"]))}


@app.put("/portal/admin/delegations")
def portal_admin_set_delegation(
    payload: PortalDelegationRequest, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    _portal_call(lambda: PORTAL_GOVERNANCE.assign_delegate(identity["user_id"],payload.manager_user_id,payload.delegate_user_id,valid_from=payload.valid_from,valid_until=payload.valid_until))
    return {"ok":True}


@app.delete("/portal/admin/delegations")
def portal_admin_remove_delegation(
    manager_user_id: str, delegate_user_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    _portal_call(lambda: PORTAL_GOVERNANCE.remove_delegate(identity["user_id"],manager_user_id,delegate_user_id))
    return {"ok":True}


@app.get("/portal/admin/users/{target_user_id}/knowledgebase-access")
def portal_admin_user_access(
    target_user_id: str, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    access = _portal_call(
        lambda: PORTAL_GOVERNANCE.access_for_user(identity["user_id"], target_user_id)
    )
    return {"access": access}


@app.put("/portal/admin/users/{target_user_id}/knowledgebase-access")
def portal_admin_grant_access(
    target_user_id: str,
    payload: PortalAccessRequest,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    _portal_call(
        lambda: PORTAL_GOVERNANCE.grant_access(
            identity["user_id"],
            target_user_id,
            payload.knowledgebase_id,
            can_read=payload.can_read,
            can_upload=payload.can_upload,
        )
    )
    return {"ok": True}


@app.get("/portal/admin/knowledgebase-changes")
def portal_admin_list_knowledgebase_changes(
    status: str | None = None, identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    changes = _portal_call(lambda: PORTAL_GOVERNANCE.list_change_requests(identity["user_id"], status))
    return {"changes": serialize_governance(changes)}


@app.post("/portal/admin/knowledgebase-changes", status_code=201)
def portal_admin_request_knowledgebase_change(
    payload: PortalKnowledgebaseChangeRequest,
    request: Request,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if identity["role"] == "portal_admin" and not payload.confirmed:
        raise HTTPException(status_code=409, detail="confirmation_required")
    access_snapshot = None
    if payload.kind in {"archive", "delete"} and payload.knowledgebase_id:
        access_snapshot = _knowledgebase_access_snapshot(payload.knowledgebase_id)
    change = _portal_call(
        lambda: PORTAL_GOVERNANCE.request_knowledgebase_change(
            identity["user_id"],
            payload.kind,
            knowledgebase_id=payload.knowledgebase_id,
            payload=payload.payload,
        )
    )
    if change.status == "approved" and access_snapshot:
        label, readers = access_snapshot
        action = "archiviert" if payload.kind == "archive" else "gelöscht"
        _notify_access_removed(
            readers, subject_type="knowledgebase", subject_id=change.knowledgebase_id or payload.knowledgebase_id,
            subject_title=label, status=f"knowledgebase_{payload.kind}",
            message=f"Der Wissensbereich wurde {action} und ist in Vinci nicht mehr abrufbar.",
            reason=str(payload.payload.get("reason") or ""),
        )
    return {"change": serialize_governance(change)}

@app.post("/portal/admin/knowledgebase-changes/{request_id}/decision")
def portal_admin_decide_knowledgebase_change(
    request_id: str,
    payload: PortalKnowledgebaseDecisionRequest,
    request: Request,
    identity: dict[str, Any] = Depends(require_portal_identity),
) -> dict[str, Any]:
    if not payload.confirmed:
        raise HTTPException(status_code=409, detail="confirmation_required")
    pending_change = _portal_call(lambda: PORTAL_GOVERNANCE.change_request(request_id))
    access_snapshot = None
    if pending_change.kind in {"archive", "delete"} and pending_change.knowledgebase_id:
        access_snapshot = _knowledgebase_access_snapshot(pending_change.knowledgebase_id)
    change = _portal_call(
        lambda: PORTAL_GOVERNANCE.decide_knowledgebase_change(
            identity["user_id"],
            request_id,
            approve=payload.approve,
            reason=payload.reason,
        )
    )
    if payload.approve and access_snapshot:
        label, readers = access_snapshot
        action = "archiviert" if pending_change.kind == "archive" else "gelöscht"
        _notify_access_removed(
            readers, subject_type="knowledgebase",
            subject_id=change.knowledgebase_id or pending_change.knowledgebase_id,
            subject_title=label, status=f"knowledgebase_{pending_change.kind}",
            message=f"Der Wissensbereich wurde {action} und ist in Vinci nicht mehr abrufbar.",
            reason=payload.reason,
        )
    return {"change": serialize_governance(change)}

@app.get("/unlock/status")
def unlock_status(
    request: Request, admin: dict[str, Any] = Depends(require_admin)
) -> dict[str, Any]:
    return {
        "enabled": bool(UNLOCK_ENABLED or DEV_AUTH_BYPASS),
        "unlocked": _has_valid_unlock(request, admin),
        "ttl_seconds": UNLOCK_TTL_SECONDS,
    }


@app.post("/unlock")
def unlock(
    payload: UnlockRequest,
    request: Request,
    response: Response,
    admin: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    if DEV_AUTH_BYPASS:
        return {"ok": True, "unlocked": True, "ttl_seconds": UNLOCK_TTL_SECONDS}
    if not UNLOCK_ENABLED:
        raise HTTPException(status_code=503, detail="admin_unlock_not_configured")
    key, recent = _attempt_state(request, admin)
    if len(recent) >= UNLOCK_MAX_ATTEMPTS:
        raise HTTPException(
            status_code=429,
            detail="admin_unlock_temporarily_blocked",
            headers={"Retry-After": str(UNLOCK_BLOCK_SECONDS)},
        )
    if not _unlock_code_matches(payload.code):
        with _unlock_failures_lock:
            _unlock_failures.setdefault(key, []).append(time.time())
        _audit(admin, "admin_unlock_failed", remote=request.client.host if request.client else "unknown")
        raise HTTPException(status_code=401, detail="admin_unlock_code_invalid")
    with _unlock_failures_lock:
        _unlock_failures.pop(key, None)
    response.set_cookie(
        key=UNLOCK_COOKIE,
        value=_issue_unlock_token(admin),
        max_age=UNLOCK_TTL_SECONDS,
        httponly=True,
        secure=True,
        samesite="strict",
        path="/admin/vector",
    )
    _audit(admin, "admin_unlocked", ttl_seconds=UNLOCK_TTL_SECONDS)
    return {"ok": True, "unlocked": True, "ttl_seconds": UNLOCK_TTL_SECONDS}


@app.post("/lock")
def lock(
    response: Response, admin: dict[str, Any] = Depends(require_admin)
) -> dict[str, Any]:
    response.delete_cookie(
        key=UNLOCK_COOKIE,
        httponly=True,
        secure=True,
        samesite="strict",
        path="/admin/vector",
    )
    _audit(admin, "admin_locked")
    return {"ok": True, "unlocked": False}

@app.get("/collections")
def list_collections(admin: dict[str, Any] = Depends(require_unlocked_admin)) -> dict[str, Any]:
    del admin
    state = _load_state()
    items: list[dict[str, Any]] = []
    for collection in _collection_names():
        root = KB_ROOT / collection
        documents = [
            path
            for path in root.rglob("*")
            if path.is_file()
            and not path.name.startswith((".", "~$"))
            and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ] if root.exists() else []
        collection_state = ((state.get("collections") or {}).get(collection) or {})
        items.append(
            {
                "id": collection,
                "label": _collection_label(collection),
                "files": len(documents),
                "last_indexed_at": collection_state.get("last_reconcile_at") or "",
                "deletable": collection not in COLLECTIONS,
            }
        )
    return {"collections": items}


@app.post("/collections", status_code=201)
def create_collection(
    payload: CreateCollectionRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    collection = payload.id.strip().lower()
    if collection in _collection_names():
        raise HTTPException(status_code=409, detail="collection_already_exists")
    root = (KB_ROOT / collection).resolve()
    try:
        root.relative_to(KB_ROOT)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="invalid_collection_path") from exc
    root.mkdir(parents=False, exist_ok=False)
    metadata = {
        "id": collection,
        "label": payload.label.strip(),
        "created_at": datetime.now().astimezone().isoformat(),
        "created_by": admin.get("email") or admin.get("name") or admin.get("id"),
    }
    _atomic_write(
        root / ".collection.json",
        json.dumps(metadata, ensure_ascii=False, indent=2).encode("utf-8"),
    )
    try:
        _qdrant(
            "PUT",
            f"/collections/{collection}",
            json={"vectors": {"size": 1024, "distance": "Cosine"}},
        )
    except HTTPException:
        shutil.rmtree(root)
        raise
    _audit(admin, "collection_created", collection=collection, label=payload.label.strip())
    return {"ok": True, "collection": {"id": collection, "label": payload.label.strip(), "files": 0, "deletable": True}}


@app.patch("/collections/{collection}")
def update_collection(
    collection: str,
    payload: UpdateCollectionRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    collection = _safe_collection(collection)
    root = KB_ROOT / collection
    if not root.exists():
        raise HTTPException(status_code=404, detail="collection_source_missing")
    metadata_path = root / ".collection.json"
    try:
        metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        if not isinstance(metadata, dict):
            metadata = {}
    except Exception:
        metadata = {}
    previous_label = _collection_label(collection)
    metadata.update(
        {
            "id": collection,
            "label": payload.label.strip(),
            "updated_at": datetime.now().astimezone().isoformat(),
            "updated_by": admin.get("email") or admin.get("name") or admin.get("id"),
        }
    )
    _atomic_write(
        metadata_path,
        json.dumps(metadata, ensure_ascii=False, indent=2).encode("utf-8"),
    )
    _audit(
        admin,
        "collection_updated",
        collection=collection,
        previous_label=previous_label,
        label=payload.label.strip(),
    )
    return {
        "ok": True,
        "collection": {
            "id": collection,
            "label": payload.label.strip(),
            "deletable": collection not in COLLECTIONS,
        },
    }


@app.delete("/collections/{collection}")
def delete_collection(
    collection: str,
    payload: DeleteCollectionRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    collection = _safe_collection(collection)
    if collection in COLLECTIONS:
        raise HTTPException(status_code=409, detail="protected_collection")
    if payload.confirm_id.strip() != collection:
        raise HTTPException(status_code=400, detail="collection_confirmation_mismatch")
    root = KB_ROOT / collection
    if not root.exists():
        raise HTTPException(status_code=404, detail="collection_source_missing")
    archived_at = datetime.now().astimezone()
    purge_at = archived_at + timedelta(days=TRASH_RETENTION_DAYS)
    archive = TRASH_ROOT / "collections" / f"{archived_at.strftime('%Y%m%d-%H%M%S')}-{collection}"
    archive.parent.mkdir(parents=True, exist_ok=True)
    label = _collection_label(collection)
    files = sum(
        1
        for path in root.rglob("*")
        if path.is_file()
        and not path.name.startswith(".")
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    shutil.move(str(root), str(archive))
    try:
        _qdrant("DELETE", f"/collections/{collection}")
    except HTTPException:
        shutil.move(str(archive), str(root))
        raise
    manifest = {
        "collection": collection,
        "label": label,
        "files": files,
        "archived_at": archived_at.isoformat(),
        "purge_at": purge_at.isoformat(),
        "retention_days": TRASH_RETENTION_DAYS,
        "deleted_by": admin.get("email") or admin.get("name") or admin.get("id"),
    }
    _atomic_write(
        archive / ".trash.json",
        json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"),
    )
    _audit(
        admin,
        "collection_deleted",
        collection=collection,
        archived_path=archive.relative_to(KB_ROOT).as_posix(),
        purge_at=purge_at.isoformat(),
    )
    return {
        "ok": True,
        "collection": collection,
        "archived_path": archive.relative_to(KB_ROOT).as_posix(),
        "purge_at": purge_at.isoformat(),
        "retention_days": TRASH_RETENTION_DAYS,
    }


@app.get("/trash/collections")
def list_trashed_collections(
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    return {
        "collections": _trash_collections(),
        "retention_days": TRASH_RETENTION_DAYS,
    }


@app.post("/trash/collections/{archive_id}/restore")
def restore_trashed_collection(
    archive_id: str,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    archive_id = _safe_archive_id(archive_id)
    archive = TRASH_ROOT / "collections" / archive_id
    if not archive.is_dir():
        raise HTTPException(status_code=404, detail="trash_collection_not_found")
    item = _trash_manifest(archive)
    collection = item["collection"]
    if collection in _collection_names():
        raise HTTPException(status_code=409, detail="collection_already_exists")
    destination = KB_ROOT / collection
    _qdrant(
        "PUT",
        f"/collections/{collection}",
        json={"vectors": {"size": 1024, "distance": "Cosine"}},
    )
    try:
        (archive / ".trash.json").unlink(missing_ok=True)
        shutil.move(str(archive), str(destination))
    except Exception:
        try:
            _qdrant("DELETE", f"/collections/{collection}")
        except HTTPException:
            pass
        raise
    reindex = _trigger_reindex(collection)
    _audit(admin, "collection_restored", archive_id=archive_id, collection=collection)
    return {
        "ok": True,
        "collection": {"id": collection, "label": item["label"], "files": item["files"]},
        "reindex": reindex,
    }


@app.delete("/trash/collections/{archive_id}")
def purge_trashed_collection(
    archive_id: str,
    payload: PurgeCollectionRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    archive_id = _safe_archive_id(archive_id)
    archive = TRASH_ROOT / "collections" / archive_id
    if not archive.is_dir():
        raise HTTPException(status_code=404, detail="trash_collection_not_found")
    item = _trash_manifest(archive)
    if payload.confirm_id.strip() != item["collection"]:
        raise HTTPException(status_code=400, detail="collection_confirmation_mismatch")
    shutil.rmtree(archive)
    _audit(admin, "collection_purged", **item)
    return {"ok": True, "archive_id": archive_id, "collection": item["collection"]}


@app.post("/maintenance/trash-cleanup", include_in_schema=False)
def maintenance_trash_cleanup(
    request: Request,
    dry_run: bool = False,
) -> dict[str, Any]:
    actor = require_maintenance(request)
    return _purge_expired_collections(dry_run=dry_run, actor=actor)

@app.get("/collections/{collection}/documents")
def list_documents(
    collection: str,
    query: str = "",
    extension: str = "",
    expiry: str = "",
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    collection = _safe_collection(collection)
    state = _load_state()
    root = KB_ROOT / collection
    documents: list[dict[str, Any]] = []
    if root.exists():
        for path in root.rglob("*"):
            if (
                not path.is_file()
                or path.name.startswith((".", "~$"))
                or path.suffix.lower() not in SUPPORTED_EXTENSIONS
            ):
                continue
            item = _document_summary(path, collection, state)
            haystack = " ".join(
                [
                    item["name"],
                    item["path"],
                    item["title"],
                    item["owner"],
                    " ".join(item["tags"] if isinstance(item["tags"], list) else []),
                ]
            ).lower()
            if query and query.lower() not in haystack:
                continue
            if extension and item["extension"] != extension.lower().lstrip("."):
                continue
            if expiry and item["expiry_status"] != expiry:
                continue
            documents.append(item)
    documents.sort(key=lambda item: item["modified_at"], reverse=True)
    return {"collection": collection, "documents": documents, "count": len(documents)}


@app.get("/collections/{collection}/documents/{relative_path:path}")
def get_document(
    collection: str,
    relative_path: str,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    path = _document_path(collection, relative_path)
    state = _load_state()
    summary = _document_summary(path, collection, state)
    preview = _extract_preview(path)
    return {
        **summary,
        "editable": path.suffix.lower() in EDITABLE_EXTENSIONS,
        "content": preview if path.suffix.lower() in EDITABLE_EXTENSIONS else "",
        "preview": preview[:200_000],
    }


@app.put("/collections/{collection}/documents/{relative_path:path}")
def save_document(
    collection: str,
    relative_path: str,
    payload: SaveDocumentRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    path = _document_path(collection, relative_path)
    if path.suffix.lower() not in EDITABLE_EXTENSIONS:
        raise HTTPException(status_code=400, detail="binary_document_replace_via_upload")
    version_id = _create_version(path, collection, relative_path, admin, "save")
    _atomic_write(path, payload.content.encode("utf-8"))
    _audit(admin, "document_saved", collection=collection, path=relative_path, version_id=version_id)
    reindex = _trigger_reindex(collection, relative_path)
    return {
        "ok": True,
        "version_id": version_id,
        "reindex": reindex,
        "document": _document_summary(path, collection, _load_state()),
    }


@app.post("/upload")
async def upload_document(
    collection: str = Form(...),
    target_path: str = Form(""),
    file: UploadFile = File(...),
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    collection = _safe_collection(collection)
    filename = Path(file.filename or "").name
    relative = target_path.strip() or filename
    destination = _document_path(collection, relative, must_exist=False)
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="upload_too_large")
    version_id = ""
    if destination.exists():
        version_id = _create_version(destination, collection, relative, admin, "replace_upload")
    _atomic_write(destination, data)
    _audit(
        admin,
        "document_uploaded",
        collection=collection,
        path=relative,
        size=len(data),
        version_id=version_id,
    )
    reindex = _trigger_reindex(collection, relative)
    return {
        "ok": True,
        "version_id": version_id,
        "reindex": reindex,
        "document": _document_summary(destination, collection, _load_state()),
    }


@app.post("/collections/{collection}/documents/{relative_path:path}/move")
def move_document(
    collection: str,
    relative_path: str,
    payload: MoveDocumentRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    source = _document_path(collection, relative_path)
    target_collection = _safe_collection(payload.target_collection)
    target = _document_path(target_collection, payload.target_path, must_exist=False)
    if target.exists():
        raise HTTPException(status_code=409, detail="target_already_exists")
    _create_version(source, collection, relative_path, admin, "move")
    target.parent.mkdir(parents=True, exist_ok=True)
    source.replace(target)
    _audit(
        admin,
        "document_moved",
        collection=collection,
        path=relative_path,
        target_collection=target_collection,
        target_path=payload.target_path,
    )
    source_reindex = _trigger_reindex(collection, relative_path)
    target_reindex = _trigger_reindex(target_collection, payload.target_path)
    return {"ok": True, "collection": target_collection, "path": payload.target_path, "reindex": {"source": source_reindex, "target": target_reindex}}


@app.delete("/collections/{collection}/documents/{relative_path:path}")
def delete_document(
    collection: str,
    relative_path: str,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    source = _document_path(collection, relative_path)
    version_id = _create_version(source, collection, relative_path, admin, "delete")
    trash_path = (
        TRASH_ROOT
        / datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        / collection
        / Path(*_safe_relative_path(relative_path).parts)
    )
    trash_path.parent.mkdir(parents=True, exist_ok=True)
    source.replace(trash_path)
    _audit(
        admin,
        "document_deleted",
        collection=collection,
        path=relative_path,
        trash_path=trash_path.relative_to(KB_ROOT).as_posix(),
        version_id=version_id,
    )
    reindex = _trigger_reindex(collection, relative_path)
    return {"ok": True, "recoverable": True, "version_id": version_id, "reindex": reindex}


@app.get("/collections/{collection}/chunks/{relative_path:path}")
def document_chunks(
    collection: str,
    relative_path: str,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    _document_path(collection, relative_path)
    doc_id = f"{collection}/{_safe_relative_path(relative_path).as_posix()}"
    result = _qdrant(
        "POST",
        f"/collections/{collection}/points/scroll",
        json={
            "limit": 256,
            "with_payload": True,
            "with_vector": False,
            "filter": {"must": [{"key": "doc_id", "match": {"value": doc_id}}]},
        },
    ).get("result") or {}
    chunks = []
    for point in result.get("points") or []:
        payload = point.get("payload") or {}
        chunks.append(
            {
                "id": point.get("id"),
                "index": int(payload.get("chunk_index") or 0),
                "content": payload.get("content") or payload.get("text") or "",
            }
        )
    chunks.sort(key=lambda item: item["index"])
    return {"doc_id": doc_id, "chunks": chunks, "count": len(chunks)}


@app.get("/collections/{collection}/versions/{relative_path:path}")
def document_versions(
    collection: str,
    relative_path: str,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    _safe_collection(collection)
    _safe_relative_path(relative_path)
    version_dir = _version_key(collection, relative_path)
    versions = []
    if version_dir.exists():
        for meta_path in version_dir.glob("*.json"):
            try:
                versions.append(json.loads(meta_path.read_text(encoding="utf-8")))
            except Exception:
                continue
    versions.sort(key=lambda item: item.get("created_at") or "", reverse=True)
    return {"versions": versions}


@app.post("/collections/{collection}/restore/{relative_path:path}")
def restore_version(
    collection: str,
    relative_path: str,
    payload: RestoreVersionRequest,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    destination = _document_path(collection, relative_path, must_exist=False)
    version_dir = _version_key(collection, relative_path)
    meta_path = version_dir / f"{payload.version_id}.json"
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="version_not_found")
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    snapshot = version_dir / str(meta.get("snapshot") or "")
    if not snapshot.exists():
        raise HTTPException(status_code=404, detail="version_snapshot_missing")
    if destination.exists():
        _create_version(destination, collection, relative_path, admin, "before_restore")
    _atomic_write(destination, snapshot.read_bytes())
    _audit(
        admin,
        "version_restored",
        collection=collection,
        path=relative_path,
        version_id=payload.version_id,
    )
    reindex = _trigger_reindex(collection, relative_path)
    return {"ok": True, "reindex": reindex, "document": _document_summary(destination, collection, _load_state())}


@app.get("/search")
def semantic_search(
    query: str,
    limit: int = 12,
    admin: dict[str, Any] = Depends(require_unlocked_admin),
) -> dict[str, Any]:
    del admin
    clean_query = " ".join(query.split())
    if len(clean_query) < 2:
        return {"query": clean_query, "results": []}
    limit = max(1, min(limit, 30))
    if not IONOS_API_KEY:
        raise HTTPException(status_code=503, detail="embedding_api_not_configured")
    try:
        embedding_response = requests.post(
            f"{IONOS_BASE_URL}/embeddings",
            headers={"Authorization": f"Bearer {IONOS_API_KEY}"},
            json={"model": EMBEDDING_MODEL, "input": [clean_query]},
            timeout=120,
        )
        embedding_response.raise_for_status()
        vector = embedding_response.json()["data"][0]["embedding"]
    except Exception as exc:
        raise HTTPException(status_code=503, detail="embedding_request_failed") from exc
    results: list[dict[str, Any]] = []
    per_collection = max(4, min(limit, 12))
    for collection in _collection_names():
        payload = _qdrant(
            "POST",
            f"/collections/{collection}/points/search",
            json={
                "vector": vector,
                "limit": per_collection,
                "with_payload": True,
                "with_vector": False,
            },
        )
        for point in (payload.get("result") or []):
            source = point.get("payload") or {}
            results.append(
                {
                    "collection": collection,
                    "path": source.get("source_path") or "",
                    "doc_id": source.get("doc_id") or "",
                    "chunk_index": int(source.get("chunk_index") or 0),
                    "content": source.get("content") or source.get("text") or "",
                    "score": float(point.get("score") or 0),
                }
            )
    results.sort(key=lambda item: item["score"], reverse=True)
    return {"query": clean_query, "results": results[:limit]}
